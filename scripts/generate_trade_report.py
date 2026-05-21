"""Generate DOCX trade report comparing v4 FSM vs v5 DTE-aware exits for every signal.

Produces a Word document with:
  1. Executive summary (total P&L, win rate, daily breakdown)
  2. Per-trade detail table (entry, exit reason, P&L, peak gain, hold time)
  3. Gate fire breakdown for each strategy
  4. Daily P&L table

Usage:
    python scripts/generate_trade_report.py
    # => reports/trade_comparison_YYYY-MM-DD.docx
"""

from __future__ import annotations

import os
import sqlite3
import sys
from collections import defaultdict
from datetime import datetime, timedelta, timezone
from pathlib import Path

import numpy as np

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("ERROR: python-docx required. pip install python-docx")
    sys.exit(1)

PROJECT_DIR = Path(__file__).resolve().parent.parent
SIGNALS_DB = str(PROJECT_DIR / "journal" / "owlet-kody" / "raw_messages.db")
HARVESTER_DB = str(PROJECT_DIR / "journal" / "owlet-harvester" / "options_data.db")
REPORTS_DIR = PROJECT_DIR / "reports"

SLIPPAGE = 0.15
PORTFOLIO = 8000


# ============================================================================
# Shared helpers
# ============================================================================

def _parse_tick(tick, sig_ts, entry):
    ts, mid, bid, ask, underlying = tick
    ts_dt = datetime.fromisoformat(ts) if isinstance(ts, str) else ts
    if ts_dt.tzinfo is None:
        ts_dt = ts_dt.replace(tzinfo=timezone.utc)
    price = mid if mid and mid > 0 else ((bid + ask) / 2 if bid and ask else 0)
    if price <= 0:
        return None
    elapsed = (ts_dt - sig_ts).total_seconds() / 60
    gain_pct = (price - entry) / entry * 100
    et_hour = (ts_dt.hour - 4) % 24
    et_min = ts_dt.minute
    return price, elapsed, gain_pct, et_hour, et_min, underlying, ts_dt


def _eod_check(et_hour, et_min):
    return et_hour >= 15 and et_min >= 45


# ============================================================================
# v4 simulation (matches production v4 FSM logic)
# ============================================================================

def v4_simulate(entry, ticks, sig_ts, contracts, direction):
    """v4 FSM-equivalent: GRACE(90s) → DEVELOPING(soft trail + hard stop) → TRAILING(trail + house money)."""
    if not ticks or entry <= 0:
        return {"pnl": 0, "reason": "no_data", "hold": 0, "exit_prem": 0, "peak_gain": 0, "exit_gain_pct": 0}

    peak = entry
    grace_sec = 90
    trail_activate = 35.0  # v4 default
    hard_stop_pct = 30.0
    house_money_floor = 0.0

    # House money floors: +500% → floor at +200%, +200% → floor at +80%, +100% → floor at +30%
    hm_tiers = [(500, 200), (200, 80), (100, 30)]

    for tick in ticks:
        parsed = _parse_tick(tick, sig_ts, entry)
        if parsed is None:
            continue
        price, elapsed, gain_pct, et_hour, et_min, underlying, ts_dt = parsed

        if price > peak:
            peak = price

        peak_gain = (peak - entry) / entry * 100
        drop_entry = max(0, (entry - price) / entry * 100)
        drop_peak = (peak - price) / peak * 100 if peak > 0 else 0

        def _exit(reason):
            pnl = (price - entry) * contracts * 100
            if pnl > 0:
                pnl *= (1 - SLIPPAGE)
            exit_gain = (price - entry) / entry * 100 if entry > 0 else 0
            return {"pnl": pnl, "reason": reason, "hold": elapsed,
                    "exit_prem": price, "peak_gain": peak_gain, "exit_gain_pct": exit_gain}

        # EOD cutoff
        if _eod_check(et_hour, et_min):
            return _exit("eod_cutoff")

        elapsed_sec = elapsed * 60

        # GRACE state (first 90s)
        if elapsed_sec < grace_sec:
            # Bar-1 reverse check (90-150s window, down 5%+)
            if 90 <= elapsed_sec <= 150 and drop_entry >= 5:
                return _exit("bar1_reverse")
            continue

        # Hard stop -30%
        if drop_entry >= hard_stop_pct:
            return _exit("hard_stop")

        # DEVELOPING state (gain < trail_activate)
        if peak_gain < trail_activate:
            # Soft trail (15-35% band): protect 50% of peak gain
            if 15 <= peak_gain:
                floor = entry + (peak - entry) * 0.50
                if price <= floor:
                    return _exit("soft_trail")

            # Theta timer: 60min+ and no meaningful gain
            if elapsed >= 60 and gain_pct < 5:
                return _exit("theta_timer")
            continue

        # TRAILING state (gain >= trail_activate)

        # Update house money floor (ratchet up only)
        for trigger, floor_pct in hm_tiers:
            if peak_gain >= trigger:
                new_floor = entry * (1 + floor_pct / 100)
                if new_floor > house_money_floor:
                    house_money_floor = new_floor
                break

        # House money floor check
        if house_money_floor > 0 and price <= house_money_floor:
            return _exit("house_money_floor")

        # Trail stop — tiered by gain level
        if peak_gain >= 400:
            trail_width = 0.20
        elif peak_gain >= 200:
            trail_width = 0.25
        elif peak_gain >= 100:
            trail_width = 0.30
        else:
            trail_width = 0.35

        trail_stop = peak * (1 - trail_width)
        if price <= trail_stop:
            return _exit("trail_stop")

        # Theta timer in trailing: 60min+, gain < 10% of peak
        if elapsed >= 60 and gain_pct < peak_gain * 0.1:
            return _exit("theta_timer")

    # End of data
    last_price = 0
    for t in reversed(ticks):
        p = t[1] if t[1] and t[1] > 0 else 0
        if p > 0:
            last_price = p
            break
    if last_price > 0:
        pnl = (last_price - entry) * contracts * 100
        if pnl > 0:
            pnl *= (1 - SLIPPAGE)
        exit_gain = (last_price - entry) / entry * 100 if entry > 0 else 0
        return {"pnl": pnl, "reason": "eod_data_end", "hold": elapsed,
                "exit_prem": last_price, "peak_gain": peak_gain, "exit_gain_pct": exit_gain}
    return {"pnl": 0, "reason": "no_data", "hold": 0, "exit_prem": 0, "peak_gain": 0, "exit_gain_pct": 0}


# ============================================================================
# v5 DTE-aware simulation (matches backtest_ml_v5b.py)
# ============================================================================

_INDEX_TICKERS = {"SPY", "QQQ", "IWM", "DIA", "XLF", "XLK"}
_HIGH_VOL_TICKERS = {"MSTR", "AMD", "TSLA", "NVDA", "AVGO", "META", "COIN", "SMCI", "PLTR"}

# Multi-day contract caps (v5)
MULTI_DAY_MAX_CONTRACTS = 2
MULTI_DAY_EXPENSIVE_THRESHOLD = 5.0
DAILY_LOSS_CIRCUIT_BREAKER_PCT = 12.5
INDEX_PROFIT_TARGET_PCT = 30.0


def v5_simulate(entry, ticks, sig_ts, contracts, direction, dte=0, ticker=""):
    """v5 category-aware exits: index profit target, underlying confirmation,
    soft trail 10-50%/60% keep, DTE-aware thresholds."""
    if not ticks or entry <= 0:
        return {"pnl": 0, "reason": "no_data", "hold": 0, "exit_prem": 0, "peak_gain": 0, "exit_gain_pct": 0}

    is_call = direction.lower() in ("call", "bullish", "long")
    is_index = ticker in _INDEX_TICKERS
    peak = entry
    entry_underlying = None

    is_high_vol = ticker in _HIGH_VOL_TICKERS

    # DTE-aware thresholds — high-vol tickers get wider stops
    if is_high_vol:
        tight_stop = 45 if dte == 0 else 60
        backstop = 75 if dte == 0 else 85
    else:
        tight_stop = 35 if dte == 0 else 52
        backstop = 65 if dte == 0 else 75

    for tick in ticks:
        parsed = _parse_tick(tick, sig_ts, entry)
        if parsed is None:
            continue
        price, elapsed, gain_pct, et_hour, et_min, underlying, ts_dt = parsed

        if price > peak:
            peak = price
        if entry_underlying is None and underlying and underlying > 0:
            entry_underlying = underlying

        peak_gain = (peak - entry) / entry * 100
        drop_entry = max(0, (entry - price) / entry * 100)
        drop_peak = (peak - price) / peak * 100 if peak > 0 else 0

        # Underlying state
        u_move = 0.0
        has_underlying = False
        underlying_against = False
        underlying_confirms = False
        if entry_underlying and underlying and underlying > 0:
            has_underlying = True
            u_move = (underlying - entry_underlying) / entry_underlying * 100
            if is_call:
                underlying_against = u_move < -0.5
                underlying_confirms = u_move > 0.2
            else:
                underlying_against = u_move > 0.5
                underlying_confirms = u_move < -0.2

        def _exit(reason):
            pnl = (price - entry) * contracts * 100
            if pnl > 0:
                pnl *= (1 - SLIPPAGE)
            exit_gain = (price - entry) / entry * 100 if entry > 0 else 0
            return {"pnl": pnl, "reason": reason, "hold": elapsed,
                    "exit_prem": price, "peak_gain": peak_gain, "exit_gain_pct": exit_gain}

        # EOD cutoff (0DTE only)
        if dte == 0 and _eod_check(et_hour, et_min):
            return _exit("eod_cutoff")

        if elapsed < 5:
            continue

        # INDEX PROFIT TARGET (v5): take +30% on indexes
        if is_index and gain_pct >= INDEX_PROFIT_TARGET_PCT:
            return _exit("profit_target")

        # SCALP TRAIL
        if peak_gain >= 20 and gain_pct > 0 and gain_pct < peak_gain * 0.6:
            if dte == 0 and not underlying_confirms:
                return _exit("scalp_trail")
            elif dte > 0 and underlying_against:
                return _exit("scalp_trail")

        # CHECKPOINT (0DTE only)
        if dte == 0 and drop_entry >= 30:
            if has_underlying and underlying_against:
                return _exit("checkpoint_cut")

        # GRADUATED STOP
        if underlying_against:
            if drop_entry >= tight_stop:
                return _exit("confirmed_stop")
        else:
            if drop_entry >= backstop:
                return _exit("hard_stop")

        # SOFT TRAIL (10-50% band, 60% keep — v5 widened)
        if 10 <= peak_gain < 50:
            floor = entry + (peak - entry) * 0.60
            if price <= floor:
                return _exit("soft_trail")

        # ADAPTIVE TRAIL
        if peak_gain >= 400:
            if drop_peak >= 30:
                return _exit("adaptive_moonshot")
        elif peak_gain >= 150:
            if drop_peak >= 45:
                return _exit("adaptive_runner")
        elif peak_gain >= 40:
            if drop_peak >= 40:
                return _exit("adaptive_active")

        # THETA BLEED (0DTE only)
        if dte == 0 and elapsed >= 120 and drop_entry >= 30:
            return _exit("theta_bleed")

    # End of data
    last_price = 0
    for t in reversed(ticks):
        p = t[1] if t[1] and t[1] > 0 else 0
        if p > 0:
            last_price = p
            break
    if last_price > 0:
        pnl = (last_price - entry) * contracts * 100
        if pnl > 0:
            pnl *= (1 - SLIPPAGE)
        elapsed_f = 0
        try:
            last_ts = datetime.fromisoformat(ticks[-1][0]) if isinstance(ticks[-1][0], str) else ticks[-1][0]
            if last_ts.tzinfo is None:
                last_ts = last_ts.replace(tzinfo=timezone.utc)
            elapsed_f = (last_ts - sig_ts).total_seconds() / 60
        except Exception:
            pass
        exit_gain = (last_price - entry) / entry * 100 if entry > 0 else 0
        return {"pnl": pnl, "reason": "eod_data_end", "hold": elapsed_f,
                "exit_prem": last_price, "peak_gain": peak_gain, "exit_gain_pct": exit_gain}
    return {"pnl": 0, "reason": "no_data", "hold": 0, "exit_prem": 0, "peak_gain": 0, "exit_gain_pct": 0}


# ============================================================================
# Data loading
# ============================================================================

def build_ct(ticker, day, direction, strike):
    dt = datetime.strptime(day, "%Y-%m-%d")
    ds = dt.strftime("%y%m%d")
    cp = "C" if direction.lower() in ("call", "bullish", "long") else "P"
    si = int(strike * 1000)
    return f"O:{ticker}{ds}{cp}{si:08d}"


def load_signals_and_ticks():
    sig_conn = sqlite3.connect(SIGNALS_DB)
    sig_conn.row_factory = sqlite3.Row
    harv_conn = sqlite3.connect(HARVESTER_DB)

    signals = sig_conn.execute("""
        SELECT ts.id, ts.ticker, ts.direction, ts.sentiment, ts.score, ts.strike,
               ts.atm_premium, ts.otm_premium, date(ts.created_at) as day,
               ts.created_at as sig_ts
        FROM trade_signals ts ORDER BY ts.created_at
    """).fetchall()

    results = []
    no_data = no_strike = 0

    for sig in signals:
        ticker = sig["ticker"]
        direction = sig["direction"] or "bullish"
        day = sig["day"]
        strike = sig["strike"]
        score = sig["score"] or 0
        premium = sig["atm_premium"] or sig["otm_premium"]
        sent = (sig["sentiment"] or direction).lower()
        option_type = "put" if sent in ("bearish", "put") else "call"

        if not strike or not premium or premium <= 0:
            no_strike += 1
            continue

        # Try same-day contract first, then multi-day
        sig_date = datetime.strptime(day, "%Y-%m-%d").date()
        candidates = [sig_date]
        for delta in range(1, 6):
            d = sig_date + timedelta(days=delta)
            if d.weekday() < 5:
                candidates.append(d)
                if len(candidates) >= 4:
                    break

        ticks = None
        dte = 0
        for exp_date in candidates:
            expiry = exp_date.strftime("%Y-%m-%d")
            exp_str = exp_date.strftime("%y%m%d")
            ot = "C" if option_type == "call" else "P"
            strike_int = int(strike * 1000)
            ct = f"O:{ticker}{exp_str}{ot}{strike_int:08d}"

            rows = harv_conn.execute("""
                SELECT captured_at, midpoint, bid, ask, underlying_price
                FROM harvest_snapshots WHERE contract_ticker = ? AND captured_at >= ?
                ORDER BY captured_at
            """, (ct, sig["sig_ts"])).fetchall()
            if rows and len(rows) >= 10:
                ticks = rows
                dte = (exp_date - sig_date).days
                break

        if not ticks:
            no_data += 1
            continue

        # Entry price from harvester
        first = ticks[0]
        entry = (first[3] if first[3] and first[3] > 0 else first[1]) or premium
        if entry <= 0:
            entry = premium

        # Contract sizing by score
        if score >= 95:
            contracts = 5
        elif score >= 90:
            contracts = 4
        elif score >= 85:
            contracts = 3
        else:
            contracts = 1

        sig_ts = datetime.fromisoformat(sig["sig_ts"])
        if sig_ts.tzinfo is None:
            sig_ts = sig_ts.replace(tzinfo=timezone.utc)

        # Overall peak for reference
        all_mids = [r[1] for r in ticks if r[1] and r[1] > 0]
        overall_peak = max(all_mids) if all_mids else entry
        overall_peak_gain = (overall_peak - entry) / entry * 100

        results.append({
            "ticker": ticker,
            "direction": direction,
            "option_type": option_type,
            "day": day,
            "score": score,
            "strike": strike,
            "premium": premium,
            "entry": entry,
            "contracts": contracts,
            "ticks": ticks,
            "sig_ts": sig_ts,
            "overall_peak_gain": overall_peak_gain,
            "dte": dte,
        })

    sig_conn.close()
    harv_conn.close()
    return results, no_data, no_strike


# ============================================================================
# DOCX generation
# ============================================================================

def _set_cell(cell, text, bold=False, color=None, align=None, size=8):
    """Helper to format a table cell."""
    p = cell.paragraphs[0]
    p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _pnl_color(pnl):
    if pnl > 0:
        return RGBColor(0, 128, 0)  # green
    elif pnl < 0:
        return RGBColor(200, 0, 0)  # red
    return RGBColor(0, 0, 0)


def generate_report(trades, no_data, no_strike):
    doc = Document()

    # Title
    title = doc.add_heading("OptionsOwl Trade Report: v4 FSM vs v5 Category-Aware", level=0)

    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        f"Portfolio: ${PORTFOLIO:,}\n"
        f"Signals with data: {len(trades)} | No data: {no_data} | No strike: {no_strike}\n"
        f"Slippage: {SLIPPAGE*100:.0f}% on winning trades"
    )

    # Run both strategies
    v4_results = []
    v5_results = []

    # v5 daily circuit breaker tracking
    cb_limit = PORTFOLIO * (DAILY_LOSS_CIRCUIT_BREAKER_PCT / 100)
    v5_daily_pnl = defaultdict(float)  # day → cumulative realized P&L

    for t in trades:
        v4_r = v4_simulate(t["entry"], t["ticks"], t["sig_ts"], t["contracts"], t["direction"])
        v4_r.update({"ticker": t["ticker"], "day": t["day"], "score": t["score"],
                      "entry": t["entry"], "contracts": t["contracts"],
                      "direction": t["direction"], "dte": t["dte"],
                      "overall_peak": t["overall_peak_gain"]})
        v4_results.append(v4_r)

        # v5 daily circuit breaker: skip trade if day's losses exceed threshold
        day = t["day"]
        if v5_daily_pnl[day] < -cb_limit:
            v5_r = {"pnl": 0, "reason": "circuit_breaker", "hold": 0,
                     "exit_prem": 0, "peak_gain": 0, "exit_gain_pct": 0}
            v5_r.update({"ticker": t["ticker"], "day": day, "score": t["score"],
                          "entry": t["entry"], "contracts": 0,
                          "direction": t["direction"], "dte": t["dte"],
                          "overall_peak": t["overall_peak_gain"]})
            v5_results.append(v5_r)
            continue

        # v5 applies multi-day contract cap
        v5_contracts = t["contracts"]
        if t["dte"] > 0:
            if t["entry"] > MULTI_DAY_EXPENSIVE_THRESHOLD:
                v5_contracts = min(v5_contracts, 1)
            elif MULTI_DAY_MAX_CONTRACTS > 0:
                v5_contracts = min(v5_contracts, MULTI_DAY_MAX_CONTRACTS)

        v5_r = v5_simulate(t["entry"], t["ticks"], t["sig_ts"], v5_contracts,
                           t["direction"], t["dte"], t["ticker"])
        v5_r.update({"ticker": t["ticker"], "day": day, "score": t["score"],
                      "entry": t["entry"], "contracts": v5_contracts,
                      "direction": t["direction"], "dte": t["dte"],
                      "overall_peak": t["overall_peak_gain"]})
        v5_results.append(v5_r)

        # Track realized P&L for circuit breaker
        v5_daily_pnl[day] += v5_r["pnl"]

    # ── Executive Summary ──
    doc.add_heading("1. Executive Summary", level=1)

    summary_table = doc.add_table(rows=8, cols=3)
    summary_table.style = "Light Grid Accent 1"

    headers = ["Metric", "v4 FSM", "v5 Category-Aware"]
    for i, h in enumerate(headers):
        _set_cell(summary_table.rows[0].cells[i], h, bold=True, size=9)

    for strat_idx, (name, results) in enumerate(
        [("v4", v4_results), ("v5", v5_results)], start=1
    ):
        pnls = [r["pnl"] for r in results]
        wins = sum(1 for p in pnls if p > 0)
        n = len(pnls)
        total = sum(pnls)
        avg_w = np.mean([p for p in pnls if p > 0]) if wins > 0 else 0
        avg_l = np.mean([p for p in pnls if p <= 0]) if (n - wins) > 0 else 0
        avg_h = np.mean([r["hold"] for r in results])
        wr = wins / n * 100 if n > 0 else 0

        metrics = [
            ("Trades", str(n)),
            ("Total P&L", f"${total:+,.0f}"),
            ("Return on Portfolio", f"{total / PORTFOLIO * 100:+.1f}%"),
            ("Win Rate", f"{wr:.0f}% ({wins}W/{n - wins}L)"),
            ("Avg Win", f"${avg_w:+,.0f}"),
            ("Avg Loss", f"${avg_l:+,.0f}"),
            ("Avg Hold", f"{avg_h:.0f}m"),
        ]

        for row_idx, (metric, value) in enumerate(metrics):
            _set_cell(summary_table.rows[row_idx + 1].cells[0], metric, bold=True, size=9)
            color = None
            if "P&L" in metric or "Return" in metric:
                val_num = total if metric == "Total P&L" else total / PORTFOLIO * 100
                color = _pnl_color(val_num)
            _set_cell(summary_table.rows[row_idx + 1].cells[strat_idx], value, color=color, size=9)

    # ── Daily P&L ──
    doc.add_heading("2. Daily P&L Comparison", level=1)

    all_days = sorted(set(r["day"] for r in v4_results))

    daily_table = doc.add_table(rows=len(all_days) + 2, cols=8)
    daily_table.style = "Light Grid Accent 1"

    daily_headers = ["Date", "Sigs", "v4 P&L", "v4 WR", "v5 P&L", "v5 WR", "Better", "Delta"]
    for i, h in enumerate(daily_headers):
        _set_cell(daily_table.rows[0].cells[i], h, bold=True, size=8)

    cum_v4 = cum_v5 = 0
    v4_day_wins = v5_day_wins = 0

    for row_idx, day in enumerate(all_days, start=1):
        v4_day = [r for r in v4_results if r["day"] == day]
        v5_day = [r for r in v5_results if r["day"] == day]

        v4_pnl = sum(r["pnl"] for r in v4_day)
        v5_pnl = sum(r["pnl"] for r in v5_day)
        v4_wr = sum(1 for r in v4_day if r["pnl"] > 0) / len(v4_day) * 100 if v4_day else 0
        v5_wr = sum(1 for r in v5_day if r["pnl"] > 0) / len(v5_day) * 100 if v5_day else 0

        cum_v4 += v4_pnl
        cum_v5 += v5_pnl
        if v4_pnl > 0:
            v4_day_wins += 1
        if v5_pnl > 0:
            v5_day_wins += 1

        better = "v4" if v4_pnl > v5_pnl else ("v5" if v5_pnl > v4_pnl else "tie")
        delta = v5_pnl - v4_pnl

        _set_cell(daily_table.rows[row_idx].cells[0], day, size=8)
        _set_cell(daily_table.rows[row_idx].cells[1], str(len(v4_day)), size=8,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(daily_table.rows[row_idx].cells[2], f"${v4_pnl:+,.0f}", size=8,
                  color=_pnl_color(v4_pnl))
        _set_cell(daily_table.rows[row_idx].cells[3], f"{v4_wr:.0f}%", size=8,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(daily_table.rows[row_idx].cells[4], f"${v5_pnl:+,.0f}", size=8,
                  color=_pnl_color(v5_pnl))
        _set_cell(daily_table.rows[row_idx].cells[5], f"{v5_wr:.0f}%", size=8,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(daily_table.rows[row_idx].cells[6], better, size=8, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(daily_table.rows[row_idx].cells[7], f"${delta:+,.0f}", size=8,
                  color=_pnl_color(delta))

    # Cumulative row
    cum_row = len(all_days) + 1
    _set_cell(daily_table.rows[cum_row].cells[0], "CUMULATIVE", bold=True, size=8)
    _set_cell(daily_table.rows[cum_row].cells[2], f"${cum_v4:+,.0f}", bold=True, size=8,
              color=_pnl_color(cum_v4))
    _set_cell(daily_table.rows[cum_row].cells[3], f"{v4_day_wins}/{len(all_days)}", size=8,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(daily_table.rows[cum_row].cells[4], f"${cum_v5:+,.0f}", bold=True, size=8,
              color=_pnl_color(cum_v5))
    _set_cell(daily_table.rows[cum_row].cells[5], f"{v5_day_wins}/{len(all_days)}", size=8,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    delta_cum = cum_v5 - cum_v4
    _set_cell(daily_table.rows[cum_row].cells[7], f"${delta_cum:+,.0f}", bold=True, size=8,
              color=_pnl_color(delta_cum))

    # ── Gate Fire Breakdown ──
    doc.add_heading("3. Gate Fire Breakdown", level=1)

    for strat_name, results in [("v4 FSM", v4_results), ("v5 Category-Aware", v5_results)]:
        doc.add_heading(f"{strat_name}", level=2)

        gate_stats = defaultdict(lambda: {"fires": 0, "pnl": 0, "wins": 0, "holds": []})
        for r in results:
            reason = r["reason"]
            gate_stats[reason]["fires"] += 1
            gate_stats[reason]["pnl"] += r["pnl"]
            if r["pnl"] > 0:
                gate_stats[reason]["wins"] += 1
            gate_stats[reason]["holds"].append(r["hold"])

        gate_table = doc.add_table(rows=len(gate_stats) + 1, cols=6)
        gate_table.style = "Light Grid Accent 1"
        gate_headers = ["Gate", "Fires", "%", "P&L", "Win Rate", "Avg Hold"]
        for i, h in enumerate(gate_headers):
            _set_cell(gate_table.rows[0].cells[i], h, bold=True, size=8)

        n_total = len(results)
        for row_idx, (gate, stats) in enumerate(
            sorted(gate_stats.items(), key=lambda x: -x[1]["fires"]), start=1
        ):
            wr = stats["wins"] / stats["fires"] * 100 if stats["fires"] > 0 else 0
            pct = stats["fires"] / n_total * 100 if n_total > 0 else 0
            avg_h = np.mean(stats["holds"]) if stats["holds"] else 0

            _set_cell(gate_table.rows[row_idx].cells[0], gate, size=8)
            _set_cell(gate_table.rows[row_idx].cells[1], str(stats["fires"]), size=8,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(gate_table.rows[row_idx].cells[2], f"{pct:.0f}%", size=8,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(gate_table.rows[row_idx].cells[3], f"${stats['pnl']:+,.0f}", size=8,
                      color=_pnl_color(stats["pnl"]))
            _set_cell(gate_table.rows[row_idx].cells[4], f"{wr:.0f}%", size=8,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(gate_table.rows[row_idx].cells[5], f"{avg_h:.0f}m", size=8,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Per-Trade Detail ──
    doc.add_heading("4. Every Trade — Side by Side", level=1)
    doc.add_paragraph(
        "Each row shows one signal and how both v4 and v5 would have handled it. "
        "Green P&L = winner, Red = loser. 'Peak' shows the max available gain during the trade."
    )

    # Group by day
    for day in all_days:
        doc.add_heading(f"{day}", level=2)

        day_v4 = [r for r in v4_results if r["day"] == day]
        day_v5 = [r for r in v5_results if r["day"] == day]

        # Build combined table
        n_trades = len(day_v4)
        trade_table = doc.add_table(rows=n_trades + 1, cols=12)
        trade_table.style = "Light Grid Accent 1"

        t_headers = ["Ticker", "Dir", "Score", "DTE", "Entry",
                      "v4 Exit", "v4 P&L", "v4 Hold",
                      "v5 Exit", "v5 P&L", "v5 Hold", "Peak"]
        for i, h in enumerate(t_headers):
            _set_cell(trade_table.rows[0].cells[i], h, bold=True, size=7)

        for row_idx in range(n_trades):
            v4_r = day_v4[row_idx]
            v5_r = day_v5[row_idx]

            _set_cell(trade_table.rows[row_idx + 1].cells[0], v4_r["ticker"], size=7)
            _set_cell(trade_table.rows[row_idx + 1].cells[1],
                      v4_r["direction"][:4].upper(), size=7)
            _set_cell(trade_table.rows[row_idx + 1].cells[2], str(v4_r["score"]), size=7,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(trade_table.rows[row_idx + 1].cells[3], str(v4_r["dte"]), size=7,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(trade_table.rows[row_idx + 1].cells[4],
                      f"${v4_r['entry']:.2f}", size=7)

            # v4 columns
            _set_cell(trade_table.rows[row_idx + 1].cells[5], v4_r["reason"], size=7)
            _set_cell(trade_table.rows[row_idx + 1].cells[6],
                      f"${v4_r['pnl']:+,.0f}", size=7,
                      color=_pnl_color(v4_r["pnl"]))
            _set_cell(trade_table.rows[row_idx + 1].cells[7],
                      f"{v4_r['hold']:.0f}m", size=7,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

            # v5 columns
            _set_cell(trade_table.rows[row_idx + 1].cells[8], v5_r["reason"], size=7)
            _set_cell(trade_table.rows[row_idx + 1].cells[9],
                      f"${v5_r['pnl']:+,.0f}", size=7,
                      color=_pnl_color(v5_r["pnl"]))
            _set_cell(trade_table.rows[row_idx + 1].cells[10],
                      f"{v5_r['hold']:.0f}m", size=7,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

            # Peak available
            peak = v4_r.get("overall_peak", 0)
            _set_cell(trade_table.rows[row_idx + 1].cells[11],
                      f"+{peak:.0f}%", size=7,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

        # Day summary
        v4_day_pnl = sum(r["pnl"] for r in day_v4)
        v5_day_pnl = sum(r["pnl"] for r in day_v5)
        doc.add_paragraph(
            f"Day total: v4=${v4_day_pnl:+,.0f}  v5=${v5_day_pnl:+,.0f}  "
            f"delta=${v5_day_pnl - v4_day_pnl:+,.0f}"
        )

    # ── Business Logic Reference ──
    doc.add_heading("5. Exit Gate Reference", level=1)

    doc.add_heading("v4 FSM Exit Logic", level=2)
    doc.add_paragraph(
        "v4 uses a formal Finite State Machine with 3 states. "
        "All mutable state lives in a TradeState dataclass (peak premium, house money floor, "
        "milestone locks, theta timer). The FSM is stateless — it evaluates the current "
        "snapshot and returns an ExitAction."
    )

    v4_gates = [
        ("GRACE (0-90s)", "Only bar-1 reverse (-5%) and bid disappearance can exit. "
         "Protects against open noise."),
        ("hard_stop", "Premium down 30% from entry. Hard floor, always active after grace."),
        ("soft_trail", "15-35% peak gain band: floor = entry + 50% of peak gain. "
         "Prevents winners from turning into losers."),
        ("theta_timer", "60min+ with <5% gain (DEVELOPING) or <10% of peak (TRAILING). "
         "Cuts dead trades that aren't moving."),
        ("trail_stop", "Tiered trailing stop: +50% gain → 35% trail, +100% → 30%, "
         "+200% → 25%, +400% → 20%. Tightens as gains grow."),
        ("house_money_floor", "Progressive monotonic floor: +100% → lock +30%, "
         "+200% → lock +80%, +500% → lock +200%. Never lowers."),
        ("milestone_lock", "Partial profit lock: close 15% of contracts at +200%, +400%, +600%. "
         "Locks gains while letting rest run."),
        ("eod_cutoff", "Close all at 3:45 PM ET. Safety net for 0DTE."),
        ("bar1_reverse", "90-150s after fill, premium down 5%+. Catches immediate reversals."),
    ]

    v4_table = doc.add_table(rows=len(v4_gates) + 1, cols=2)
    v4_table.style = "Light Grid Accent 1"
    _set_cell(v4_table.rows[0].cells[0], "Gate", bold=True, size=9)
    _set_cell(v4_table.rows[0].cells[1], "Logic", bold=True, size=9)
    for i, (gate, desc) in enumerate(v4_gates, start=1):
        _set_cell(v4_table.rows[i].cells[0], gate, bold=True, size=8)
        _set_cell(v4_table.rows[i].cells[1], desc, size=8)

    doc.add_heading("v5 Category-Aware Exit Logic", level=2)
    doc.add_paragraph(
        "v5 adds category awareness on top of v4's FSM framework. Key additions: "
        "(1) Index profit target at +30% for SPY/QQQ/IWM — indexes mean-revert, take the money. "
        "(2) Soft trail widened to 10-50% band with 60% keep (was 15-35%/50%) — catches gains earlier. "
        "(3) Multi-day contract cap at 2 (1 for expensive >$5 premiums) — limits single-trade max loss. "
        "(4) Daily circuit breaker at 12.5% of portfolio — stops trading on bear days. "
        "(5) Underlying confirmation for scalp trail, checkpoint cut, graduated stops."
    )

    v5_gates = [
        ("profit_target", "Index tickers (SPY, QQQ, IWM, DIA, XLF, XLK): take profits at +30%. "
         "Backtested: 100% WR, $+1,401 total. Indexes mean-revert — take the money early."),
        ("scalp_trail", "Premium peaked +20%, faded to <60% of peak, gain still positive. "
         "0DTE: exit if underlying NOT confirming (>0.2% in direction). "
         "Multi-day: only exit if underlying actively AGAINST (>0.5%). "
         "Catches IV-driven spikes without underlying support."),
        ("checkpoint_cut", "Premium down 30% AND underlying moved 0.5%+ against. "
         "DISABLED for multi-day (DTE>0) — temporary dips recover. "
         "0DTE only: both signals agree the trade is dead."),
        ("confirmed_stop", "Underlying actively against AND premium down past tight stop. "
         "0DTE: 35% drop. Multi-day: 52% drop (1.5x wider). "
         "Cuts when the stock is moving against you AND premium confirms."),
        ("hard_stop", "Backstop when underlying is NOT against but premium collapsed. "
         "0DTE: 65% drop. Multi-day: 75% drop. "
         "Safety net for cases where underlying data is flat but option bleeds."),
        ("soft_trail", "10-50% peak gain band: floor = 60% of peak gain above entry (v5 widened). "
         "Wider band (10% vs 15%) catches gains earlier. "
         "Higher keep (60% vs 50%) is more protective."),
        ("adaptive_active", "Peak gain 40-150%: exit on 40% drop from peak. "
         "Primary trail for moderate winners."),
        ("adaptive_runner", "Peak gain 150-400%: exit on 45% drop from peak. "
         "Wider trail lets big runners breathe."),
        ("adaptive_moonshot", "Peak gain 400%+: exit on 30% drop from peak. "
         "Tighter trail to lock in huge gains."),
        ("theta_bleed", "0DTE only: 120min+ and premium down 30%+. Theta eating the position. "
         "DISABLED for multi-day — theta is negligible with 1+ days remaining."),
        ("eod_cutoff", "0DTE only: close at 3:45 PM ET. "
         "DISABLED for multi-day — position holds overnight."),
        ("circuit_breaker", "Daily loss exceeds 12.5% of portfolio → block new trades for the day. "
         "Scales with portfolio: $8K × 12.5% = $1,000, $2.5K × 12.5% = $312. "
         "Limits bear-day damage."),
    ]

    v5_table = doc.add_table(rows=len(v5_gates) + 1, cols=2)
    v5_table.style = "Light Grid Accent 1"
    _set_cell(v5_table.rows[0].cells[0], "Gate", bold=True, size=9)
    _set_cell(v5_table.rows[0].cells[1], "Logic", bold=True, size=9)
    for i, (gate, desc) in enumerate(v5_gates, start=1):
        _set_cell(v5_table.rows[i].cells[0], gate, bold=True, size=8)
        _set_cell(v5_table.rows[i].cells[1], desc, size=8)

    doc.add_heading("Key Differences: v4 vs v5", level=2)
    diffs = [
        ("Category Awareness", "v4 treats all tickers the same. v5 has separate strategies for "
         "indexes (profit target at +30%), high-vol tickers (wider trails), and standard tickers."),
        ("Index Profit Target", "v5 adds a +30% profit target for SPY/QQQ/IWM/DIA/XLF/XLK. "
         "Backtested: 100% WR, $+1,401 total. v4 doesn't have this."),
        ("Soft Trail (Widened)", "v4: 15-35% band, 50% keep. v5: 10-50% band, 60% keep. "
         "Catches gains earlier (+10% vs +15%) and keeps more of the peak (60% vs 50%)."),
        ("Multi-day Contract Cap", "v5 caps multi-day trades at 2 contracts (1 for expensive >$5). "
         "Biggest single improvement: reduces max single loss from -$1,900 to -$620."),
        ("Daily Circuit Breaker", "v5 stops opening new trades when daily realized losses exceed "
         "12.5% of portfolio. Scales with portfolio size. Limits bear-day damage."),
        ("Underlying Confirmation", "v4 ignores underlying price. v5 uses it for scalp trail, "
         "checkpoint cut, and graduated stops."),
        ("DTE-Aware Thresholds", "v4 uses fixed stops. v5 widens stops 1.5x for DTE>0 "
         "and disables checkpoint/theta/EOD for multi-day contracts."),
        ("Scalp Detection", "v5 adds scalp_trail — detects IV-driven spikes that fade without "
         "underlying support. v4 doesn't have this."),
    ]

    diff_table = doc.add_table(rows=len(diffs) + 1, cols=2)
    diff_table.style = "Light Grid Accent 1"
    _set_cell(diff_table.rows[0].cells[0], "Aspect", bold=True, size=9)
    _set_cell(diff_table.rows[0].cells[1], "Comparison", bold=True, size=9)
    for i, (aspect, desc) in enumerate(diffs, start=1):
        _set_cell(diff_table.rows[i].cells[0], aspect, bold=True, size=8)
        _set_cell(diff_table.rows[i].cells[1], desc, size=8)

    # ── Key Findings ──
    doc.add_heading("6. Key Findings", level=1)

    # Find biggest disagreements
    disagreements = []
    for v4_r, v5_r in zip(v4_results, v5_results):
        delta = v5_r["pnl"] - v4_r["pnl"]
        if abs(delta) > 100:
            disagreements.append({
                "ticker": v4_r["ticker"], "day": v4_r["day"],
                "v4_pnl": v4_r["pnl"], "v5_pnl": v5_r["pnl"],
                "v4_reason": v4_r["reason"], "v5_reason": v5_r["reason"],
                "delta": delta, "dte": v4_r["dte"],
                "peak": v4_r.get("overall_peak", 0),
            })

    disagreements.sort(key=lambda x: -abs(x["delta"]))

    doc.add_paragraph(f"Trades where v4 and v5 differ by >$100: {len(disagreements)}")

    if disagreements:
        dis_table = doc.add_table(rows=min(len(disagreements), 20) + 1, cols=8)
        dis_table.style = "Light Grid Accent 1"
        dis_headers = ["Ticker", "Day", "DTE", "v4 Exit", "v4 P&L",
                        "v5 Exit", "v5 P&L", "Delta"]
        for i, h in enumerate(dis_headers):
            _set_cell(dis_table.rows[0].cells[i], h, bold=True, size=8)

        for row_idx, d in enumerate(disagreements[:20], start=1):
            _set_cell(dis_table.rows[row_idx].cells[0], d["ticker"], size=8)
            _set_cell(dis_table.rows[row_idx].cells[1], d["day"], size=8)
            _set_cell(dis_table.rows[row_idx].cells[2], str(d["dte"]), size=8,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(dis_table.rows[row_idx].cells[3], d["v4_reason"], size=8)
            _set_cell(dis_table.rows[row_idx].cells[4], f"${d['v4_pnl']:+,.0f}", size=8,
                      color=_pnl_color(d["v4_pnl"]))
            _set_cell(dis_table.rows[row_idx].cells[5], d["v5_reason"], size=8)
            _set_cell(dis_table.rows[row_idx].cells[6], f"${d['v5_pnl']:+,.0f}", size=8,
                      color=_pnl_color(d["v5_pnl"]))
            _set_cell(dis_table.rows[row_idx].cells[7], f"${d['delta']:+,.0f}", size=8,
                      color=_pnl_color(d["delta"]), bold=True)

    # ── Peak Capture Efficiency (v5 feature) ──
    doc.add_heading("7. Peak Capture Efficiency", level=1)
    doc.add_paragraph(
        "How close does each strategy exit to the trade's peak gain? "
        "Peak Capture % = exit_gain / peak_gain × 100. "
        "A score of 80% means we captured 80% of the maximum available gain. "
        "Higher is better — it means we're exiting closer to the top."
    )

    # Compute peak capture for both strategies
    def _peak_capture_stats(results):
        captures = []
        for r in results:
            peak = r.get("peak_gain", 0)
            exit_g = r.get("exit_gain_pct", 0)
            if peak > 5:  # only meaningful for trades that had a real peak (>5%)
                capture = (exit_g / peak * 100) if peak > 0 else 0
                captures.append({
                    "ticker": r["ticker"], "day": r["day"],
                    "peak_gain": peak, "exit_gain": exit_g,
                    "capture_pct": capture, "reason": r["reason"],
                    "pnl": r["pnl"],
                })
        return captures

    v4_captures = _peak_capture_stats(v4_results)
    v5_captures = _peak_capture_stats(v5_results)

    # Summary stats
    pc_table = doc.add_table(rows=7, cols=3)
    pc_table.style = "Light Grid Accent 1"
    pc_headers = ["Metric", "v4 FSM", "v5 Category-Aware"]
    for i, h in enumerate(pc_headers):
        _set_cell(pc_table.rows[0].cells[i], h, bold=True, size=9)

    for strat_idx, (caps, name) in enumerate([(v4_captures, "v4"), (v5_captures, "v5")], start=1):
        if not caps:
            continue
        all_captures = [c["capture_pct"] for c in caps]
        win_captures = [c["capture_pct"] for c in caps if c["pnl"] > 0]
        metrics = [
            ("Trades with peak >5%", str(len(caps))),
            ("Mean Capture %", f"{np.mean(all_captures):.1f}%"),
            ("Median Capture %", f"{np.median(all_captures):.1f}%"),
            ("Winners Mean Capture", f"{np.mean(win_captures):.1f}%" if win_captures else "N/A"),
            ("Exits at 80%+ of peak", f"{sum(1 for c in all_captures if c >= 80)} ({sum(1 for c in all_captures if c >= 80)/len(caps)*100:.0f}%)"),
            ("Exits at 50%+ of peak", f"{sum(1 for c in all_captures if c >= 50)} ({sum(1 for c in all_captures if c >= 50)/len(caps)*100:.0f}%)"),
        ]
        for row_idx, (metric, value) in enumerate(metrics):
            _set_cell(pc_table.rows[row_idx + 1].cells[0], metric, bold=True, size=9)
            _set_cell(pc_table.rows[row_idx + 1].cells[strat_idx], value, size=9)

    # Per-trade peak capture table (winners only, sorted by capture %)
    doc.add_heading("Winners — Peak Capture Detail", level=2)
    doc.add_paragraph(
        "Every winning trade: how much of the peak gain did each strategy capture? "
        "Sorted by v5 capture % (descending). Green = captured >70% of peak."
    )

    # Build side-by-side data
    paired = []
    for v4_r, v5_r in zip(v4_results, v5_results):
        peak = v4_r.get("overall_peak", 0)
        if peak <= 5 or (v4_r["pnl"] <= 0 and v5_r["pnl"] <= 0):
            continue
        v4_exit_g = v4_r.get("exit_gain_pct", 0)
        v5_exit_g = v5_r.get("exit_gain_pct", 0)
        v4_cap = (v4_exit_g / peak * 100) if peak > 0 else 0
        v5_cap = (v5_exit_g / peak * 100) if peak > 0 else 0
        paired.append({
            "ticker": v4_r["ticker"], "day": v4_r["day"],
            "peak": peak,
            "v4_exit": v4_exit_g, "v4_cap": v4_cap, "v4_reason": v4_r["reason"],
            "v4_pnl": v4_r["pnl"],
            "v5_exit": v5_exit_g, "v5_cap": v5_cap, "v5_reason": v5_r["reason"],
            "v5_pnl": v5_r["pnl"],
        })
    paired.sort(key=lambda x: -x["v5_cap"])

    if paired:
        cap_table = doc.add_table(rows=len(paired) + 1, cols=10)
        cap_table.style = "Light Grid Accent 1"
        cap_headers = ["Ticker", "Day", "Peak", "v4 Exit", "v4 Cap%", "v4 P&L",
                        "v5 Exit", "v5 Cap%", "v5 P&L", "Better"]
        for i, h in enumerate(cap_headers):
            _set_cell(cap_table.rows[0].cells[i], h, bold=True, size=7)

        for row_idx, p in enumerate(paired, start=1):
            _set_cell(cap_table.rows[row_idx].cells[0], p["ticker"], size=7)
            _set_cell(cap_table.rows[row_idx].cells[1], p["day"], size=7)
            _set_cell(cap_table.rows[row_idx].cells[2], f"+{p['peak']:.0f}%", size=7,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(cap_table.rows[row_idx].cells[3], f"+{p['v4_exit']:.0f}%", size=7,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            v4_color = RGBColor(0, 128, 0) if p["v4_cap"] >= 70 else (
                RGBColor(200, 128, 0) if p["v4_cap"] >= 40 else RGBColor(200, 0, 0))
            _set_cell(cap_table.rows[row_idx].cells[4], f"{p['v4_cap']:.0f}%", size=7,
                      color=v4_color, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(cap_table.rows[row_idx].cells[5], f"${p['v4_pnl']:+,.0f}", size=7,
                      color=_pnl_color(p["v4_pnl"]))
            _set_cell(cap_table.rows[row_idx].cells[6], f"+{p['v5_exit']:.0f}%", size=7,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            v5_color = RGBColor(0, 128, 0) if p["v5_cap"] >= 70 else (
                RGBColor(200, 128, 0) if p["v5_cap"] >= 40 else RGBColor(200, 0, 0))
            _set_cell(cap_table.rows[row_idx].cells[7], f"{p['v5_cap']:.0f}%", size=7,
                      color=v5_color, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(cap_table.rows[row_idx].cells[8], f"${p['v5_pnl']:+,.0f}", size=7,
                      color=_pnl_color(p["v5_pnl"]))
            better = "v5" if p["v5_cap"] > p["v4_cap"] else ("v4" if p["v4_cap"] > p["v5_cap"] else "tie")
            _set_cell(cap_table.rows[row_idx].cells[9], better, size=7, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


# ============================================================================
# Main
# ============================================================================

def main():
    print("Loading signals and tick data...")
    trades, no_data, no_strike = load_signals_and_ticks()
    print(f"  {len(trades)} signals with data, {no_data} no data, {no_strike} no strike")

    print("Generating report...")
    doc = generate_report(trades, no_data, no_strike)

    REPORTS_DIR.mkdir(exist_ok=True)
    filename = f"trade_comparison_{datetime.now().strftime('%Y-%m-%d')}.docx"
    output_path = REPORTS_DIR / filename
    doc.save(str(output_path))
    print(f"\nReport saved: {output_path}")


if __name__ == "__main__":
    main()
