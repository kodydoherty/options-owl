#!/usr/bin/env python3
"""Generate the updated Gold Standard + PUT Track report as .docx."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

doc = Document()

# Styles
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── Title ─────────────────────────────────────────────────────────────────
doc.add_heading("OptionsOwl Combined Strategy Report", level=0)
doc.add_heading("CALL (Gold Standard) + PUT (Scalp) — Full Writeup", level=1)

meta = doc.add_paragraph()
meta.add_run(
    f"Generated: {datetime.now().strftime('%Y-%m-%d')}  |  "
    f"Period: Feb 28 – May 21, 2026 (60 trading days, 49 active)  |  "
    f"Starting Capital: $23,000"
)

# ── Executive Summary ─────────────────────────────────────────────────────
doc.add_heading("Executive Summary", level=2)
doc.add_paragraph(
    "$23,000 → $118,161 (+413.7%) in 60 trading days. "
    "84 total trades: 73 winners, 11 losers, 86.9% combined win rate. "
    "Max drawdown 7.5%. Profit factor 6.23."
)
doc.add_paragraph(
    "This is the first combined CALL + PUT backtest. The CALL track (Gold Standard ML pipeline) "
    "produced 58 trades at 96.6% WR with PF=96.22. The PUT scalp track added 26 trades at 65.4% WR "
    "with PF=2.25, contributing an additional $21,743. Together they are complementary — "
    "PUTs profit on down days when CALLs don't enter, and CALLs dominate on trending days."
)

# Key metrics table
doc.add_heading("Key Metrics", level=3)
t = doc.add_table(rows=1, cols=5)
t.style = "Light Grid Accent 1"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t.rows[0].cells
for i, h in enumerate(["Metric", "CALL Track", "PUT Track", "Combined", "Prior Report"]):
    hdr[i].text = h

data = [
    ("Trades", "58", "26", "84", "62 (CALL only)"),
    ("Win Rate", "96.6%", "65.4%", "86.9%", "93.5%"),
    ("Total P&L", "+$73,418", "+$21,743", "+$95,161", "+$57,138"),
    ("Return", "+319.2%", "+94.5%", "+413.7%", "+248%"),
    ("Profit Factor", "96.22", "2.25", "6.23", "17.96"),
    ("Max Drawdown", "—", "—", "7.5%", "10.0%"),
    ("Avg Win", "$1,325", "$2,304", "$1,553", "$991"),
    ("Avg Loss", "-$386", "-$1,936", "-$1,654", "-$842"),
    ("Avg Hold", "~25 min", "~30 min", "~27 min", "~30 min"),
]
for row_data in data:
    row = t.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

# ── What Changed Since Last Report ────────────────────────────────────────
doc.add_heading("What Changed Since Last Report (May 24)", level=2)

changes = [
    ("PUT Scalp Track Added",
     "ATM 0DTE PUTs at 1:00-2:30 PM ET with fixed +50%/-60%/60min rules. "
     "Targets cheap PUTs ($0.05-$0.50) on green days to catch intraday reversals. "
     "5 tickers: SPY, QQQ, TSLA, META, IWM."),
    ("PUT Ticker Exclusion Gate",
     "PLTR (-$48K), AMD (-$9K), MSTR (breakeven), AVGO (breakeven), AAPL, GOOGL, NVDA, AMZN "
     "excluded from PUT trading based on 3+ year backtest showing consistent losses. "
     "These tickers are still traded on the CALL track."),
    ("PUT Market Direction Gate",
     "PUTs only enter when SPY is green or flat (cheap premiums catch reversals). "
     "When SPY drops below -0.5% (bear mode), PUTs expand to all tickers and double concurrent slots. "
     "CALLs are skipped in bear mode."),
    ("Bear Mode",
     "When SPY is down 0.5%+ from open at ~10:30 AM: skip new CALL entries, "
     "expand PUT tickers to all 13, increase PUT concurrent slots from 2 to 4. "
     "Triggered on 1 of 60 days in this period."),
    ("Signal Source Migration",
     "All 4 trading bots now run in sourcing-only mode (ENABLE_DISCORD_SIGNALS=false). "
     "Signals come from the ML sourcing pipeline via PostgreSQL, not Discord."),
    ("PostgreSQL Tick Data Capture",
     "Harvester now writes stock_ticks, option_ticks, and stock_candles to PostgreSQL "
     "for future backtesting and ML training. Fire-and-forget pattern (never blocks trading)."),
    ("Redis Cross-Agent Coordination",
     "Signal dedup (prevents 4 bots entering same signal), regime sharing, "
     "rate limiting, daily loss tracking. Graceful degradation if Redis is down."),
    ("E2E Pipeline Test",
     "scripts/e2e_signal_test.py — injects synthetic signals into PG and verifies "
     "all 4 agents consume them. Validates the full sourcing → consumer → pipeline path."),
    ("Infrastructure Deployed",
     "PostgreSQL 16, Redis 7, ML sourcing scanner, signal consumer — all running on droplet. "
     "Trading bots in paper mode for validation."),
]
for title, desc in changes:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p.add_run(f" — {desc}")

# ── System Architecture ───────────────────────────────────────────────────
doc.add_heading("System Architecture", level=2)

doc.add_heading("CALL Track — ML Gold Standard Pipeline", level=3)
doc.add_paragraph(
    "Morning scan (9:35-11:00 ET), every minute, for each of 13 tickers:"
)
steps = [
    "REGIME FILTER (daily): Evaluate SPY's first 15 minutes. If regime score < 0.19, skip the entire day for CALLs. "
    "Uses 18 features: morning range, direction, overnight gap, GEX, prev-day stats. Skipped 10 of 60 days.",
    "PATTERN ENTRY MODEL (AUC=0.890): LightGBM classifier scans ATM call premiums. "
    "Triggers on: premium drops 14%+ from open, volume surges 7.5x, IV expands, underlying stabilizes. "
    "Threshold 0.85 — passes 263 of 47,077 scanned (0.6%).",
    "ENTRY TIMING MODEL (AUC=0.839): Second LightGBM with 30 features evaluates entry timing quality. "
    "Threshold 0.70 — blocks 110 of 263 pattern-approved signals (42%).",
    "ENTRY GATES: Premium $0.20-$6.00, bid-ask spread < 15%, max 2 same-direction, max 1 index, "
    "5-min spacing, GFV buffer 15%, $5K dollar cap.",
    "DCA: After entry, monitor 8-20 minutes. If premium dips 15-35% and underlying stable (< 0.5% move), "
    "auto-double position at lower price.",
    "V5 FSM EXIT: 10-gate state machine per trade, every 5 seconds. Category-aware (HIGH_VOL/INDEX/STANDARD), "
    "DTE-aware (0DTE vs multi-day). Breakeven ratchet at +20%, scaleout 1/3 at +20%, "
    "2PM trail tightening (30% tighter).",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("PUT Track — Afternoon Scalp", level=3)
doc.add_paragraph(
    "Afternoon scan at fixed slots: 1:00, 1:30, 2:00, 2:30 PM ET."
)
put_steps = [
    "MARKET DIRECTION CHECK: SPY must be green or flat (cheap PUT premiums). "
    "If SPY < -0.5% (bear mode), expand tickers and double slots.",
    "TICKER FILTER: Only SPY, QQQ, TSLA, META, IWM. "
    "PLTR, AMD, MSTR, AVGO, AAPL, GOOGL, NVDA, AMZN excluded (3+ year losers on PUTs).",
    "PREMIUM GATE: $0.05-$0.50 (cheap OTM puts, high gamma).",
    "FIXED EXIT RULES: +50% target, -60% stop, 60-minute max hold. "
    "No FSM complexity — PUTs are binary (spike or expire).",
]
for i, s in enumerate(put_steps, 1):
    doc.add_paragraph(f"{i}. {s}")

# ── Position Sizing ──────────────────────────────────────────────────────
doc.add_heading("Position Sizing", level=2)
doc.add_paragraph(
    "deployable = portfolio x 75%\n"
    "per_slot = deployable / 6 (4 CALL + 2 PUT max concurrent)\n"
    "scaled = per_slot x 85%\n"
    "contracts = min(scaled/cost, $5K cap, 15% portfolio cap, GFV remaining, 200 max)\n\n"
    "For $23K start: deployable=$17,250, per_slot=$2,875, scaled=$2,443. "
    "As portfolio compounds to $100K+, positions scale up but $5K dollar cap prevents catastrophic sizing."
)

# ── Loss Mitigation ──────────────────────────────────────────────────────
doc.add_heading("Loss Mitigation Strategy", level=2)
doc.add_paragraph("CALL track had only 2 losses out of 58 trades. How:")
mitigations = [
    "Regime filter skips bad days entirely (10 days skipped)",
    "High pattern threshold (0.85) — only strong dip+volume signals",
    "Entry timing gate (0.70) — rejects 42% of pattern-approved signals",
    "$5K dollar cap — prevents catastrophic sizing on cheap premiums",
    "Index correlation guard — max 1 of SPY/QQQ/IWM at a time",
    "Bad day mode — after 1st hard stop, raise pattern threshold to 0.90",
    "Breakeven ratchet — once +20%, stop floor = entry price (can never lose)",
    "Fast exits — avg hold ~25 min. Get in, take profit, get out",
]
for m in mitigations:
    doc.add_paragraph(m, style="List Bullet")

doc.add_paragraph("\nPUT track had 9 losses out of 26 trades — expected for a scalp strategy. "
                   "Key: PUT losses are capped at -60% of a cheap premium ($0.05-$0.50), "
                   "so max loss per PUT ~$6K vs avg PUT win ~$2,304.")

# ── The 11 Losing Trades ─────────────────────────────────────────────────
doc.add_heading("The 11 Losing Trades", level=2)

doc.add_heading("CALL Losses (2 trades, -$771 total)", level=3)
call_losses = [
    ("2026-03-18", "PLTR", "$2.31", "21", "-$651", "theta_timer",
     "Low-conviction entry, premium bled for 231 minutes. Theta decay killed it."),
    ("2026-05-21", "AMD", "$4.00", "24", "-$120", "soft_trail",
     "Tiny loss, entered at 19.2% peak but soft trail caught the reversal early."),
]
t2 = doc.add_table(rows=1, cols=7)
t2.style = "Light Grid Accent 1"
for i, h in enumerate(["Date", "Ticker", "Entry", "Contracts", "P&L", "Exit Reason", "Notes"]):
    t2.rows[0].cells[i].text = h
for row_data in call_losses:
    row = t2.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_heading("PUT Losses (9 trades, -$17,426 total)", level=3)
doc.add_paragraph(
    "7 of 9 PUT losses were 'put_maxhold' — premium never hit the +50% target within 60 minutes "
    "and expired with small-to-moderate losses. 2 were 'put_stop' — premium crashed through -60%. "
    "IWM was the worst PUT performer (5 losses out of 11 trades), suggesting it may need to be excluded."
)

# ── Per-Ticker Breakdown ─────────────────────────────────────────────────
doc.add_heading("Per-Ticker Breakdown", level=2)
ticker_data = [
    ("AMZN", "6", "0", "+$17,158", "+$0", "+$17,158", "Top earner. AMZN $0.27 entry on 3/4 → $12,559 single trade"),
    ("SPY", "9", "6", "+$8,566", "+$5,960", "+$14,526", "Strong on both tracks"),
    ("TSLA", "7", "3", "+$7,138", "+$6,868", "+$14,006", "100% WR on PUTs, 100% on CALLs"),
    ("AAPL", "5", "0", "+$12,378", "+$0", "+$12,378", "AAPL $1.26 entry on 4/9 → $6,492"),
    ("QQQ", "2", "6", "+$1,893", "+$6,323", "+$8,216", "PUTs outperformed CALLs"),
    ("NVDA", "5", "0", "+$6,969", "+$0", "+$6,969", "Consistent CALL winner"),
    ("AMD", "8", "0", "+$5,717", "+$0", "+$5,717", "Excluded from PUTs (3yr loser)"),
    ("MSTR", "6", "0", "+$5,415", "+$0", "+$5,415", "Excluded from PUTs"),
    ("IWM", "1", "11", "+$1,466", "+$2,592", "+$4,058", "PUT-heavy. 6W/5L on PUTs — borderline"),
    ("GOOGL", "2", "0", "+$2,516", "+$0", "+$2,516", ""),
    ("META", "3", "0", "+$2,458", "+$0", "+$2,458", ""),
    ("PLTR", "4", "0", "+$1,744", "+$0", "+$1,744", "Excluded from PUTs"),
]
t3 = doc.add_table(rows=1, cols=7)
t3.style = "Light Grid Accent 1"
for i, h in enumerate(["Ticker", "CALLs", "PUTs", "CALL P&L", "PUT P&L", "Total", "Notes"]):
    t3.rows[0].cells[i].text = h
for row_data in ticker_data:
    row = t3.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

# ── Exit Reason Breakdown ────────────────────────────────────────────────
doc.add_heading("Exit Reason Breakdown", level=2)
exit_data = [
    ("soft_trail", "35", "97.1%", "+$33,278", "Workhorse exit — lets winners run, locks 60-70% of peak gain"),
    ("put_target", "17", "100.0%", "+$39,169", "PUT scalp hit +50% target — perfect hit rate"),
    ("scalp_target", "9", "100.0%", "+$13,508", "CALL scalp target (index 0DTE, +25%)"),
    ("sideways_scalp", "7", "100.0%", "+$3,916", "Small gains in choppy conditions"),
    ("put_maxhold", "7", "0.0%", "-$11,328", "PUT didn't hit target in 60 min — all losses"),
    ("adaptive_trail", "3", "100.0%", "+$15,851", "Big runners caught by adaptive trail"),
    ("put_stop", "2", "0.0%", "-$6,098", "PUT crashed through -60% stop"),
    ("breakeven_ratchet", "2", "100.0%", "+$1,024", "Was up 20%+, pulled back to entry, exited at breakeven"),
    ("theta_timer", "1", "0.0%", "-$651", "Position bled for 231 min — theta decay"),
    ("eod_data_end", "1", "100.0%", "+$6,492", "Held to close — AAPL 4/9 massive runner"),
]
t4 = doc.add_table(rows=1, cols=5)
t4.style = "Light Grid Accent 1"
for i, h in enumerate(["Exit Reason", "Count", "Win Rate", "P&L", "Description"]):
    t4.rows[0].cells[i].text = h
for row_data in exit_data:
    row = t4.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

# ── Complete Trade Log ────────────────────────────────────────────────────
doc.add_heading("Complete Trade Log (84 Trades)", level=2)

trades = [
    ("2026-03-04", "CALL", "AMZN", "$0.27", "180", "+$12,559", "580.0%", "28m", "adaptive_trail"),
    ("2026-03-04", "CALL", "MSTR", "$3.05", "12", "+$960", "36.4%", "9m", "soft_trail"),
    ("2026-03-05", "CALL", "AAPL", "$1.38", "28", "+$1,746", "118.8%", "103m", "adaptive_trail"),
    ("2026-03-06", "CALL", "NVDA", "$0.80", "50", "+$1,278", "48.7%", "17m", "soft_trail"),
    ("2026-03-06", "CALL", "AAPL", "$0.59", "71", "+$1,546", "94.9%", "31m", "adaptive_trail"),
    ("2026-03-09", "CALL", "TSLA", "$0.49", "89", "+$890", "24.5%", "8m", "soft_trail"),
    ("2026-03-10", "CALL", "AMZN", "$1.33", "33", "+$1,004", "29.3%", "8m", "soft_trail"),
    ("2026-03-10", "CALL", "MSTR", "$4.50", "10", "+$1,015", "37.8%", "46m", "scalp_target"),
    ("2026-03-11", "PUT", "TSLA", "$0.45", "103", "+$2,194", "50.0%", "34m", "put_target"),
    ("2026-03-12", "CALL", "SPY", "$2.36", "20", "+$120", "12.7%", "14m", "sideways_scalp"),
    ("2026-03-12", "CALL", "TSLA", "$5.30", "9", "+$630", "17.0%", "37m", "soft_trail"),
    ("2026-03-13", "CALL", "SPY", "$4.16", "22", "+$1,401", "21.0%", "20m", "soft_trail"),
    ("2026-03-13", "PUT", "SPY", "$0.45", "110", "+$2,343", "52.3%", "2m", "put_target"),
    ("2026-03-13", "PUT", "QQQ", "$0.38", "131", "+$2,346", "62.7%", "2m", "put_target"),
    ("2026-03-16", "CALL", "AMD", "$3.80", "13", "+$585", "18.4%", "85m", "soft_trail"),
    ("2026-03-16", "CALL", "MSTR", "$4.85", "10", "+$990", "35.1%", "12m", "soft_trail"),
    ("2026-03-16", "PUT", "IWM", "$0.28", "179", "+$2,340", "51.0%", "30m", "put_target"),
    ("2026-03-17", "CALL", "NVDA", "$2.05", "24", "+$1,071", "30.2%", "17m", "soft_trail"),
    ("2026-03-17", "CALL", "AMZN", "$0.97", "51", "+$864", "24.7%", "10m", "soft_trail"),
    ("2026-03-17", "PUT", "SPY", "$0.48", "103", "+$2,344", "54.9%", "17m", "put_target"),
    ("2026-03-17", "PUT", "QQQ", "$0.47", "105", "-$494", "22.4%", "60m", "put_maxhold"),
    ("2026-03-18", "CALL", "PLTR", "$2.31", "21", "-$651", "3.9%", "231m", "theta_timer"),
    ("2026-03-19", "CALL", "NVDA", "$1.38", "36", "+$1,668", "37.7%", "7m", "scalp_target"),
    ("2026-03-19", "CALL", "TSLA", "$4.40", "11", "+$880", "18.2%", "56m", "soft_trail"),
    ("2026-03-20", "CALL", "SPY", "$2.38", "21", "+$147", "16.4%", "5m", "soft_trail"),
    ("2026-03-20", "PUT", "IWM", "$0.38", "131", "+$2,346", "94.2%", "4m", "put_target"),
    ("2026-03-23", "PUT", "SPY", "$0.33", "151", "-$2,340", "9.2%", "60m", "put_maxhold"),
    ("2026-03-23", "PUT", "QQQ", "$0.32", "156", "-$2,560", "3.4%", "60m", "put_maxhold"),
    ("2026-03-24", "CALL", "AMD", "$4.15", "12", "+$480", "16.9%", "26m", "soft_trail"),
    ("2026-03-24", "CALL", "PLTR", "$3.70", "13", "+$1,055", "35.1%", "59m", "soft_trail"),
    ("2026-03-25", "PUT", "TSLA", "$0.46", "107", "+$2,331", "111.4%", "46m", "put_target"),
    ("2026-03-26", "CALL", "NVDA", "$2.19", "22", "+$704", "18.7%", "12m", "soft_trail"),
    ("2026-03-26", "PUT", "SPY", "$0.37", "134", "-$1,069", "45.6%", "60m", "put_maxhold"),
    ("2026-03-26", "PUT", "QQQ", "$0.29", "173", "+$2,345", "87.2%", "55m", "put_target"),
    ("2026-03-27", "CALL", "AMZN", "$1.66", "60", "+$1,240", "30.5%", "34m", "scalp_target"),
    ("2026-03-30", "CALL", "AMD", "$5.20", "9", "+$765", "19.2%", "23m", "soft_trail"),
    ("2026-03-30", "PUT", "IWM", "$0.47", "105", "+$2,338", "56.2%", "9m", "put_target"),
    ("2026-03-31", "CALL", "TSLA", "$0.91", "54", "+$1,209", "45.1%", "18m", "soft_trail"),
    ("2026-03-31", "CALL", "AAPL", "$0.79", "63", "+$1,398", "38.0%", "19m", "scalp_target"),
    ("2026-04-01", "CALL", "PLTR", "$1.89", "26", "+$260", "15.3%", "5m", "soft_trail"),
    ("2026-04-06", "PUT", "IWM", "$0.45", "110", "+$2,343", "105.2%", "4m", "put_target"),
    ("2026-04-07", "CALL", "SPY", "$0.64", "78", "+$312", "14.1%", "64m", "sideways_scalp"),
    ("2026-04-08", "CALL", "SPY", "$1.08", "46", "+$937", "30.6%", "10m", "scalp_target"),
    ("2026-04-08", "CALL", "GOOGL", "$0.70", "71", "+$1,746", "42.9%", "7m", "soft_trail"),
    ("2026-04-08", "PUT", "TSLA", "$0.45", "110", "+$2,343", "162.6%", "1m", "put_target"),
    ("2026-04-09", "CALL", "AAPL", "$1.26", "39", "+$6,492", "217.5%", "361m", "eod_data_end"),
    ("2026-04-09", "CALL", "GOOGL", "$2.18", "44", "+$770", "15.2%", "17m", "sideways_scalp"),
    ("2026-04-10", "PUT", "QQQ", "$0.47", "105", "+$2,338", "64.6%", "12m", "put_target"),
    ("2026-04-14", "CALL", "AMD", "$2.35", "21", "+$735", "16.6%", "9m", "soft_trail"),
    ("2026-04-14", "PUT", "IWM", "$0.28", "179", "-$3,047", "4.3%", "56m", "put_stop"),
    ("2026-04-15", "CALL", "AMD", "$4.20", "11", "+$605", "15.5%", "14m", "soft_trail"),
    ("2026-04-15", "CALL", "IWM", "$0.50", "200", "+$1,466", "28.6%", "19m", "scalp_target"),
    ("2026-04-16", "CALL", "TSLA", "$5.50", "18", "+$2,215", "31.0%", "21m", "soft_trail"),
    ("2026-04-16", "CALL", "META", "$4.90", "10", "+$850", "18.8%", "34m", "soft_trail"),
    ("2026-04-17", "PUT", "IWM", "$0.24", "200", "-$1,246", "30.9%", "60m", "put_maxhold"),
    ("2026-04-20", "CALL", "SPY", "$1.62", "60", "+$2,090", "26.6%", "31m", "scalp_target"),
    ("2026-04-20", "PUT", "IWM", "$0.34", "147", "+$2,348", "103.0%", "21m", "put_target"),
    ("2026-04-21", "PUT", "SPY", "$0.26", "194", "+$2,348", "55.3%", "47m", "put_target"),
    ("2026-04-21", "PUT", "QQQ", "$0.49", "101", "+$2,347", "63.8%", "18m", "put_target"),
    ("2026-04-22", "CALL", "SPY", "$0.93", "53", "+$841", "22.6%", "12m", "sideways_scalp"),
    ("2026-04-22", "CALL", "META", "$1.84", "27", "+$488", "21.7%", "7m", "soft_trail"),
    ("2026-04-24", "CALL", "SPY", "$0.44", "200", "+$1,296", "16.0%", "6m", "soft_trail"),
    ("2026-04-27", "CALL", "QQQ", "$0.86", "58", "+$629", "54.7%", "31m", "breakeven_ratchet"),
    ("2026-04-27", "CALL", "AAPL", "$0.57", "87", "+$1,196", "28.1%", "9m", "soft_trail"),
    ("2026-04-29", "CALL", "AMD", "$4.60", "10", "+$395", "20.7%", "25m", "breakeven_ratchet"),
    ("2026-04-29", "CALL", "PLTR", "$2.82", "34", "+$1,080", "20.2%", "9m", "soft_trail"),
    ("2026-04-30", "CALL", "META", "$6.00", "16", "+$1,120", "15.1%", "19m", "soft_trail"),
    ("2026-04-30", "CALL", "AMZN", "$2.34", "21", "+$1,386", "34.6%", "15m", "soft_trail"),
    ("2026-05-01", "CALL", "AMD", "$0.77", "128", "+$2,272", "38.0%", "19m", "scalp_target"),
    ("2026-05-01", "PUT", "SPY", "$0.37", "134", "+$2,335", "53.7%", "34m", "put_target"),
    ("2026-05-04", "PUT", "IWM", "$0.48", "103", "-$3,052", "28.1%", "37m", "put_stop"),
    ("2026-05-05", "CALL", "TSLA", "$3.95", "24", "+$504", "15.4%", "30m", "sideways_scalp"),
    ("2026-05-05", "PUT", "IWM", "$0.40", "124", "-$1,974", "41.9%", "60m", "put_maxhold"),
    ("2026-05-07", "PUT", "IWM", "$0.20", "200", "+$1,840", "53.3%", "6m", "put_target"),
    ("2026-05-11", "PUT", "IWM", "$0.39", "127", "-$1,645", "32.9%", "60m", "put_maxhold"),
    ("2026-05-13", "CALL", "MSTR", "$3.05", "16", "+$1,325", "36.1%", "8m", "soft_trail"),
    ("2026-05-14", "CALL", "AMZN", "$2.31", "21", "+$105", "10.8%", "22m", "sideways_scalp"),
    ("2026-05-15", "CALL", "NVDA", "$1.00", "100", "+$2,248", "30.4%", "6m", "soft_trail"),
    ("2026-05-18", "CALL", "MSTR", "$5.25", "9", "+$675", "19.6%", "9m", "soft_trail"),
    ("2026-05-19", "CALL", "SPY", "$0.40", "125", "+$1,422", "35.0%", "9m", "scalp_target"),
    ("2026-05-20", "CALL", "QQQ", "$1.38", "72", "+$1,264", "29.6%", "13m", "sideways_scalp"),
    ("2026-05-20", "CALL", "MSTR", "$4.60", "10", "+$450", "17.4%", "36m", "soft_trail"),
    ("2026-05-21", "CALL", "TSLA", "$5.15", "9", "+$810", "18.4%", "38m", "soft_trail"),
    ("2026-05-21", "CALL", "AMD", "$4.00", "24", "-$120", "19.2%", "5m", "soft_trail"),
]

# Split into chunks for readability
t5 = doc.add_table(rows=1, cols=9)
t5.style = "Light Grid Accent 1"
for i, h in enumerate(["Date", "Type", "Ticker", "Entry", "Ct", "P&L", "Peak", "Hold", "Exit Reason"]):
    t5.rows[0].cells[i].text = h
for row_data in trades:
    row = t5.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        # Color P&L cells
        if i == 5:
            if val.startswith("-"):
                row[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 0, 0)
            elif val.startswith("+"):
                row[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)

# Make table font smaller for readability
for row in t5.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)

# ── Known Limitations & Holes ────────────────────────────────────────────
doc.add_heading("Known Limitations & Risks", level=2)

holes = [
    ("Backtest != Production (Execution Gaps)",
     "Backtest uses ask for entry, bid for exit — realistic but doesn't account for partial fills, "
     "Webull rejections, or 100-500ms WebSocket latency. Dip-confirm entry gate (60s wait) "
     "is NOT simulated in backtest — production will get slightly better entries."),
    ("IWM PUT Performance is Borderline",
     "IWM is 6W/5L on PUTs with net +$2,592. It's the only PUT ticker with more than 2 losses. "
     "Consider excluding or reducing PUT allocation on IWM."),
    ("PUT Maxhold Losses are Predictable",
     "7 of 9 PUT losses are 'put_maxhold' — premium never hit +50% in 60 min. "
     "These are structural: cheap OTM PUTs that don't spike. Could add early exit on "
     "no-movement detection (if +0% after 30 min, exit)."),
    ("Sample Size",
     "84 trades over 60 days is informative but not statistically robust. "
     "The 3+ year ThetaData backtest (889 days, 1539 trades) confirms the general patterns hold."),
    ("Bull Market Bias",
     "Feb-May 2026 was generally bullish. CALL track 96.6% WR may not hold in a sustained bear market. "
     "Bear mode helps but only triggers when SPY drops 0.5%+ from open — gradual grinds down "
     "may not trigger it."),
    ("ML Model Staleness",
     "Models trained on 2023-2026 data. Market microstructure changes require periodic retraining "
     "(monthly or quarterly)."),
    ("Portfolio Compounding Creates Sizing Drift",
     "As portfolio grows from $23K to $118K, positions scale up. The $5K dollar cap prevents "
     "catastrophic sizing, but uncapped positions (cheaper premiums) can grow large."),
    ("Regime Model is Blunt (AUC=0.616)",
     "Only catches the most obvious chop days. A stronger regime model could skip more losing days "
     "and improve PF further."),
    ("No Overnight Risk Handling",
     "Multi-day options held overnight can gap down on open. The 5-min grace period + backstop "
     "limits damage but doesn't prevent it."),
]
for title, desc in holes:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p.add_run(f" — {desc}")

# ── Production Status ────────────────────────────────────────────────────
doc.add_heading("Production Deployment Status", level=2)

doc.add_heading("Infrastructure (All Deployed)", level=3)
infra = [
    "PostgreSQL 16 — trade data, ML signals, tick data capture",
    "Redis 7 — signal dedup, regime sharing, rate limiting",
    "ML Sourcing Scanner — pattern + entry timing + regime models, scans 9:35-11:00 ET",
    "Signal Consumer — polls PG ml_signals every 30s, routes to entry pipeline",
    "Harvester — captures options chain snapshots + stock candles to PG",
    "4 Trading Bots — owlet-kody/adam/vinny/yank, each with own Webull creds",
    "E2E Test Script — validates full pipeline from signal injection to agent consumption",
]
for item in infra:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Current Mode: Paper Trading", level=3)
doc.add_paragraph(
    "All 4 bots are running with PAPER_TRADE=true, WEBULL_KILL_SWITCH=true, "
    "ENABLE_DISCORD_SIGNALS=false. They consume signals from the ML sourcing pipeline only. "
    "Validation in progress — switch to live trading after confirming pipeline works correctly."
)

doc.add_heading("Deployment Commands", level=3)
doc.add_paragraph(
    "# Full E2E test\n"
    "python3 scripts/e2e_signal_test.py --droplet\n\n"
    "# Quick status check\n"
    "python3 scripts/e2e_signal_test.py --droplet --status\n\n"
    "# Deploy code changes\n"
    "./scripts/rebuild.sh\n\n"
    "# Check bot logs\n"
    "./scripts/trade-log.sh logs\n\n"
    "# P&L analysis\n"
    "python scripts/trade-pnl.py --droplet"
)

# ── Next Steps ────────────────────────────────────────────────────────────
doc.add_heading("Next Steps", level=2)
next_steps = [
    "Validate paper trading Monday — run E2E test, confirm all 4 agents pick up signals",
    "Monitor first week of sourcing-only signals — compare to Discord baseline",
    "Consider excluding IWM from PUT track (borderline performance)",
    "Add PUT early-exit rule: if +0% after 30 min, exit (avoids maxhold losses)",
    "Walk-forward validation: run backtest on unseen data (June+) as it becomes available",
    "Retrain regime model with PUT-aware features (currently CALL-only)",
    "Switch to live trading after 1-2 weeks of paper validation",
]
for i, s in enumerate(next_steps, 1):
    doc.add_paragraph(f"{i}. {s}")

# Save
output_path = "/Users/kody/dev/options-owl/reports/Combined_Strategy_Report_2026-05-25.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
