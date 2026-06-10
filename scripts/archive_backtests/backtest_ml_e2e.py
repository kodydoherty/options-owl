"""End-to-end backtest: ML signal sourcing → full filter pipeline → V5 FSM exit.

This simulates the FULL production pipeline without Discord/Neverland:
  1. Load stock candles from harvester DB → compute_indicators() → _infer_direction()
  2. Scan harvester option snapshots for ATM contracts matching inferred direction
  3. Compute ML features from real option data (bid/ask/IV/delta/volume)
  4. Run ML model (predict_entry_confidence) for entry confidence
  5. Run scoring engine (5 tiers) → quality gate → penalty veto
  6. Apply entry gates: morning cutoff, premium cap, spread gate
  7. Run V5 FSM against subsequent tick data for exit simulation
     (with candle_data for ENRG/scalp_target/resistance awareness)
  8. Track portfolio with position sizing, concurrent limits, etc.

All times are in ET (America/New_York). Never uses UTC-4 hacks.

Usage:
    python scripts/backtest_ml_e2e.py
    python scripts/backtest_ml_e2e.py --no-morning-cutoff
    python scripts/backtest_ml_e2e.py --no-scalp-target
"""

from __future__ import annotations

import argparse
import sqlite3
import sys
from collections import defaultdict
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from types import SimpleNamespace
from zoneinfo import ZoneInfo

import numpy as np
import pandas as pd

PROJECT_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_DIR))

from options_owl.risk.exit_v5.config import V5Config, get_ticker_config
from options_owl.risk.exit_v5.fsm import ExitFSM, TradeState
from options_owl.sourcing.data.indicator_engine import IndicatorSet, compute_indicators
from options_owl.sourcing.scoring.engine import compute_score
from options_owl.sourcing.scoring.types import Direction, SignalContext, SignalState
from options_owl.sourcing.filters.quality_gate import check_quality_gate
from options_owl.sourcing.filters.penalty_veto import check_penalty_veto
from options_owl.sourcing.scoring.ml_gates.signal_model import (
    compute_option_features_from_live,
    predict_entry_confidence,
)

ET = ZoneInfo("America/New_York")
UTC = ZoneInfo("UTC")

HARVESTER_DB = str(PROJECT_DIR / "journal" / "owlet-harvester" / "options_data.db")
THETADATA_DB = str(PROJECT_DIR / "journal" / "thetadata_options.db")

TICKERS = [
    "SPY", "QQQ", "NVDA", "TSLA", "META", "AAPL", "AMZN",
    "GOOGL", "MSFT", "AMD", "MSTR", "PLTR", "AVGO", "IWM",
]

# ── Configuration ────────────────────────────────────────────────────────────

PORTFOLIO_START = 20_000
MAX_CONCURRENT = 4
MAX_POSITION_PCT = 0.15
MAX_RISK_PCT = 0.75
SCORE_THRESHOLD = 60  # Sourcing agent score threshold

# Cash account constraints (GFV protection)
GFV_BUFFER_PCT = 15.0  # only deploy 85% of start-of-day balance
DAILY_LOSS_CIRCUIT_BREAKER_PCT = 15.0  # stop trading after losing 15% of SOD balance
MAX_SAME_DIRECTION = 3  # max positions in the same direction (correlation guard)

# Sampling: check every N minutes for new entries
ENTRY_SAMPLE_INTERVAL_MIN = 5

# V6 settings matching production
_V6_SETTINGS = SimpleNamespace(
    ENABLE_V6_BREAKEVEN_RATCHET=True,
    V6_BREAKEVEN_TRIGGER_PCT=20.0,
    ENABLE_V6_SCALEOUT=True,
    V6_SCALEOUT_GAIN_PCT=20.0,
    V6_SCALEOUT_FRACTION=0.333,
    V6_SCALEOUT_MIN_CONTRACTS=3,
    ENABLE_V6_2PM_TIGHTEN=True,
    V6_2PM_TRAIL_TIGHTEN_FACTOR=0.7,
    V6_2PM_SOFT_TRAIL_BOOST=0.15,
    ENABLE_V6_PER_TICKER_CONFIG=True,
    ENABLE_V6_PREMIUM_CAP=True,
    V6_PREMIUM_CAP=6.0,
    V6_PREMIUM_CAP_MID=7.0,
    V6_PREMIUM_CAP_HIGH=9.0,
    ENABLE_V6_SPREAD_GATE=True,
    V6_MAX_SPREAD_PCT=15.0,
    ENABLE_V6_EARLY_POP_GATE=True,
    ENABLE_V6_SIDEWAYS_SCALP=True,
    ENABLE_SCALP_TARGET=True,
    SCALP_TARGET_PCT=25.0,
    SCALP_RUNNER_CONFIRM_PCT=40.0,
)


# ── Data loading ─────────────────────────────────────────────────────────────


def load_trading_days(conn) -> list[str]:
    """Get all unique trading days in harvester DB."""
    rows = conn.execute(
        "SELECT DISTINCT date(captured_at) as d FROM harvest_snapshots ORDER BY d"
    ).fetchall()
    return [r[0] for r in rows]


def load_stock_candles(conn, ticker: str, date_str: str) -> list[dict]:
    """Load 5m stock candles for a ticker on a given day from harvester DB.

    Returns list of OHLCV dicts (oldest-first) matching compute_indicators() input.
    """
    day_dt = datetime.strptime(date_str, "%Y-%m-%d").replace(tzinfo=ET)
    # Market hours: 9:30 AM - 4:00 PM ET
    start_et = day_dt.replace(hour=9, minute=30)
    end_et = day_dt.replace(hour=16, minute=0)
    # Convert to UTC epoch ms (bar_start_ts is UTC epoch ms)
    start_ms = int(start_et.astimezone(UTC).timestamp() * 1000)
    end_ms = int(end_et.astimezone(UTC).timestamp() * 1000)

    rows = conn.execute("""
        SELECT bar_start_ts, open, high, low, close, volume, vwap, bar_start
        FROM stock_candles
        WHERE ticker = ? AND timeframe = '5m'
          AND bar_start_ts >= ? AND bar_start_ts <= ?
        ORDER BY bar_start_ts
    """, (ticker, start_ms, end_ms)).fetchall()

    candles = []
    for r in rows:
        candles.append({
            "timestamp": r[0],
            "open": float(r[1]),
            "high": float(r[2]),
            "low": float(r[3]),
            "close": float(r[4]),
            "volume": float(r[5] or 0),
            "vwap": float(r[6] or 0),
        })
    return candles


def load_stock_candles_multi_day(conn, ticker: str, date_str: str, lookback_days: int = 5) -> list[dict]:
    """Load stock candles for multiple days (for PDH/PDL/PWH/PWL institutional levels).

    Returns candles from up to `lookback_days` prior days + today.
    """
    day_dt = datetime.strptime(date_str, "%Y-%m-%d").replace(tzinfo=ET)
    start_et = (day_dt - timedelta(days=lookback_days + 2)).replace(hour=9, minute=30)
    end_et = day_dt.replace(hour=16, minute=0)
    start_ms = int(start_et.astimezone(UTC).timestamp() * 1000)
    end_ms = int(end_et.astimezone(UTC).timestamp() * 1000)

    rows = conn.execute("""
        SELECT bar_start_ts, open, high, low, close, volume, vwap, bar_start
        FROM stock_candles
        WHERE ticker = ? AND timeframe = '5m'
          AND bar_start_ts >= ? AND bar_start_ts <= ?
        ORDER BY bar_start_ts
    """, (ticker, start_ms, end_ms)).fetchall()

    candles = []
    for r in rows:
        candles.append({
            "timestamp": r[0],
            "open": float(r[1]),
            "high": float(r[2]),
            "low": float(r[3]),
            "close": float(r[4]),
            "volume": float(r[5] or 0),
            "vwap": float(r[6] or 0),
        })
    return candles


def load_day_snapshots(conn, date_str: str, ticker: str):
    """Load all option snapshots for a ticker on a given day."""
    rows = conn.execute("""
        SELECT hs.captured_at, hs.underlying_price, hs.bid, hs.ask, hs.midpoint,
               hs.implied_volatility, hs.delta, hs.theta, hs.vega,
               hs.day_volume, hc.strike, hc.option_type, hc.expiry_date,
               hs.contract_ticker
        FROM harvest_snapshots hs
        JOIN harvest_contracts hc ON hs.contract_ticker = hc.contract_ticker
        WHERE hc.underlying = ? AND date(hs.captured_at) = ?
          AND hs.bid > 0 AND hs.ask > 0
          AND ABS(hc.strike - hs.underlying_price) < 3
        ORDER BY hs.captured_at, ABS(hc.strike - hs.underlying_price)
    """, (ticker, date_str)).fetchall()
    return rows


def load_contract_ticks(conn, contract_ticker: str, after_ts: str):
    """Load all ticks for a specific contract after a given timestamp."""
    rows = conn.execute("""
        SELECT captured_at, midpoint, bid, ask, underlying_price,
               implied_volatility, delta, theta, vega, day_volume
        FROM harvest_snapshots
        WHERE contract_ticker = ? AND captured_at > ?
        ORDER BY captured_at
    """, (contract_ticker, after_ts)).fetchall()
    if not rows or len(rows) < 5:
        return None
    df = pd.DataFrame(rows, columns=[
        "captured_at", "midpoint", "bid", "ask", "underlying_price",
        "iv", "delta", "theta", "vega", "volume"
    ])
    df["premium"] = df["midpoint"].where(df["midpoint"] > 0, (df["bid"] + df["ask"]) / 2)
    df["premium"] = df["premium"].where(df["premium"] > 0, np.nan)
    df = df.dropna(subset=["premium"])
    if len(df) < 5:
        return None
    df["ts"] = pd.to_datetime(df["captured_at"], format="ISO8601")
    df = df.sort_values("ts").reset_index(drop=True)
    return df


# ── ThetaData loaders (1-min resolution) ─────────────────────────────────────


def load_trading_days_theta(conn, start: str = "2026-03-27", end: str = "2026-05-21") -> list[str]:
    """Get trading days from ThetaData DB within date range."""
    rows = conn.execute("""
        SELECT DISTINCT date(timestamp) as d FROM option_ohlc
        WHERE ticker = 'SPY' AND date(timestamp) >= ? AND date(timestamp) <= ?
        ORDER BY d
    """, (start, end)).fetchall()
    return [r[0] for r in rows]


def load_stock_candles_theta(conn, ticker: str, date_str: str) -> list[dict]:
    """Load 1-min stock candles from ThetaData, aggregate to 5m for indicators."""
    rows = conn.execute("""
        SELECT timestamp, open, high, low, close, volume, vwap
        FROM stock_ohlc
        WHERE ticker = ? AND date(timestamp) = ?
        ORDER BY timestamp
    """, (ticker, date_str)).fetchall()

    if not rows:
        return []

    # Aggregate 1-min bars to 5-min bars
    candles_5m = []
    batch = []
    for r in rows:
        batch.append(r)
        if len(batch) == 5:
            ts = datetime.fromisoformat(batch[0][0])
            candles_5m.append({
                "timestamp": int(ts.timestamp() * 1000),
                "open": float(batch[0][1]),
                "high": max(float(b[2]) for b in batch),
                "low": min(float(b[3]) for b in batch),
                "close": float(batch[-1][4]),
                "volume": sum(float(b[5] or 0) for b in batch),
                "vwap": float(batch[-1][6] or 0),
            })
            batch = []
    return candles_5m


def load_stock_candles_theta_multi_day(conn, ticker: str, date_str: str, lookback_days: int = 5) -> list[dict]:
    """Load multi-day 5m candles from ThetaData for institutional levels."""
    day_dt = datetime.strptime(date_str, "%Y-%m-%d")
    start = (day_dt - timedelta(days=lookback_days + 2)).strftime("%Y-%m-%d")

    rows = conn.execute("""
        SELECT timestamp, open, high, low, close, volume, vwap
        FROM stock_ohlc
        WHERE ticker = ? AND date(timestamp) >= ? AND date(timestamp) <= ?
        ORDER BY timestamp
    """, (ticker, start, date_str)).fetchall()

    if not rows:
        return []

    candles_5m = []
    batch = []
    for r in rows:
        batch.append(r)
        if len(batch) == 5:
            ts = datetime.fromisoformat(batch[0][0])
            candles_5m.append({
                "timestamp": int(ts.timestamp() * 1000),
                "open": float(batch[0][1]),
                "high": max(float(b[2]) for b in batch),
                "low": min(float(b[3]) for b in batch),
                "close": float(batch[-1][4]),
                "volume": sum(float(b[5] or 0) for b in batch),
                "vwap": float(batch[-1][6] or 0),
            })
            batch = []
    return candles_5m


def load_day_snapshots_theta(conn, date_str: str, ticker: str):
    """Load ATM option snapshots from ThetaData for a ticker on a given day.

    Returns rows in the same format as load_day_snapshots() from harvester:
    (captured_at, underlying_price, bid, ask, midpoint, iv, delta, theta, vega,
     volume, strike, option_type, expiry_date, contract_ticker)
    """
    # Find expiry dates available for this day (0DTE or nearest)
    expiries = conn.execute("""
        SELECT DISTINCT expiration FROM option_ohlc
        WHERE ticker = ? AND date(timestamp) = ?
        ORDER BY expiration
    """, (ticker, date_str)).fetchall()
    if not expiries:
        return []

    # Prefer 0DTE, fall back to nearest
    target_exp = None
    for exp_row in expiries:
        if exp_row[0] == date_str:
            target_exp = date_str
            break
    if not target_exp:
        target_exp = expiries[0][0]

    rows = conn.execute("""
        SELECT oohlc.timestamp,
               COALESCE(og.underlying_price, 0) as underlying_price,
               COALESCE(oq.bid, 0) as bid,
               COALESCE(oq.ask, 0) as ask,
               CASE
                   WHEN oq.bid > 0 AND oq.ask > 0 THEN (oq.bid + oq.ask) / 2.0
                   ELSE oohlc.close
               END as midpoint,
               COALESCE(og.implied_vol, 0) as iv,
               COALESCE(og.delta, 0) as delta,
               COALESCE(og.theta, 0) as theta,
               COALESCE(og.vega, 0) as vega,
               oohlc.volume,
               oohlc.strike,
               LOWER(oohlc.right) as option_type,
               oohlc.expiration as expiry_date,
               oohlc.ticker || '_' || oohlc.expiration || '_' || CAST(oohlc.strike AS TEXT) || '_' || oohlc.right as contract_ticker
        FROM option_ohlc oohlc
        LEFT JOIN option_quotes oq
          ON oohlc.ticker = oq.ticker AND oohlc.expiration = oq.expiration
          AND oohlc.strike = oq.strike AND oohlc.right = oq.right
          AND oohlc.timestamp = oq.timestamp
        LEFT JOIN option_greeks og
          ON oohlc.ticker = og.ticker AND oohlc.expiration = og.expiration
          AND oohlc.strike = og.strike AND oohlc.right = og.right
          AND oohlc.timestamp = og.timestamp
        WHERE oohlc.ticker = ? AND date(oohlc.timestamp) = ?
          AND oohlc.expiration = ?
        ORDER BY oohlc.timestamp, ABS(COALESCE(og.underlying_price, 0) - oohlc.strike)
    """, (ticker, date_str, target_exp)).fetchall()

    return rows


def load_contract_ticks_theta(conn, ticker: str, expiration: str, strike: float,
                               right: str, after_ts: str):
    """Load 1-min ticks for a specific contract from ThetaData.

    Returns DataFrame in the same format as load_contract_ticks().
    """
    rows = conn.execute("""
        SELECT oohlc.timestamp,
               CASE
                   WHEN oq.bid > 0 AND oq.ask > 0 THEN (oq.bid + oq.ask) / 2.0
                   ELSE oohlc.close
               END as midpoint,
               COALESCE(oq.bid, oohlc.close) as bid,
               COALESCE(oq.ask, oohlc.close) as ask,
               COALESCE(og.underlying_price, 0) as underlying_price,
               COALESCE(og.implied_vol, 0) as iv,
               COALESCE(og.delta, 0) as delta,
               COALESCE(og.theta, 0) as theta,
               COALESCE(og.vega, 0) as vega,
               oohlc.volume
        FROM option_ohlc oohlc
        LEFT JOIN option_quotes oq
          ON oohlc.ticker = oq.ticker AND oohlc.expiration = oq.expiration
          AND oohlc.strike = oq.strike AND oohlc.right = oq.right
          AND oohlc.timestamp = oq.timestamp
        LEFT JOIN option_greeks og
          ON oohlc.ticker = og.ticker AND oohlc.expiration = og.expiration
          AND oohlc.strike = og.strike AND oohlc.right = og.right
          AND oohlc.timestamp = og.timestamp
        WHERE oohlc.ticker = ? AND oohlc.expiration = ?
          AND oohlc.strike = ? AND oohlc.right = ?
          AND oohlc.timestamp > ?
        ORDER BY oohlc.timestamp
    """, (ticker, expiration, strike, right, after_ts)).fetchall()

    if not rows or len(rows) < 5:
        return None

    df = pd.DataFrame(rows, columns=[
        "captured_at", "midpoint", "bid", "ask", "underlying_price",
        "iv", "delta", "theta", "vega", "volume"
    ])
    df["premium"] = df["midpoint"].where(df["midpoint"] > 0, np.nan)
    df = df.dropna(subset=["premium"])
    if len(df) < 5:
        return None
    df["ts"] = pd.to_datetime(df["captured_at"], format="ISO8601")
    df = df.sort_values("ts").reset_index(drop=True)
    return df


# ── Direction inference (matches production scanner.py) ──────────────────────


def infer_direction(indicators: IndicatorSet) -> Direction:
    """Infer CALL/PUT from technical indicators.

    Matches production _infer_direction() in scanner.py:
    EMA cross (primary, +/-2), MACD (+/-1), VWAP (+/-1).
    """
    bullish = 0
    bearish = 0

    # EMA cross (primary — weight 2)
    if indicators.ema_cross_strength > 0.05:
        bullish += 2
    elif indicators.ema_cross_strength < -0.05:
        bearish += 2

    # MACD
    if indicators.macd_line > 0:
        bullish += 1
    elif indicators.macd_line < 0:
        bearish += 1

    # VWAP
    if indicators.vwap > 0 and indicators.last_close > indicators.vwap:
        bullish += 1
    elif indicators.vwap > 0 and indicators.last_close < indicators.vwap:
        bearish += 1

    return Direction.CALL if bullish >= bearish else Direction.PUT


def infer_direction_from_underlying(underlying_history: list[float]) -> Direction:
    """Fallback direction inference when no candle data available.

    Uses simple EMA cross from underlying price series.
    """
    if len(underlying_history) < 21:
        return Direction.CALL  # default bullish if insufficient data

    prices = np.array(underlying_history[-78:], dtype=np.float64)

    # Simple EMA 9 vs 21
    k9 = 2.0 / 10
    k21 = 2.0 / 22
    ema9 = prices[0]
    ema21 = prices[0]
    for p in prices[1:]:
        ema9 = p * k9 + ema9 * (1 - k9)
        ema21 = p * k21 + ema21 * (1 - k21)

    return Direction.CALL if ema9 > ema21 else Direction.PUT


# ── Scoring pipeline (matches production scanner.py) ─────────────────────────


def run_scoring_pipeline(
    ticker: str,
    candles_5m: list[dict],
    indicators: IndicatorSet,
    direction: Direction,
    option_snap: dict | None = None,
    score_threshold: int = SCORE_THRESHOLD,
) -> tuple[bool, int, str, SignalContext]:
    """Run the full sourcing scoring pipeline.

    Returns (passed, score, reject_reason, ctx).
    """
    ctx = SignalContext(
        ticker=ticker,
        scan_time=datetime.now(ET).isoformat(),
        state=SignalState.INDICATED,
        direction=direction,
        candles_5m=candles_5m,
        candle_source="harvester_db",
        indicators=indicators,
    )

    # Stage 3: Score
    scored = compute_score(ctx)
    ctx.score_total = scored.score
    ctx.tier1_direction = scored.breakdown.get("direction")
    ctx.tier2_timing = scored.breakdown.get("timing")
    ctx.tier3_amplifiers = scored.breakdown.get("amplifiers")
    ctx.tier4_risk = scored.breakdown.get("risk")
    ctx.tier5_calibration = scored.breakdown.get("calibration")

    if scored.rejected:
        return False, 0, scored.reject_reason or "scoring_rejected", ctx

    if scored.score < score_threshold:
        return False, scored.score, f"score {scored.score} < {score_threshold}", ctx

    # Stage 4b: Quality gate
    if not check_quality_gate(ctx, score_threshold):
        return False, scored.score, ctx.filter_reason, ctx

    # Stage 4c: Penalty veto
    if check_penalty_veto(ctx):
        return False, scored.score, ctx.filter_reason, ctx

    return True, scored.score, "", ctx


# ── ML Entry Signal ──────────────────────────────────────────────────────────


def check_ml_entry(
    ticker: str,
    snapshot: tuple,
    premium_history: list[float],
    volume_history: list[int],
    underlying_history: list[float],
    minutes_since_open: int,
    is_call: bool,
) -> dict | None:
    """Run ML model on a snapshot to check if it's a valid entry signal."""
    captured_at, underlying, bid, ask, midpoint, iv, delta, theta, vega, volume = (
        snapshot[0], float(snapshot[1] or 0), float(snapshot[2] or 0),
        float(snapshot[3] or 0), float(snapshot[4] or 0),
        float(snapshot[5] or 0), float(snapshot[6] or 0),
        float(snapshot[7] or 0), float(snapshot[8] or 0), int(snapshot[9] or 0),
    )

    premium = midpoint if midpoint > 0 else (bid + ask) / 2
    if premium <= 0:
        return None

    features = compute_option_features_from_live(
        ticker=ticker,
        premium=premium,
        bid=bid,
        ask=ask,
        iv=iv,
        delta=delta,
        theta=theta,
        vega=vega,
        volume=volume,
        underlying_price=underlying,
        minutes_since_open=minutes_since_open,
        is_call=is_call,
        premium_history=premium_history,
        volume_history=volume_history,
        underlying_history=underlying_history,
    )

    direction = "CALL" if is_call else "PUT"
    result = predict_entry_confidence(ticker, features, direction)

    if result["is_signal"] and result["confidence"] > 0:
        result["premium"] = premium
        result["bid"] = bid
        result["ask"] = ask
        result["underlying"] = underlying
        result["iv"] = iv
        result["delta"] = delta
        return result

    return None


# ── FSM Simulation ───────────────────────────────────────────────────────────


def build_candle_data_for_fsm(indicators: IndicatorSet | None) -> dict:
    """Build candle_data dict matching what position_monitor passes to FSM.

    The FSM expects: {"indicators": {"5m": {...}, "15m": {...}, ...}}
    For backtest we only have 5m data, which is sufficient for ENRG evaluation.
    """
    if indicators is None:
        return {}

    # Build the indicator dict matching candle_cache format
    ind_5m = {
        "ema_cross_strength": indicators.ema_cross_strength,
        "macd_line": indicators.macd_line,
        "macd_histogram": indicators.macd_histogram,
        "rsi9": indicators.rsi9,
        "vwap": indicators.vwap,
        "bb_squeeze": indicators.bb_squeeze,
        "atr_expanding": indicators.atr_expanding,
        "volume_ratio": indicators.volume_ratio,
        "adx": indicators.adx,
        "last_close": indicators.last_close,
        "last_high": indicators.last_high,
        "last_low": indicators.last_low,
        "session_high": indicators.session_high,
        "session_low": indicators.session_low,
        "pdh": indicators.pdh,
        "pdl": indicators.pdl,
        "pwh": indicators.pwh,
        "pwl": indicators.pwl,
    }

    return {"indicators": {"5m": ind_5m}}


def simulate_exit(
    df, entry_premium, contracts, option_type, dte, expiry_date, ticker,
    candle_data: dict | None = None,
    v6_settings=None,
):
    """Run production V5 FSM against tick data."""
    if entry_premium <= 0:
        return {"pnl": 0, "reason": "no_data", "hold": 0, "exit_prem": 0, "peak_gain": 0}

    tcfg = get_ticker_config(ticker, use_per_ticker=True)
    fsm = ExitFSM(tcfg, settings=v6_settings or _V6_SETTINGS)

    entry_ts = df["ts"].iloc[0]
    if hasattr(entry_ts, 'to_pydatetime'):
        entry_ts = entry_ts.to_pydatetime()
    if entry_ts.tzinfo is not None:
        # Convert to ET first, then strip tzinfo (FSM expects naive ET)
        entry_ts = entry_ts.astimezone(ET).replace(tzinfo=None)
    else:
        # Harvester: timestamps are UTC, convert to ET
        entry_ts_utc = entry_ts.replace(tzinfo=UTC)
        entry_ts = entry_ts_utc.astimezone(ET).replace(tzinfo=None)

    first_underlying = 0.0
    for i in range(min(5, len(df))):
        u = df["underlying_price"].iloc[i]
        if u and u > 0:
            first_underlying = float(u)
            break

    state = TradeState(
        trade_id=1, ticker=ticker, option_type=option_type,
        entry_premium=entry_premium, entry_time=entry_ts,
        contracts=contracts, peak_premium=entry_premium,
        entry_underlying_price=first_underlying,
        dte=dte, expiry_date=expiry_date or "",
    )

    locked_pnl = 0.0
    remaining = contracts

    for idx in range(1, len(df)):
        premium = df["premium"].iloc[idx]
        if np.isnan(premium) or premium <= 0:
            continue

        raw_bid = df["bid"].iloc[idx]
        raw_ask = df["ask"].iloc[idx]
        bid = float(raw_bid) if raw_bid and not pd.isna(raw_bid) else premium
        ask = float(raw_ask) if raw_ask and not pd.isna(raw_ask) else premium

        now = df["ts"].iloc[idx]
        if hasattr(now, 'to_pydatetime'):
            now = now.to_pydatetime()

        underlying = df["underlying_price"].iloc[idx] or 0.0

        # Convert to ET for FSM (handles both tz-aware and naive-UTC timestamps)
        if now.tzinfo is not None:
            now_et = now.astimezone(ET)
        else:
            now_et = now.replace(tzinfo=UTC).astimezone(ET)
        minutes_to_close = max(0, (16 * 60) - (now_et.hour * 60 + now_et.minute))
        now_naive_et = now_et.replace(tzinfo=None)

        action = fsm.evaluate(
            state, premium, bid, ask, now_naive_et,
            current_underlying=underlying,
            minutes_to_close=minutes_to_close,
            candle_data=candle_data or {},
        )

        if action.should_exit:
            if action.contracts_to_close > 0 and action.contracts_to_close < remaining:
                closed = action.contracts_to_close
                locked_pnl += (premium - entry_premium) * closed * 100
                remaining -= closed
                state.contracts = remaining
                continue

            elapsed = (now_naive_et - entry_ts).total_seconds() / 60
            peak_gain = (state.peak_premium - entry_premium) / entry_premium * 100
            pnl = locked_pnl + (premium - entry_premium) * remaining * 100
            return {
                "pnl": pnl,
                "reason": action.reason.value,
                "hold": elapsed,
                "exit_prem": premium,
                "peak_gain": peak_gain,
            }

    last_prem = df["premium"].iloc[-1]
    last_ts = df["ts"].iloc[-1]
    if hasattr(last_ts, 'to_pydatetime'):
        last_ts = last_ts.to_pydatetime()
    if last_ts.tzinfo is not None:
        last_ts = last_ts.replace(tzinfo=None)
    elapsed = (last_ts - entry_ts).total_seconds() / 60
    peak_gain = (state.peak_premium - entry_premium) / entry_premium * 100
    pnl = locked_pnl + (last_prem - entry_premium) * remaining * 100
    return {
        "pnl": pnl, "reason": "eod_data_end", "hold": elapsed,
        "exit_prem": last_prem, "peak_gain": peak_gain,
    }


# ── Main ─────────────────────────────────────────────────────────────────────


@dataclass
class BacktestConfig:
    """All tunable parameters for a backtest run."""
    portfolio_start: int = PORTFOLIO_START
    max_concurrent: int = MAX_CONCURRENT
    max_position_pct: float = MAX_POSITION_PCT
    max_risk_pct: float = MAX_RISK_PCT
    score_threshold: int = SCORE_THRESHOLD
    gfv_buffer_pct: float = GFV_BUFFER_PCT
    daily_loss_cb_pct: float = DAILY_LOSS_CIRCUIT_BREAKER_PCT
    max_same_direction: int = MAX_SAME_DIRECTION
    enable_morning_cutoff: bool = False
    morning_cutoff_hour: int = 11
    enable_scoring: bool = True
    enable_scalp_target: bool = True
    scalp_target_pct: float = 25.0
    scalp_runner_confirm_pct: float = 40.0
    premium_cap: float = 6.0
    spread_gate_pct: float = 15.0
    # After-hour scalp-only mode: after this hour, only take +scalp_target_pct scalps
    after_hour_scalp_only: int | None = None  # e.g. 11 = after 11 AM, scalp only
    quiet: bool = False  # suppress per-day output
    # Data source: "harvester" (90s snapshots) or "thetadata" (1-min bars)
    data_source: str = "harvester"
    theta_start: str = "2026-03-27"
    theta_end: str = "2026-05-21"
    # Filter to specific tickers (None = all TICKERS)
    tickers: list | None = None
    # Fixed sizing: always size from initial portfolio, never compound
    fixed_sizing: bool = False
    # Direction filter: "both", "call", or "put"
    direction_filter: str = "both"


@dataclass
class BacktestResult:
    """Results from a single backtest run."""
    total_pnl: float = 0.0
    final_portfolio: float = 0.0
    trades: int = 0
    wins: int = 0
    losses: int = 0
    win_rate: float = 0.0
    avg_win: float = 0.0
    avg_loss: float = 0.0
    max_win: float = 0.0
    max_loss: float = 0.0
    max_drawdown: float = 0.0
    avg_hold_min: float = 0.0
    sharpe: float = 0.0
    profit_factor: float = 0.0
    trading_days: int = 0
    signals_checked: int = 0
    signals_passed_ml: int = 0
    signals_blocked_scoring: int = 0
    signals_blocked_morning: int = 0
    trade_list: list = field(default_factory=list)
    daily_stats: list = field(default_factory=list)
    per_ticker: dict = field(default_factory=dict)
    config: BacktestConfig | None = None


def run_backtest(cfg: BacktestConfig, conn=None) -> BacktestResult:
    """Run a full backtest with the given configuration. Returns BacktestResult."""
    use_theta = cfg.data_source == "thetadata"
    own_conn = conn is None
    if own_conn:
        db_path = THETADATA_DB if use_theta else HARVESTER_DB
        conn = sqlite3.connect(db_path)
    # For ThetaData, we also need harvester DB for stock candles (if no theta stock data)
    theta_conn = conn if use_theta else None
    # Harvester conn for stock candles (always needed for indicators)
    harvester_conn = None
    if use_theta:
        harvester_conn = sqlite3.connect(HARVESTER_DB)

    # Build V6 settings from config
    v6 = SimpleNamespace(
        ENABLE_V6_BREAKEVEN_RATCHET=True,
        V6_BREAKEVEN_TRIGGER_PCT=20.0,
        ENABLE_V6_SCALEOUT=True,
        V6_SCALEOUT_GAIN_PCT=20.0,
        V6_SCALEOUT_FRACTION=0.333,
        V6_SCALEOUT_MIN_CONTRACTS=3,
        ENABLE_V6_2PM_TIGHTEN=True,
        V6_2PM_TRAIL_TIGHTEN_FACTOR=0.7,
        V6_2PM_SOFT_TRAIL_BOOST=0.15,
        ENABLE_V6_PER_TICKER_CONFIG=True,
        ENABLE_V6_PREMIUM_CAP=True,
        V6_PREMIUM_CAP=cfg.premium_cap,
        V6_PREMIUM_CAP_MID=7.0,
        V6_PREMIUM_CAP_HIGH=9.0,
        ENABLE_V6_SPREAD_GATE=True,
        V6_MAX_SPREAD_PCT=cfg.spread_gate_pct,
        ENABLE_V6_EARLY_POP_GATE=True,
        ENABLE_V6_SIDEWAYS_SCALP=True,
        ENABLE_SCALP_TARGET=cfg.enable_scalp_target,
        SCALP_TARGET_PCT=cfg.scalp_target_pct,
        SCALP_RUNNER_CONFIRM_PCT=cfg.scalp_runner_confirm_pct,
    )

    if use_theta:
        trading_days = load_trading_days_theta(conn, cfg.theta_start, cfg.theta_end)
    else:
        trading_days = load_trading_days(conn)
    if not cfg.quiet:
        sizing_mode = "FIXED (no compounding)" if cfg.fixed_sizing else "COMPOUNDING"
        print(f"Portfolio: ${cfg.portfolio_start:,} | Sizing: {sizing_mode}")
        print(f"Score threshold: {cfg.score_threshold}")
        print(f"Data source: {cfg.data_source}")
        print(f"Trading days: {len(trading_days)} ({trading_days[0] if trading_days else '?'} to {trading_days[-1] if trading_days else '?'})")

    portfolio = cfg.portfolio_start
    results = []
    daily_stats = []
    signals_checked = 0
    signals_passed_ml = 0
    signals_blocked_morning = 0
    signals_blocked_premium = 0
    signals_blocked_spread = 0
    signals_blocked_concurrent = 0
    signals_blocked_scoring = 0
    signals_blocked_direction = 0
    signals_blocked_gfv = 0
    signals_blocked_circuit_breaker = 0
    signals_blocked_correlation = 0
    peak_portfolio = portfolio
    max_drawdown = 0.0

    for day_idx, date_str in enumerate(trading_days):
        day_results = []
        open_trades: set[str] = set()  # tickers with open trades this day

        # ── Cash account tracking (GFV + circuit breaker) ──
        sod_balance = portfolio
        gfv_limit = sod_balance * (1 - cfg.gfv_buffer_pct / 100)
        daily_spent = 0.0
        daily_pnl_realized = 0.0
        circuit_breaker_tripped = False
        open_directions: dict[str, str] = {}

        # Build market open time for this day (9:30 AM ET)
        day_dt = datetime.strptime(date_str, "%Y-%m-%d")
        market_open_et = day_dt.replace(hour=9, minute=30, tzinfo=ET)

        if not cfg.quiet:
            print(f"\n[Day {day_idx+1}/{len(trading_days)}] {date_str}:", end="", flush=True)

        active_tickers = cfg.tickers if cfg.tickers else TICKERS
        for ticker in active_tickers:
            # ── Stage 1: Load candle data for direction + indicators ──
            if use_theta:
                candles_5m = load_stock_candles_theta(conn, ticker, date_str)
                candles_multi = load_stock_candles_theta_multi_day(conn, ticker, date_str) if candles_5m else []
            else:
                candles_5m = load_stock_candles(conn, ticker, date_str)
                candles_multi = load_stock_candles_multi_day(conn, ticker, date_str) if candles_5m else []
            has_candle_data = len(candles_5m) >= 10

            # Pre-compute indicators if we have candle data
            indicators = None
            direction = None
            candle_data_fsm = {}

            if has_candle_data:
                # Use multi-day candles for institutional levels (PDH/PDL)
                ind_candles = candles_multi if len(candles_multi) >= 30 else candles_5m
                indicators = compute_indicators(ind_candles)
                direction = infer_direction(indicators)
                candle_data_fsm = build_candle_data_for_fsm(indicators)

            # ── Stage 2: Load option snapshots ──
            if use_theta:
                snapshots = load_day_snapshots_theta(conn, date_str, ticker)
            else:
                snapshots = load_day_snapshots(conn, date_str, ticker)
            if not snapshots:
                continue

            # Build underlying history from snapshots for fallback direction
            underlying_hist_all: list[float] = []
            for s in snapshots:
                u = float(s[1] or 0)
                if u > 0 and (not underlying_hist_all or u != underlying_hist_all[-1]):
                    underlying_hist_all.append(u)

            # If no candle data, infer direction from underlying price history
            if direction is None and len(underlying_hist_all) >= 10:
                direction = infer_direction_from_underlying(underlying_hist_all)

            if direction is None:
                signals_blocked_direction += 1
                continue

            # Apply direction filter (--direction call/put)
            if cfg.direction_filter == "call" and direction != Direction.CALL:
                continue
            if cfg.direction_filter == "put" and direction != Direction.PUT:
                continue

            # Filter snapshots to matching direction
            dir_str = "call" if direction == Direction.CALL else "put"
            filtered_snaps = [s for s in snapshots if s[11] == dir_str]
            if not filtered_snaps:
                continue

            # ── Stage 3: Run scoring pipeline (if candle data available) ──
            scoring_passed = True
            score = 0
            scoring_reason = ""
            if cfg.enable_scoring and has_candle_data:
                scoring_passed, score, scoring_reason, _ = run_scoring_pipeline(
                    ticker, candles_5m, indicators, direction,
                    score_threshold=cfg.score_threshold,
                )
                if not scoring_passed:
                    signals_blocked_scoring += 1
                    continue

            # ── Stage 4: Sample snapshots for ML entry ──
            premium_hist: list[float] = []
            volume_hist: list[int] = []
            underlying_hist: list[float] = []
            last_check_ts = None
            already_entered_today = False
            is_call = direction == Direction.CALL

            for snap in filtered_snaps:
                # Update histories
                mid = float(snap[4] or 0)
                vol = int(snap[9] or 0)
                und = float(snap[1] or 0)
                if mid > 0:
                    premium_hist.append(mid)
                if vol > 0:
                    volume_hist.append(vol)
                if und > 0:
                    underlying_hist.append(und)

                if already_entered_today:
                    continue

                # Parse timestamp → ET
                try:
                    snap_dt = datetime.fromisoformat(snap[0])
                    if snap_dt.tzinfo is None:
                        snap_dt = snap_dt.replace(tzinfo=UTC)
                    snap_et = snap_dt.astimezone(ET)
                except (ValueError, TypeError):
                    continue

                # Rate limit
                if last_check_ts is not None:
                    elapsed = (snap_et - last_check_ts).total_seconds() / 60
                    if elapsed < ENTRY_SAMPLE_INTERVAL_MIN:
                        continue
                last_check_ts = snap_et

                # Market hours check
                if snap_et.hour < 9 or (snap_et.hour == 9 and snap_et.minute < 30):
                    continue
                if snap_et.hour >= 15 and snap_et.minute >= 30:
                    continue

                # Morning cutoff
                if cfg.enable_morning_cutoff:
                    cutoff_min = cfg.morning_cutoff_hour * 60
                    snap_min = snap_et.hour * 60 + snap_et.minute
                    if snap_min >= cutoff_min:
                        signals_blocked_morning += 1
                        break

                minutes_since_open = int((snap_et - market_open_et.replace(
                    year=snap_et.year, month=snap_et.month, day=snap_et.day
                )).total_seconds() / 60)
                if minutes_since_open < 0:
                    continue

                signals_checked += 1

                # Daily loss circuit breaker
                if circuit_breaker_tripped:
                    signals_blocked_circuit_breaker += 1
                    continue

                # Concurrent positions check
                if len(open_trades) >= cfg.max_concurrent:
                    signals_blocked_concurrent += 1
                    continue

                # Already have this ticker open today?
                if ticker in open_trades:
                    continue

                # Correlation guard: max N positions same direction
                same_dir_count = sum(1 for d in open_directions.values() if d == dir_str)
                if same_dir_count >= cfg.max_same_direction:
                    signals_blocked_correlation += 1
                    continue

                # ── ML entry check ──
                ml_result = check_ml_entry(
                    ticker, snap[1:11],
                    premium_hist[-15:], volume_hist[-15:],
                    underlying_hist[-15:],
                    minutes_since_open, is_call,
                )
                if ml_result is None:
                    continue

                signals_passed_ml += 1
                option_type = dir_str

                entry_premium = ml_result["ask"]
                if entry_premium <= 0:
                    entry_premium = ml_result["premium"]
                if entry_premium <= 0:
                    continue

                # Premium cap gate
                if entry_premium > cfg.premium_cap:
                    signals_blocked_premium += 1
                    continue

                # Spread gate
                if ml_result["bid"] > 0 and ml_result["ask"] > 0:
                    spread_pct = (ml_result["ask"] - ml_result["bid"]) / ml_result["ask"] * 100
                    if spread_pct > cfg.spread_gate_pct:
                        signals_blocked_spread += 1
                        continue

                # ── Position sizing (with GFV constraint) ──
                contract_ticker = snap[13]
                expiry_date = snap[12]

                sizing_base = cfg.portfolio_start if cfg.fixed_sizing else portfolio
                deployable = sizing_base * cfg.max_risk_pct
                per_slot = deployable / cfg.max_concurrent
                confidence = ml_result["confidence"]
                if confidence >= 0.90:
                    mult = 0.95
                elif confidence >= 0.80:
                    mult = 0.60
                else:
                    mult = 1.00

                cost_per = entry_premium * 100
                scaled = per_slot * mult
                position_cap = sizing_base * cfg.max_position_pct

                # GFV check: can we afford this trade with settled funds?
                gfv_remaining = gfv_limit - daily_spent
                if gfv_remaining < cost_per:
                    signals_blocked_gfv += 1
                    continue
                raw_contracts = int(scaled / cost_per) if cost_per > 0 else 1
                cap_contracts = int(position_cap / cost_per) if cost_per > 0 else 1
                # Also cap by GFV remaining settled funds
                gfv_max_contracts = int(gfv_remaining / cost_per) if cost_per > 0 else 1
                contracts = max(1, min(raw_contracts, cap_contracts, gfv_max_contracts))

                try:
                    exp_dt = datetime.strptime(expiry_date, "%Y-%m-%d").date()
                    dte = max(0, (exp_dt - day_dt.date()).days)
                except (ValueError, TypeError):
                    dte = 0

                # ── Load tick data and simulate exit ──
                if use_theta:
                    # ThetaData: parse contract info from snapshot
                    snap_strike = float(snap[10])
                    snap_right = snap[11].upper()  # 'call'→'CALL', 'put'→'PUT'
                    tick_df = load_contract_ticks_theta(
                        conn, ticker, expiry_date, snap_strike, snap_right, snap[0]
                    )
                else:
                    tick_df = load_contract_ticks(conn, contract_ticker, snap[0])
                if tick_df is None:
                    continue

                result = simulate_exit(
                    tick_df, entry_premium, contracts, option_type,
                    dte, expiry_date, ticker,
                    candle_data=candle_data_fsm,
                    v6_settings=v6,
                )

                # Track GFV spend (cost of this trade)
                trade_cost = contracts * cost_per
                daily_spent += trade_cost

                trade = {
                    "ticker": ticker,
                    "direction": dir_str,
                    "day": date_str,
                    "entry_time_et": snap_et.strftime("%H:%M"),
                    "entry": entry_premium,
                    "contracts": contracts,
                    "cost": trade_cost,
                    "ml_confidence": round(confidence, 3),
                    "ml_threshold": round(ml_result["threshold"], 3),
                    "runner_score": round(ml_result.get("runner_score", 0), 3),
                    "model_source": ml_result.get("model_source", "unknown"),
                    "score": score,
                    "had_candles": has_candle_data,
                    "dte": dte,
                    "pnl": result["pnl"],
                    "reason": result["reason"],
                    "hold": result["hold"],
                    "exit_prem": result["exit_prem"],
                    "peak_gain": result["peak_gain"],
                }
                day_results.append(trade)
                results.append(trade)
                portfolio += result["pnl"]

                # Track realized P&L for circuit breaker
                daily_pnl_realized += result["pnl"]
                loss_pct = abs(daily_pnl_realized) / sod_balance * 100 if daily_pnl_realized < 0 else 0
                if loss_pct >= cfg.daily_loss_cb_pct:
                    circuit_breaker_tripped = True

                open_trades.add(ticker)
                open_directions[ticker] = dir_str
                already_entered_today = True
                # Track drawdown
                if portfolio > peak_portfolio:
                    peak_portfolio = portfolio
                dd = (peak_portfolio - portfolio) / peak_portfolio * 100 if peak_portfolio > 0 else 0
                if dd > max_drawdown:
                    max_drawdown = dd

                if not cfg.quiet:
                    pnl_str = f"+${result['pnl']:,.0f}" if result["pnl"] >= 0 else f"-${abs(result['pnl']):,.0f}"
                    print(f" {ticker}{'C' if is_call else 'P'} {pnl_str}", end="", flush=True)

        # Daily summary
        if day_results:
            day_pnl = sum(t["pnl"] for t in day_results)
            day_wins = sum(1 for t in day_results if t["pnl"] > 0)
            day_cost = sum(t["cost"] for t in day_results)
            daily_stats.append({
                "day": date_str,
                "trades": len(day_results),
                "pnl": day_pnl,
                "wins": day_wins,
                "portfolio": portfolio,
                "daily_spent": day_cost,
                "gfv_limit": gfv_limit,
                "circuit_breaker": circuit_breaker_tripped,
            })
            if not cfg.quiet:
                cb_flag = " CB!" if circuit_breaker_tripped else ""
                print(f" | Day: ${day_pnl:,.0f} Port: ${portfolio:,.0f} "
                      f"(spent ${day_cost:,.0f}/{gfv_limit:,.0f}){cb_flag}")
        else:
            if not cfg.quiet:
                print(" (no trades)")

    if own_conn:
        conn.close()
    if harvester_conn:
        harvester_conn.close()

    # Build result
    res = BacktestResult(config=cfg)
    res.trade_list = results
    res.daily_stats = daily_stats
    res.trading_days = len(trading_days)
    res.signals_checked = signals_checked
    res.signals_passed_ml = signals_passed_ml
    res.signals_blocked_scoring = signals_blocked_scoring
    res.signals_blocked_morning = signals_blocked_morning
    res.max_drawdown = max_drawdown

    if not results:
        res.final_portfolio = portfolio
        return res

    df_r = pd.DataFrame(results)
    pnls = df_r["pnl"]
    wins_n = int((pnls > 0).sum())
    losses_n = int((pnls <= 0).sum())

    res.trades = len(results)
    res.wins = wins_n
    res.losses = losses_n
    res.total_pnl = float(pnls.sum())
    res.final_portfolio = portfolio
    res.win_rate = wins_n / len(pnls) * 100 if len(pnls) > 0 else 0
    res.avg_win = float(pnls[pnls > 0].mean()) if wins_n > 0 else 0
    res.avg_loss = float(pnls[pnls <= 0].mean()) if losses_n > 0 else 0
    res.max_win = float(pnls.max())
    res.max_loss = float(pnls.min())
    res.avg_hold_min = float(df_r["hold"].mean())

    # Profit factor
    gross_win = float(pnls[pnls > 0].sum()) if wins_n > 0 else 0
    gross_loss = abs(float(pnls[pnls <= 0].sum())) if losses_n > 0 else 1
    res.profit_factor = gross_win / gross_loss if gross_loss > 0 else float("inf")

    # Sharpe (daily returns)
    if daily_stats:
        daily_pnls = [d["pnl"] for d in daily_stats]
        if len(daily_pnls) > 1:
            mean_r = np.mean(daily_pnls)
            std_r = np.std(daily_pnls, ddof=1)
            res.sharpe = (mean_r / std_r * np.sqrt(252)) if std_r > 0 else 0

    # Per-ticker breakdown
    for tkr, group in df_r.groupby("ticker"):
        gpnl = group["pnl"]
        res.per_ticker[tkr] = {
            "trades": len(gpnl),
            "pnl": float(gpnl.sum()),
            "win_rate": float((gpnl > 0).sum() / len(gpnl) * 100),
        }

    return res


def print_results(res: BacktestResult):
    """Print full backtest results to console."""
    if res.trades == 0:
        print("\nNo trades generated. Check ML models exist in journal/models/signal_ml_v2/")
        return

    df_r = pd.DataFrame(res.trade_list)
    pnls = df_r["pnl"]

    print(f"\n{'=' * 90}")
    print(f"ML E2E BACKTEST — {res.trades} trades over {len(res.daily_stats)} days")
    print(f"{'=' * 90}")
    print(f"Starting Portfolio: ${res.config.portfolio_start:,}")
    print(f"Final Portfolio:    ${res.final_portfolio:,.2f}")
    print(f"Total P&L:          ${res.total_pnl:,.2f} ({res.total_pnl/res.config.portfolio_start*100:+.1f}%)")
    print(f"Win Rate:           {res.win_rate:.1f}% ({res.wins}W / {res.losses}L)")
    print(f"Avg Win:            ${res.avg_win:,.2f}")
    print(f"Avg Loss:           ${res.avg_loss:,.2f}")
    print(f"Max Win:            ${res.max_win:,.2f}")
    print(f"Max Loss:           ${res.max_loss:,.2f}")
    print(f"Max Drawdown:       {res.max_drawdown:.1f}%")
    print(f"Profit Factor:      {res.profit_factor:.2f}")
    print(f"Sharpe (ann.):      {res.sharpe:.2f}")
    print(f"Avg Hold:           {res.avg_hold_min:.0f} min")
    print(f"\nSignals checked:    {res.signals_checked:,}")
    print(f"ML signals passed:  {res.signals_passed_ml:,}")
    print(f"Blocked (scoring):  {res.signals_blocked_scoring:,}")

    # Per-direction breakdown
    if "direction" in df_r.columns:
        print(f"\n{'Direction':<8} {'Trades':>6} {'Total P&L':>12} {'Avg P&L':>10} {'Win%':>6} {'PF':>6}")
        print("-" * 55)
        for d, group in df_r.groupby("direction"):
            gpnl = group["pnl"]
            gwins = (gpnl > 0).sum()
            gwr = gwins / len(gpnl) * 100
            gross_win = gpnl[gpnl > 0].sum()
            gross_loss = abs(gpnl[gpnl < 0].sum()) or 1
            pf = gross_win / gross_loss
            print(f"{d:<8} {len(gpnl):>6} ${gpnl.sum():>10,.2f} ${gpnl.mean():>8,.2f} {gwr:>5.0f}% {pf:>5.2f}")

    # Exit reasons
    print(f"\n{'Reason':<25} {'Count':>6} {'Total P&L':>12} {'Avg P&L':>10} {'Win%':>6}")
    print("-" * 62)
    for reason, group in df_r.groupby("reason"):
        gpnl = group["pnl"]
        gwins = (gpnl > 0).sum()
        gwr = gwins / len(gpnl) * 100
        print(f"{reason:<25} {len(gpnl):>6} ${gpnl.sum():>10,.2f} ${gpnl.mean():>8,.2f} {gwr:>5.0f}%")

    # Per-ticker
    print(f"\n{'Ticker':<8} {'Trades':>6} {'Total P&L':>12} {'Avg P&L':>10} {'Win%':>6}")
    print("-" * 50)
    for tkr, group in sorted(df_r.groupby("ticker"), key=lambda x: x[1]["pnl"].sum(), reverse=True):
        gpnl = group["pnl"]
        gwins = (gpnl > 0).sum()
        gwr = gwins / len(gpnl) * 100
        print(f"{tkr:<8} {len(gpnl):>6} ${gpnl.sum():>10,.2f} ${gpnl.mean():>8,.2f} {gwr:>5.0f}%")

    # Daily P&L
    print(f"\n{'Day':<12} {'Trades':>6} {'Day P&L':>10} {'Portfolio':>12} {'W/L':>6}")
    print("-" * 55)
    df_daily = pd.DataFrame(res.daily_stats)
    for _, row in df_daily.iterrows():
        losses_d = row["trades"] - row["wins"]
        print(f"{row['day']:<12} {row['trades']:>6} ${row['pnl']:>8,.2f} "
              f"${row['portfolio']:>10,.2f} {int(row['wins'])}/{int(losses_d)}")

    # Per-trade log
    print(f"\n{'Day':<12} {'Time':>5} {'Ticker':<6} {'Dir':<5} {'Scr':>4} {'Conf':>5} {'Entry':>6} "
          f"{'Ct':>3} {'P&L':>9} {'Peak%':>6} {'Hold':>5} {'Reason':<20} {'Candle':>6}")
    print("-" * 115)
    for _, t in df_r.iterrows():
        print(f"{t['day']:<12} {t['entry_time_et']:>5} {t['ticker']:<6} {t['direction'][:4]:<5} "
              f"{int(t['score']):>4} {t['ml_confidence']:>5.2f} ${t['entry']:>4.2f} {t['contracts']:>3} "
              f"${t['pnl']:>8,.2f} {t['peak_gain']:>5.0f}% {t['hold']:>4.0f}m "
              f"{t['reason']:<20} {'Y' if t['had_candles'] else 'N':>6}")



def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--no-morning-cutoff", action="store_true")
    parser.add_argument("--no-scalp-target", action="store_true")
    parser.add_argument("--no-scoring", action="store_true", help="Skip scoring pipeline")
    parser.add_argument("--morning-cutoff-hour", type=int, default=11)
    parser.add_argument("--portfolio", type=int, default=PORTFOLIO_START)
    parser.add_argument("--score-threshold", type=int, default=SCORE_THRESHOLD)
    parser.add_argument("--data-source", choices=["harvester", "thetadata"], default="harvester",
                        help="Data source: harvester (90s snapshots) or thetadata (1-min bars)")
    parser.add_argument("--theta-start", type=str, default="2026-03-27")
    parser.add_argument("--theta-end", type=str, default="2026-05-21")
    parser.add_argument("--fixed-sizing", action="store_true",
                        help="Use fixed position sizing (no compounding)")
    parser.add_argument("--direction", choices=["call", "put", "both"], default="both",
                        help="Filter trades by direction (default: both)")
    parser.add_argument("--ticker", type=str, default=None,
                        help="Filter to specific ticker(s), comma-separated (e.g. NVDA,TSLA)")
    parser.add_argument("--max-concurrent", type=int, default=MAX_CONCURRENT,
                        help="Max concurrent positions (default: 4)")
    parser.add_argument("--max-position-pct", type=float, default=MAX_POSITION_PCT * 100,
                        help="Max position size as %% of portfolio (default: 15)")
    args = parser.parse_args()

    tickers_list = [t.strip().upper() for t in args.ticker.split(",")] if args.ticker else None
    cfg = BacktestConfig(
        portfolio_start=args.portfolio,
        max_concurrent=args.max_concurrent,
        max_position_pct=args.max_position_pct / 100.0,
        score_threshold=args.score_threshold,
        enable_morning_cutoff=not args.no_morning_cutoff,
        morning_cutoff_hour=args.morning_cutoff_hour,
        enable_scoring=not args.no_scoring,
        enable_scalp_target=not args.no_scalp_target,
        data_source=args.data_source,
        theta_start=args.theta_start,
        theta_end=args.theta_end,
        fixed_sizing=args.fixed_sizing,
        direction_filter=args.direction,
        tickers=tickers_list,
    )

    res = run_backtest(cfg)
    print_results(res)

    # Chart
    if res.daily_stats:
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            import matplotlib.dates as mdates

            df_daily = pd.DataFrame(res.daily_stats)
            fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 8), height_ratios=[2, 1])
            fig.suptitle("ML E2E Backtest — Portfolio Equity Curve", fontsize=14, fontweight="bold")

            dates = pd.to_datetime(df_daily["day"])
            ax1.plot(dates, df_daily["portfolio"], "b-o", markersize=5, linewidth=2)
            ax1.axhline(y=cfg.portfolio_start, color="gray", linestyle="--", alpha=0.5,
                        label=f"Start ${cfg.portfolio_start:,}")
            ax1.fill_between(dates, df_daily["portfolio"], cfg.portfolio_start,
                             where=df_daily["portfolio"] >= cfg.portfolio_start, alpha=0.15, color="green")
            ax1.fill_between(dates, df_daily["portfolio"], cfg.portfolio_start,
                             where=df_daily["portfolio"] < cfg.portfolio_start, alpha=0.15, color="red")
            ax1.set_ylabel("Portfolio Value ($)")
            ax1.legend()
            ax1.grid(True, alpha=0.3)

            colors = ["green" if p >= 0 else "red" for p in df_daily["pnl"]]
            ax2.bar(dates, df_daily["pnl"], color=colors, alpha=0.7, width=0.8)
            ax2.axhline(y=0, color="gray", linestyle="--", alpha=0.5)
            ax2.set_ylabel("Daily P&L ($)")
            ax2.set_xlabel("Date (ET)")
            ax2.grid(True, alpha=0.3)

            for ax in (ax1, ax2):
                ax.xaxis.set_major_formatter(mdates.DateFormatter("%m/%d"))
                plt.setp(ax.xaxis.get_majorticklabels(), rotation=45, ha="right")

            plt.tight_layout()
            chart_path = str(PROJECT_DIR / "ml_e2e_backtest.png")
            plt.savefig(chart_path, dpi=150, bbox_inches="tight")
            print(f"\nChart saved: {chart_path}")
            plt.close()
        except ImportError:
            pass


def _export_docx(df_r, df_daily, total_pnl, win_rate, wins, losses, final_portfolio,
                  signals_checked, signals_passed_ml, signals_blocked_morning,
                  signals_blocked_scoring,
                  enable_morning_cutoff, morning_cutoff_hour, enable_scoring):
    """Export comprehensive DOCX report."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    doc.add_heading("OptionsOwl ML E2E Backtest Report", 0)
    doc.add_paragraph(
        f"Generated: {datetime.now(ET).strftime('%Y-%m-%d %H:%M ET')} | "
        f"Start: ${PORTFOLIO_START:,} | Final: ${final_portfolio:,.2f} | "
        f"Period: {df_daily['day'].iloc[0]} to {df_daily['day'].iloc[-1]} "
        f"({len(df_daily)} trading days)"
    )
    doc.add_paragraph(
        "This backtest uses the ML signal sourcing agent (LightGBM models trained on "
        "ThetaData) as the entry source — NOT Discord/Neverland signals. "
        "Direction is inferred from stock candle indicators (EMA cross, MACD, VWAP) "
        "matching the production sourcing scanner. "
        "Entry scoring uses the full 5-tier pipeline with quality gate and penalty veto. "
        "Exits use the production V5 FSM with all V6 enhancements including "
        "scalp target, ENRG candle confirmation, and institutional level awareness. "
        "All times are in ET (America/New_York)."
    )

    # Summary
    doc.add_heading("Executive Summary", level=1)
    summary = [
        ("Total P&L", f"${total_pnl:,.2f} ({total_pnl/PORTFOLIO_START*100:+.1f}%)"),
        ("Win Rate", f"{win_rate:.1f}% ({wins}W / {losses}L)"),
        ("Trades", f"{len(df_r)}"),
        ("Avg Win", f"${df_r[df_r['pnl']>0]['pnl'].mean():,.2f}" if wins else "N/A"),
        ("Avg Loss", f"${df_r[df_r['pnl']<=0]['pnl'].mean():,.2f}" if losses else "N/A"),
        ("Max Win", f"${df_r['pnl'].max():,.2f}"),
        ("Max Loss", f"${df_r['pnl'].min():,.2f}"),
        ("Avg Hold", f"{df_r['hold'].mean():.0f} min"),
        ("Signals Checked", f"{signals_checked:,}"),
        ("ML Signals Passed", f"{signals_passed_ml:,}"),
        ("Scoring Pipeline Blocked", f"{signals_blocked_scoring:,}"),
        ("Morning Cutoff", f"{'ON at ' + str(morning_cutoff_hour) + ':00 AM ET' if enable_morning_cutoff else 'OFF'}"),
        ("Scoring Pipeline", f"{'ON (5-tier + quality + penalty)' if enable_scoring else 'OFF'}"),
        ("Scalp Target", f"{'ON (+25%, runner +40%)' if _V6_SETTINGS.ENABLE_SCALP_TARGET else 'OFF'}"),
    ]
    tbl = doc.add_table(rows=len(summary), cols=2)
    for i, (k, v) in enumerate(summary):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v

    # Exit reasons
    doc.add_heading("Exit Reason Breakdown", level=1)
    doc.add_paragraph(
        "Each exit reason corresponds to a V5 FSM gate. The gate priority determines "
        "which fires first. Scalp target (+25%) only fires if candle data does NOT "
        "confirm a runner (peak < 40%, no ENRG HOLD)."
    )
    reason_groups = df_r.groupby("reason")
    tbl = doc.add_table(rows=len(reason_groups) + 1, cols=5)
    for i, h in enumerate(["Reason", "Count", "Total P&L", "Avg P&L", "Win%"]):
        tbl.rows[0].cells[i].text = h
        for run in tbl.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for ri, (reason, group) in enumerate(reason_groups, 1):
        gpnl = group["pnl"]
        gwins = (gpnl > 0).sum()
        tbl.rows[ri].cells[0].text = str(reason)
        tbl.rows[ri].cells[1].text = str(len(gpnl))
        tbl.rows[ri].cells[2].text = f"${gpnl.sum():,.2f}"
        tbl.rows[ri].cells[3].text = f"${gpnl.mean():,.2f}"
        tbl.rows[ri].cells[4].text = f"{gwins/len(gpnl)*100:.0f}%"

    # Per-ticker
    doc.add_heading("Per-Ticker Breakdown", level=1)
    ticker_groups = sorted(df_r.groupby("ticker"), key=lambda x: x[1]["pnl"].sum(), reverse=True)
    tbl = doc.add_table(rows=len(ticker_groups) + 1, cols=5)
    for i, h in enumerate(["Ticker", "Trades", "Total P&L", "Avg P&L", "Win%"]):
        tbl.rows[0].cells[i].text = h
        for run in tbl.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for ri, (tkr, group) in enumerate(ticker_groups, 1):
        gpnl = group["pnl"]
        gwins = (gpnl > 0).sum()
        tbl.rows[ri].cells[0].text = str(tkr)
        tbl.rows[ri].cells[1].text = str(len(gpnl))
        tbl.rows[ri].cells[2].text = f"${gpnl.sum():,.2f}"
        tbl.rows[ri].cells[3].text = f"${gpnl.mean():,.2f}"
        tbl.rows[ri].cells[4].text = f"{gwins/len(gpnl)*100:.0f}%"

    # Daily + trade log
    doc.add_heading("Daily Trade Log", level=1)
    for day_str, day_group in df_r.groupby("day"):
        day_pnl = day_group["pnl"].sum()
        day_wins = (day_group["pnl"] > 0).sum()
        day_losses = len(day_group) - day_wins
        doc.add_heading(
            f"{day_str} — ${day_pnl:,.2f} ({day_wins}W/{day_losses}L)", level=2
        )
        tbl = doc.add_table(rows=len(day_group) + 1, cols=11)
        headers = ["Time ET", "Ticker", "Dir", "Score", "ML Conf", "Entry$",
                    "Ct", "P&L", "Peak%", "Hold", "Exit Reason"]
        for i, h in enumerate(headers):
            tbl.rows[0].cells[i].text = h
            for run in tbl.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)
        for ti, (_, t) in enumerate(day_group.iterrows(), 1):
            vals = [
                t["entry_time_et"], t["ticker"], t["direction"][:4],
                str(int(t["score"])), f"{t['ml_confidence']:.2f}",
                f"${t['entry']:.2f}", str(int(t["contracts"])),
                f"${t['pnl']:,.2f}", f"{t['peak_gain']:.0f}%",
                f"{t['hold']:.0f}m", t["reason"],
            ]
            for i, val in enumerate(vals):
                cell = tbl.rows[ti].cells[i]
                cell.text = val
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8)
                    if i == 7 and t["pnl"] > 0:
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif i == 7 and t["pnl"] < 0:
                        run.font.color.rgb = RGBColor(200, 0, 0)

    # Chart
    chart_path = str(PROJECT_DIR / "ml_e2e_backtest.png")
    if Path(chart_path).exists():
        doc.add_heading("Equity Curve", level=1)
        doc.add_picture(chart_path, width=Inches(6.5))

    # Methodology
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        "Entry Pipeline (matches production sourcing scanner):\n"
        "1. Stock candle data loaded from harvester DB (5m bars)\n"
        "2. compute_indicators() computes EMA9/21, MACD(5,13,4), RSI9, "
        "Bollinger, Keltner, ATR, VWAP, ADX, institutional levels (PDH/PDL/PWH/PWL)\n"
        "3. Direction inferred: EMA cross (weight 2), MACD (weight 1), VWAP (weight 1)\n"
        "4. 5-tier scoring engine: Direction, Timing, Amplifiers, Risk, Calibration\n"
        "5. Quality gate: score >= 60, 2+ tiers contributing, direction tier >= 10\n"
        "6. Penalty veto: RSI overextend + choppy, wide spread + low vol, weak dir + heavy penalties\n"
        "7. ML model (LightGBM per-ticker): predict_entry_confidence on real option features\n"
        "8. Entry gates: premium cap ($6), spread gate (15%), morning cutoff\n"
    )
    doc.add_paragraph(
        "Exit Pipeline (production V5 FSM with V6 enhancements):\n"
        "1. EOD cutoff (0DTE: 15min before close)\n"
        "2. Bid disappearance (30s zero bid)\n"
        "3. Grace period (5min, backstop at -65% 0DTE / -75% multi-day, "
        "with ENRG candle confirmation)\n"
        "4. Profit target (index 0DTE: +30%)\n"
        "5. Break-even ratchet (once +20%, floor = entry)\n"
        "6. Scale-out (sell 1/3 at +20%)\n"
        "7. Scalp target (+25% unless candle-confirmed runner past +40%)\n"
        "8. Scalp trail, Checkpoint cut, Graduated stop\n"
        "9. Soft trail, Adaptive trail (category-aware: HIGH_VOL/INDEX/STANDARD)\n"
        "10. Theta exit (stale losers)\n"
    )
    doc.add_paragraph(
        "Position Sizing:\n"
        f"Portfolio: ${PORTFOLIO_START:,} | Max concurrent: {MAX_CONCURRENT} | "
        f"Max risk: {int(MAX_RISK_PCT*100)}% | Max position: {int(MAX_POSITION_PCT*100)}%\n"
        "Confidence-weighted: 0.70-0.80 = 100%, 0.80-0.90 = 60%, 0.90+ = 95%"
    )

    docx_path = str(PROJECT_DIR / "ml_e2e_backtest_report.docx")
    doc.save(docx_path)
    print(f"\nDOCX report saved: {docx_path}")


if __name__ == "__main__":
    main()
