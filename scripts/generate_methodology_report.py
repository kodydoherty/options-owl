"""Generate Backtesting Methodology Write-Up as a .docx file.

Explains how we backtest, data sources, gotchas, and all signals.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import sqlite3
import datetime

PROJECT_DIR = Path(__file__).resolve().parent.parent
OUTPUT = PROJECT_DIR / "Backtesting_Methodology.docx"
SIGNALS_DB = str(PROJECT_DIR / "journal" / "owlet-kody" / "raw_messages.db")
HARVESTER_DB = str(PROJECT_DIR / "journal" / "owlet-harvester" / "options_data.db")


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '2F5496')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                if c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(table.rows[r_idx + 1].cells[c_idx], 'D9E2F3')
    return table


def main():
    doc = Document()

    # ── Title ──
    title = doc.add_heading('OptionsOwl Backtesting Methodology', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f'How We Backtest, Data Sources, Pitfalls & Lessons Learned\n'
        f'Generated {datetime.date.today().strftime("%B %d, %Y")}'
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(89, 89, 89)

    # ── Overview ──
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        'This document explains exactly how our backtesting works, what data we use, '
        'what assumptions we make, and the critical pitfalls we discovered. '
        'If you are running your own backtest, you MUST follow these guidelines or your '
        'results will be misleading.'
    )

    # ── Data Sources ──
    doc.add_heading('2. Data Sources', level=1)

    doc.add_heading('2.1 Signal Source: Discord (Neverland Pirates)', level=2)
    doc.add_paragraph(
        'All trade signals come from the Neverland Pirates Discord server. '
        'Our discord_collector.py bot captures every message in real-time and parses '
        'actionable trade signals with ticker, strike, expiry, direction, and score.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Important: ')
    run.bold = True
    p.add_run(
        'We only backtest signals with score >= 78 (our entry threshold). '
        'Signals are timestamped in UTC when captured. All time-of-day logic must convert UTC to ET.'
    )

    doc.add_heading('2.2 Price Data: Polygon Harvester', level=2)
    doc.add_paragraph(
        'Our harvester bot (owlet-harvester) runs during market hours and captures options chain '
        'snapshots from the Polygon API every ~60 seconds. Each snapshot includes:'
    )
    fields = [
        'midpoint (mid of bid/ask)',
        'bid and ask (for spread calculation)',
        'underlying_price (stock price at that moment)',
        'implied_volatility, delta, gamma, theta, vega (Greeks)',
        'day_volume (options volume)',
    ]
    for f in fields:
        doc.add_paragraph(f, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('Coverage: ')
    run.bold = True
    p.add_run(
        '19 trading days from March 27 to May 1, 2026. '
        'The harvester captures snapshots for all tickers we trade across multiple strikes and expiries.'
    )

    doc.add_heading('2.3 Contract Ticker Construction', level=2)
    doc.add_paragraph(
        'To find the right options data, we construct the OCC contract ticker format: '
        'O:{TICKER}{YYMMDD}{C/P}{STRIKE*1000}. For example, O:NVDA260413C00115000 = '
        'NVDA Apr 13 2026 $115 Call.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Expiry matching: ')
    run.bold = True
    p.add_run(
        'Not all tickers have daily 0DTE options. If no data exists for the signal date, '
        'we try the next 1-5 business days. This mimics what our smart entry does in production. '
        'We record the actual DTE (days to expiry) for each matched trade.'
    )

    # ── How the Backtest Works ──
    doc.add_heading('3. How the Backtest Works', level=1)

    doc.add_heading('3.1 Entry Simulation', level=2)
    steps = [
        ('1. Match signal to harvester data',
         'Find the contract ticker in the harvester DB. Use the first snapshot after the signal timestamp '
         'as entry. If no data for that date, try next business days.'),
        ('2. Entry price = ask price',
         'We use the ASK price (not mid, not bid) as our entry because that is the realistic fill price '
         'when buying options. If ask is missing, we fall back to midpoint.'),
        ('3. Position sizing',
         'Same formula as production: deployable = portfolio * 0.75, per_slot = deployable / 5, '
         'then score-based multiplier. $8,000 portfolio baseline.'),
    ]
    for title_text, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(title_text + ': ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('3.2 Exit Simulation (FSM Replay)', level=2)
    doc.add_paragraph(
        'We replay the V5 Finite State Machine against the harvester tick data. '
        'For each snapshot after entry, we feed the premium, bid, ask, timestamp, and underlying '
        'price into the FSM evaluate() function. The first gate that triggers causes the exit.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Critical detail: ')
    run.bold = True
    p.add_run(
        'We use the ACTUAL FSM code from production (options_owl/risk/exit_v5/fsm.py), not a simplified '
        'reimplementation. This ensures our backtest results match what the bot would actually do. '
        'If you rewrite the exit logic for backtesting, your results will diverge from production.'
    )

    doc.add_heading('3.3 What We Track Per Trade', level=2)
    metrics = [
        ('P&L', '(exit_premium - entry_premium) * contracts * 100'),
        ('Exit Reason', 'Which FSM gate triggered the exit'),
        ('Hold Time', 'Minutes from entry to exit'),
        ('Peak Gain (MFE)', 'Maximum favorable excursion — the best unrealized P&L'),
        ('Scale-out P&L', 'Profits locked from partial sells'),
        ('DCA', 'Whether additional contracts were added'),
    ]
    for name, desc in metrics:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}: ')
        run.bold = True
        p.add_run(desc)

    # ── Critical Pitfalls ──
    doc.add_heading('4. Critical Pitfalls (Lessons Learned)', level=1)
    doc.add_paragraph(
        'These are mistakes we made or nearly made. If your backtest does not account for these, '
        'your results are wrong.'
    )

    pitfalls = [
        ('4.1 Entry Price: Use ASK, Not Mid',
         'When you BUY an option, you pay the ask price. Using the midpoint makes every trade look '
         '$0.05-$0.50 better than reality. Over 133 trades, this inflates P&L by thousands of dollars. '
         'Always use ask for entry, bid for exit (or midpoint for exit if you want to be generous).',
         'HIGH — can inflate total P&L by 30-50%'),

        ('4.2 Timestamps Are UTC',
         'All timestamps in our database are UTC. Market hours are 9:30 AM - 4:00 PM ET (UTC-4). '
         'If your backtest uses raw timestamps for time-of-day logic (EOD cutoff, 2PM tightening, '
         'theta timer), you MUST convert to ET first. A trade at 18:00 UTC is 2:00 PM ET, not 6 PM.',
         'HIGH — wrong timezone makes EOD/theta gates fire at wrong times'),

        ('4.3 Not All Tickers Have Daily 0DTE',
         'SPY/QQQ have daily options. NVDA/TSLA/META have Mon/Wed/Fri options. AMD/PLTR have weekly only. '
         'If your backtest assumes 0DTE exists every day for every ticker, you are testing contracts '
         'that do not exist. The actual contract might be 1-2 DTE with higher premium and different behavior.',
         'HIGH — leads to phantom trades that cannot be executed'),

        ('4.4 Harvester Gaps',
         'Our harvester captures snapshots every ~60 seconds, but it can miss intervals '
         '(API errors, rate limits, restarts). A trade might have 60 ticks or 200 ticks per day. '
         'If you assume uniform 60s intervals, your hold time and exit timing will be slightly off.',
         'MEDIUM — minor timing differences'),

        ('4.5 Bid-Ask Spread Matters',
         'A $2.00 option with $1.80 bid / $2.20 ask has a 20% spread. You enter at $2.20, '
         'but the midpoint shows $2.00. To break even, mid must reach $2.20 (a 10% gain from mid). '
         'Wide spreads make profitable exits much harder. Our spread gate (>15% = reject) addresses this.',
         'HIGH — wide spreads destroy profitability of short-term trades'),

        ('4.6 Survivorship Bias in Signal Selection',
         'We only backtest signals that scored >= 78. The signals that scored lower were rejected '
         'before entry. Do not add those back into the backtest to "increase sample size" — '
         'they would have been rejected in production.',
         'MEDIUM — including rejected signals overstates trade count'),

        ('4.7 In-Sample Overfitting',
         'When we optimize per-ticker configs on the same data we test on, we risk overfitting. '
         'The configs might look great on these 133 trades but fail on new data. '
         'True validation requires out-of-sample testing on signals we have not yet seen.',
         'HIGH — per-ticker configs may not generalize. Must validate with new data.'),

        ('4.8 Premium vs Underlying Price',
         'Options premium is NOT a linear function of underlying price. A 0DTE call with 30 minutes '
         'to expiry loses value from theta even if the stock goes up. Do not backtest options by '
         'looking at stock price alone — you must use actual options premium data.',
         'CRITICAL — stock price backtests are meaningless for options'),

        ('4.9 Gamma Death Zone (After 2 PM ET)',
         'In the last 2 hours before expiry, gamma accelerates rapidly on 0DTE options. '
         'A $0.50 move in the underlying can cause a $1.00 swing in premium. '
         'Backtests that do not account for this will show unrealistic afternoon gains. '
         'Our 2PM trail tightening addresses this.',
         'HIGH — afternoon trades behave very differently from morning trades'),

        ('4.10 Volume and Liquidity',
         'Low-volume options may show midpoints that are not actually tradeable. '
         'If day_volume is 0, the quotes may be stale market maker quotes with wide spreads. '
         'Always check that the option had actual volume before trusting the price.',
         'MEDIUM — stale quotes create phantom profitability'),

        ('4.11 Scale-Out Requires Multiple Contracts',
         'Scale-out at +20% only works if you have >= 3 contracts (sell 1/3). '
         'With a small portfolio ($500), most trades are 1 contract — no scale-out possible. '
         'Do not assume scale-out benefits apply to all portfolio sizes.',
         'MEDIUM — scale-out benefits depend on portfolio size'),

        ('4.12 DCA Amplifies Losses Too',
         'DCA (adding contracts on a dip) doubles your exposure. If the trade recovers, great. '
         'If it keeps falling, your loss is 2x what it would have been. '
         'We saw this on IWM Apr 21: -$790 without DCA became -$1,270 with DCA. '
         'Only enable DCA for tickers where backtest shows NET positive across all signals.',
         'HIGH — DCA can turn small losses into large losses'),
    ]

    for title_text, description, severity in pitfalls:
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(description)
        p = doc.add_paragraph()
        run = p.add_run(f'Severity: {severity}')
        run.bold = True

    # ── Our Signal Dataset ──
    doc.add_heading('5. Complete Signal Dataset', level=1)
    doc.add_paragraph(
        'Below is every signal we have with score >= 78. These are the signals used in all our backtests. '
        '160 total signals, 133 had matching harvester data, 9 had no data, 18 were below our '
        'data quality threshold (< 10 ticks).'
    )

    # Load all signals
    conn = sqlite3.connect(SIGNALS_DB)
    rows = conn.execute("""
        SELECT id, ticker, direction, score, strike, expiry, atm_premium, otm_premium,
               entry_price, created_at
        FROM trade_signals
        WHERE score >= 78
        ORDER BY created_at
    """).fetchall()
    conn.close()

    doc.add_heading('5.1 Summary Statistics', level=2)
    from collections import Counter
    ticker_counts = Counter(r[1] for r in rows)
    day_counts = Counter(r[9][:10] for r in rows)

    stats = [
        ['Metric', 'Value'],
        ['Total Signals', str(len(rows))],
        ['Date Range', f'{rows[0][9][:10]} to {rows[-1][9][:10]}'],
        ['Trading Days', str(len(day_counts))],
        ['Unique Tickers', str(len(ticker_counts))],
        ['Avg Signals/Day', f'{len(rows)/len(day_counts):.1f}'],
        ['Score >= 95', str(sum(1 for r in rows if (r[3] or 0) >= 95))],
        ['Score 90-94', str(sum(1 for r in rows if 90 <= (r[3] or 0) < 95))],
        ['Score 85-89', str(sum(1 for r in rows if 85 <= (r[3] or 0) < 89))],
        ['Score 78-84', str(sum(1 for r in rows if 78 <= (r[3] or 0) < 85))],
    ]
    add_styled_table(doc, stats[0], stats[1:])

    doc.add_heading('5.2 Signals Per Ticker', level=2)
    ticker_data = [['Ticker', 'Count', '% of Total']]
    for t, c in ticker_counts.most_common():
        ticker_data.append([t, str(c), f'{c/len(rows)*100:.1f}%'])
    add_styled_table(doc, ticker_data[0], ticker_data[1:])

    doc.add_heading('5.3 Signals Per Day', level=2)
    day_data = [['Date', 'Count', 'Notes']]
    for d in sorted(day_counts):
        weekday = datetime.datetime.strptime(d, '%Y-%m-%d').strftime('%A')
        day_data.append([d, str(day_counts[d]), weekday])
    add_styled_table(doc, day_data[0], day_data[1:])

    doc.add_heading('5.4 All Signals (Full Detail)', level=2)
    doc.add_paragraph(
        'Every signal in our database. Premium = ATM premium (or OTM if ATM unavailable). '
        'Timestamps are UTC.'
    )

    signal_rows = [['#', 'Date', 'Time (UTC)', 'Ticker', 'Dir', 'Score', 'Strike', 'Expiry', 'Premium']]
    for i, r in enumerate(rows, 1):
        sig_id, ticker, direction, score, strike, expiry, atm_prem, otm_prem, entry, created = r
        premium = atm_prem or otm_prem or 0
        dt_str = created[:10]
        time_str = created[11:19] if len(created) > 19 else ''
        dir_str = (direction or 'call')[:4]
        strike_str = f'${strike:.0f}' if strike else 'N/A'
        prem_str = f'${premium:.2f}' if premium else 'N/A'
        exp_str = expiry or 'N/A'
        signal_rows.append([
            str(i), dt_str, time_str, ticker, dir_str,
            str(score or ''), strike_str, exp_str, prem_str
        ])
    add_styled_table(doc, signal_rows[0], signal_rows[1:])

    # ── Harvester Coverage ──
    doc.add_heading('5.5 Harvester Data Coverage', level=2)

    conn2 = sqlite3.connect(HARVESTER_DB)
    # Get date + ticker coverage
    coverage = conn2.execute("""
        SELECT date(captured_at) as d,
               COUNT(DISTINCT contract_ticker) as contracts,
               COUNT(*) as snapshots
        FROM harvest_snapshots
        GROUP BY d
        ORDER BY d
    """).fetchall()
    conn2.close()

    cov_data = [['Date', 'Unique Contracts', 'Total Snapshots']]
    for d, contracts, snaps in coverage:
        cov_data.append([d, str(contracts), f'{snaps:,}'])
    add_styled_table(doc, cov_data[0], cov_data[1:])

    # ── Backtest vs Production Differences ──
    doc.add_heading('6. Backtest vs Production Differences', level=1)
    doc.add_paragraph(
        'No backtest perfectly replicates live trading. Here are the known differences:'
    )

    diffs = [
        ['Aspect', 'Backtest', 'Production', 'Impact'],
        ['Data granularity', '~60s Polygon snapshots', '5s WebSocket + REST', 'Backtest misses intra-minute spikes'],
        ['Entry fill', 'Ask price from snapshot', 'Webull market order fill', 'Actual fill may be better or worse'],
        ['Exit fill', 'Premium at trigger tick', 'Webull market order fill', 'Slippage on exit not captured'],
        ['Multi-source pricing', 'Polygon only', 'Polygon WS + REST + yfinance + delta approx', 'Backtest has cleaner data'],
        ['ENRG gate', 'Not simulated', 'Active (multi-TF candle voting)', 'Backtest may differ on stop decisions'],
        ['Webull rejection', 'All trades execute', 'Some rejected (invalid contract, etc)', 'Backtest has more trades'],
        ['Portfolio balance', 'Fixed $8,000', 'Dynamic (sync from Webull daily)', 'Sizing changes intraday'],
        ['Concurrent trades', 'No limit enforced', '5 max concurrent in production', 'Backtest may overcount capacity'],
    ]
    add_styled_table(doc, diffs[0], diffs[1:])

    # ── How to Run Your Own Backtest ──
    doc.add_heading('7. How to Run Your Own Backtest', level=1)

    doc.add_heading('7.1 Prerequisites', level=2)
    prereqs = [
        'Python 3.12+ with options_owl package installed (pip install -e ".[dev]")',
        'Access to journal/owlet-kody/raw_messages.db (signals)',
        'Access to journal/owlet-harvester/options_data.db (price data)',
        'The V5 FSM code in options_owl/risk/exit_v5/',
    ]
    for p_text in prereqs:
        doc.add_paragraph(p_text, style='List Bullet')

    doc.add_heading('7.2 Step-by-Step', level=2)
    steps = [
        '1. Load signals from trade_signals table (score >= 78)',
        '2. For each signal, construct the OCC contract ticker and find harvester data',
        '3. Use ASK price as entry (not mid)',
        '4. Compute contracts using score-based sizing formula',
        '5. Create an ExitFSM with V5Config (or per-ticker config for V6)',
        '6. Replay each harvester snapshot through fsm.evaluate()',
        '7. Record exit reason, P&L, hold time, MFE',
        '8. DO NOT forget to convert timestamps to ET for time-based gates',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading('7.3 Reference Scripts', level=2)
    scripts = [
        ['Script', 'Purpose'],
        ['scripts/backtest_v5_production.py', 'Baseline V5 backtest (start here)'],
        ['scripts/backtest_per_ticker_tuning.py', '12 configs per ticker comparison'],
        ['scripts/backtest_dca.py', 'Early DCA strategies (all lose money)'],
        ['scripts/backtest_dca_delayed.py', 'Delayed DCA strategies (MID_TRADE wins)'],
        ['scripts/backtest_v6_combined.py', 'V6 full combined strategy backtest'],
    ]
    add_styled_table(doc, scripts[0], scripts[1:])

    # ── Key Findings Summary ──
    doc.add_heading('8. Key Findings From Backtesting', level=1)

    findings = [
        ('Entry filtering saves more than exit optimization',
         'Blocking 13 bad entries (premium cap + spread gate) saved $3,282. '
         'The META $25.35 call alone lost -$2,520. No exit strategy can fix a bad entry.'),

        ('One-size-fits-all configs leave money on the table',
         'Per-ticker optimized configs improved P&L by $5,569. NVDA needs early profit-taking, '
         'GOOGL needs wider stops, META needs defensive configs. Each ticker has its own personality.'),

        ('Early DCA (0-5 min) loses money',
         'All 4 early DCA strategies tested were net negative. Doubling down in the first 3-5 minutes '
         'catches trades that are dying. The thesis has not been proven yet.'),

        ('Delayed DCA (8-20 min) works for specific tickers',
         'MID_TRADE_DIP DCA (8-20 min window) added $1,442 across MSFT, IWM, SPY, QQQ, AMZN, NVDA. '
         'By minute 8, the instant-death trades have already been stopped out.'),

        ('Scale-out at +20% is the single biggest improvement',
         '46 partial fills locked $4,224 in profits. Even if the trade reverses after, '
         'the scale-out profits are banked. This is the most reliable P&L improvement.'),

        ('MFE capture ratio is our biggest weakness',
         'We capture only ~16% of peak gains on average. Target is >60%. '
         'Most of our trades reach +30-50% at peak but we exit at +5-10%. '
         'The adaptive trail and per-ticker configs aim to fix this.'),

        ('Afternoon trades are fundamentally different',
         'After 2 PM ET, gamma accelerates on 0DTE options. The same trailing stop width that works '
         'in the morning is too loose in the afternoon. We tighten trails by 30% after 2 PM.'),

        ('META and AMZN are structurally harder',
         'META has 11% avg bid-ask spread and low volume. AMZN has high premiums with wide spreads. '
         'Both need defensive configs or should be avoided when premiums are too high.'),
    ]

    for title_text, desc in findings:
        p = doc.add_paragraph()
        run = p.add_run(title_text + ': ')
        run.bold = True
        p.add_run(desc)

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('-- End of Methodology Report --')
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.size = Pt(10)

    doc.save(str(OUTPUT))
    print(f"Report saved to: {OUTPUT}")


if __name__ == "__main__":
    main()
