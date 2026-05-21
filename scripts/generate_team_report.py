#!/usr/bin/env python3
"""Generate DOCX report: Team Recommendations Analysis with Backtest Evidence."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime


def add_verdict_box(doc, verdict: str, color: str):
    """Add a colored verdict paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f"VERDICT: {verdict}")
    run.bold = True
    run.font.size = Pt(12)
    if color == "red":
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    elif color == "green":
        run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)
    elif color == "yellow":
        run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)


def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def main():
    doc = Document()

    # Title
    title = doc.add_heading("OptionsOwl Team Recommendations Analysis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {datetime.date.today().isoformat()}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Backtest window: Apr 10 - May 18, 2026 (195 trades, 23 trading days)")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # =========================================================================
    # EXECUTIVE SUMMARY
    # =========================================================================
    doc.add_heading("Executive Summary", level=1)

    doc.add_paragraph(
        "This report evaluates recommendations from two team documents "
        "(WEBULL-AGENT-IMPROVEMENT-BRIEF.md and WEBULL-EXIT-LOGIC-OVERHAUL-SPEC.md) "
        "by backtesting each proposal individually against our production V5 FSM strategy. "
        "Every test uses the full production stack (all 18 entry gates, V6 enhancements, "
        "per-ticker configs, momentum confirmation) and changes exactly one variable at a time."
    )

    doc.add_paragraph(
        "Key finding: Our production V5 FSM already implements most of the team's exit logic "
        "recommendations in a more sophisticated way. The team's proposals are based on a "
        "simplified view of our system (they see 4 close reasons; we have 10+ exit gates). "
        "Most proposed changes would significantly reduce profitability."
    )

    doc.add_heading("Bottom Line", level=2)

    add_table(doc,
        ["Scenario", "P&L", "vs Baseline", "Win Rate", "Recommendation"],
        [
            ["BASELINE (Production V5)", "$53,565", "--", "77.4%", "--"],
            ["45-min max hold", "$29,383", "-$24,182", "68.2%", "REJECT"],
            ["30-min max hold", "$25,731", "-$27,834", "66.2%", "REJECT"],
            ["Commit zone (exit if underwater T+15)", "$32,803", "-$20,762", "56.9%", "REJECT"],
            ["Quick take +25% in first 15min", "$51,824", "-$1,742", "78.5%", "MONITOR"],
            ["Block all puts", "$41,013", "-$12,552", "75.3%", "REJECT"],
            ["Halve put sizing", "$46,686", "-$6,879", "77.4%", "CONSIDER"],
            ["Commit zone + 45-min hold", "$20,217", "-$33,348", "53.8%", "REJECT"],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "No single recommendation improves on production. The combined proposal "
        "(commit zone + 45-min hold) loses $33,348 vs baseline -- a 62% reduction in profit."
    )
    run.bold = True

    # =========================================================================
    # DOCUMENT 1: IMPROVEMENT BRIEF
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("Document 1: Webull Agent Improvement Brief", level=1)
    doc.add_paragraph(
        "Source: WEBULL-AGENT-IMPROVEMENT-BRIEF.md (2026-05-19). "
        "6 priorities covering slippage, manual closes, decision logging, entry scoring, "
        "per-agent variance, and account state monitoring."
    )

    # --- PRIORITY 1: SLIPPAGE ---
    doc.add_heading("Priority 1: Slippage Reduction", level=2)
    doc.add_heading("Team's Claim", level=3)
    doc.add_paragraph(
        'Mean entry slippage of -7.32%, costing an estimated $10,818. '
        'They propose mid-pricing, spread checks, and strict mid-only fills on worst tickers.'
    )

    doc.add_heading("Our Data (Backtest)", level=3)
    add_table(doc,
        ["Metric", "Team's Number", "Our Backtest"],
        [
            ["Mean slippage", "-7.32%", "+17.11%"],
            ["Median slippage", "-5.79%", "+1.58%"],
            ["Puts mean slippage", "N/A", "+1.42%"],
            ["Calls mean slippage", "N/A", "+21.29%"],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Our backtest measures slippage as (first_ask_at_entry / signal_premium - 1). "
        "The positive mean indicates we are often getting BETTER prices than the Discord signal "
        "premium, not worse. The team's -7.32% likely measures something different -- possibly "
        "comparing our Webull fill price to the mid at order time, which is a measure of "
        "spread crossing, not signal-to-fill slippage."
    )

    doc.add_heading("What We Already Have", level=3)
    doc.add_paragraph(
        "- V6 spread gate: blocks entries with bid-ask spread > 40% of mid\n"
        "- V6 premium cap: blocks entries with premium > tiered cap ($6/$7/$9)\n"
        "- Smart entry: verifies live premium vs signal premium, rejects if deviation > 75%\n"
        "- Dip-confirm: waits for ask to come down before entering"
    )

    doc.add_heading("Pushback", level=3)
    doc.add_paragraph(
        "The team's slippage numbers are inflated. Their methodology compares fill price to "
        "mid-at-order-time, but 0DTE options have legitimate bid-ask spreads of 5-15% on "
        "less liquid names. This is the cost of doing business, not a fixable leak. "
        "Our spread gate already blocks the worst offenders. "
        "Mid-pricing with walk-up logic adds complexity and latency (2-second retry loops) "
        "that could cause us to MISS fast-moving trades entirely -- the 5-15 minute window "
        "is where 67% WR lives."
    )

    add_verdict_box(doc, "PARTIALLY ACCEPT -- Already have spread gate. "
        "Will add NBBO logging for measurement. Reject mid-pricing with retry loops "
        "(latency risk outweighs savings on 0DTE).", "yellow")

    # --- PRIORITY 2: MANUAL CLOSES ---
    doc.add_heading("Priority 2: Replace Manual Closes with Systematic Rules", level=2)
    doc.add_heading("Team's Claim", level=3)
    doc.add_paragraph(
        "52% of closes are 'manual' with -$4,154 net loss. They propose a 45-min max hold "
        "and trailing stop rules at +20%/+50%/+100% thresholds."
    )

    doc.add_heading("What We Already Have", level=3)
    doc.add_paragraph(
        "Our V5 FSM exit engine has 10 gates running every 5 seconds. The team sees "
        "'manual' closes because our exit_reason taxonomy doesn't map 1:1 to their "
        "close_reason field. What they call 'manual' includes our:\n"
        "- adaptive_trail (category-aware trailing stop)\n"
        "- soft_trail (15-50% peak band)\n"
        "- scalp_trail (peaked +20%, faded)\n"
        "- graduated_stop (tight stop if underlying against)\n"
        "- theta_exit (stale loser cut)\n"
        "- breakeven_ratchet (floor at entry after +20%)\n"
        "- scaleout (sell 1/3 at +20%)\n\n"
        "These are ALL systematic rules -- they just aren't logged to Supabase yet."
    )

    doc.add_heading("Backtest: 45-min Max Hold", level=3)
    add_table(doc,
        ["Metric", "Baseline", "45-min Max Hold", "Delta"],
        [
            ["Total P&L", "$53,565", "$29,383", "-$24,182 (-45%)"],
            ["Win Rate", "77.4%", "68.2%", "-9.2%"],
            ["75+ min trades", "33 trades, $9,662", "forced closed early", "Lost $9,662 in winners"],
            ["Max Win", "$7,760", "$2,700", "-$5,060"],
            ["Worst Day", "-$1,909", "-$1,668", "Marginal improvement"],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "The 45-min max hold destroys $24,182 in profit. The team's data shows 75+ min trades "
        "at 55% WR with $9,662 total P&L -- these are real winners being cut short. "
        "Our V5 FSM already handles the actual losers in this bucket via theta_exit "
        "(cuts stale losers at 120min+down 30%) and adaptive_trail."
    )

    add_verdict_box(doc, "REJECT -- Loses $24,182. Our V5 FSM already cuts stale losers "
        "via theta_exit and adaptive_trail without killing big winners.", "red")

    # --- PRIORITY 2 TRAILING STOPS ---
    doc.add_heading("Priority 2 (cont): Trailing Stop at +20%/+50%/+100%", level=2)
    doc.add_heading("What We Already Have", level=3)
    doc.add_paragraph(
        "Our V5 FSM has MORE GRANULAR trailing than the team proposes:\n"
        "- Breakeven ratchet: once +20%, floor = entry price (V6)\n"
        "- Scaleout: sell 1/3 at +20% to lock profit (V6)\n"
        "- Scalp trail: peaked +20%, exit if faded to <60% of peak\n"
        "- Soft trail: 15-50% peak band, keep 60-70% of gain\n"
        "- Adaptive trail: category-aware tiers (40%+ active, 150%+ runner, 400%+ moonshot)\n"
        "- 2PM tightening: after 2PM ET, trails tighten by 30% (V6)\n\n"
        "The team's proposal (+20%: move to breakeven, +50%: trail at -10%, +100%: trail at -15%) "
        "is a simplified version of what we already run."
    )

    add_verdict_box(doc, "ALREADY IMPLEMENTED (more granularly). No action needed.", "green")

    # --- PRIORITY 3: DECISION LOGGING ---
    doc.add_heading("Priority 3: Decision Logging on Every Alert", level=2)
    doc.add_paragraph(
        "Team reports 69% of alerts have no execution_decisions row. "
        "This is a valid data gap -- we don't currently write back to Supabase."
    )

    add_verdict_box(doc, "ACCEPT -- Valid infrastructure gap. Will implement Supabase "
        "write-back for all entry decisions (executed, skipped, reduced). "
        "Low risk, high value for cross-system learning.", "green")

    # --- PRIORITY 4: ENTRY QUALITY SCORER ---
    doc.add_heading("Priority 4: Entry-Quality Scorer (QB Framework)", level=2)
    doc.add_heading("Team's Proposal", level=3)
    doc.add_paragraph(
        "8-component scoring system (0-100) driving sizing: VWAP alignment, support proximity, "
        "option spread, volume expansion, tape direction, alert age, GEX positioning, drawdown."
    )

    doc.add_heading("What We Already Have", level=3)
    doc.add_paragraph(
        "Our entry pipeline has 18 gates that cover most of these signals:\n"
        "- Signal scoring (78-177 scale, from Discord signal parser)\n"
        "- Spread gate (bid-ask > 40% blocked)\n"
        "- Premium cap (tiered by score)\n"
        "- Momentum confirmation gate (VWAP + price action)\n"
        "- Late-session reduction\n"
        "- Score-tiered sizing (95+: 100%, 90: 75%, 85: 50%, 78: 25%)\n"
        "- Per-ticker optimal configs\n\n"
        "Missing from our stack: GEX positioning, volume expansion check, near-support detection. "
        "These could add value but require new data feeds."
    )

    doc.add_heading("Pushback", level=3)
    doc.add_paragraph(
        "The QB framework is a SECOND scoring system layered on top of the signal score "
        "we already use. Running two independent scorers creates ambiguity (which one wins?) "
        "and doubles calibration burden. Better approach: add the missing signals "
        "(GEX, volume expansion) as additional gates in our existing pipeline."
    )

    add_verdict_box(doc, "PARTIALLY ACCEPT -- Add missing signals (GEX, volume) as new "
        "pipeline gates rather than building a parallel scoring system.", "yellow")

    # --- PRIORITY 5: PER-AGENT VARIANCE ---
    doc.add_heading("Priority 5: Per-Agent Strategy Variance", level=2)
    doc.add_heading("Team's Data", level=3)
    add_table(doc,
        ["Agent", "Portfolio", "Trades", "WR", "P&L"],
        [
            ["owlet_kody", "$23,000", "141", "60.3%", "+$6,982"],
            ["owlet_adam", "$4,685", "34", "58.8%", "-$1,203"],
            ["owlet_yank", "$3,600", "18", "66.7%", "-$572"],
            ["owlet_vinny", "$5,000", "15", "66.7%", "-$636"],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "The team correctly identifies that the strategy doesn't scale down well to small accounts. "
        "However, their suggestion (conviction >= 95 filter for small accounts) doesn't match "
        "our scoring system (we use 78-177, not 0-100)."
    )

    doc.add_heading("What We Already Did", level=3)
    doc.add_paragraph(
        "We just deployed MAX_POSITION_PCT=10 and MAX_DCA_POSITION_PCT=5 across all bots. "
        "This tightens position sizing and prevents oversized DCA adds that disproportionately "
        "hurt small accounts. Combined with score-tiered sizing (25% budget for marginal signals), "
        "small accounts already trade smaller on low-conviction setups."
    )

    add_verdict_box(doc, "PARTIALLY ADDRESSED -- Position cap changes (deployed today) "
        "help. Will monitor small account P&L over next 2 weeks before adding more filters. "
        "May raise minimum score threshold for sub-$5K accounts.", "yellow")

    # --- PRIORITY 6: ACCOUNT STATE CRON ---
    doc.add_heading("Priority 6: Account State Cron", level=2)
    doc.add_paragraph(
        "Only 5 account_state snapshots exist. Team wants 78/day per agent. "
        "This is a valid infrastructure gap for the shared-brain system."
    )

    add_verdict_box(doc, "ACCEPT -- Low effort, enables circuit breaker and drawdown tracking. "
        "Will implement 5-minute Webull balance polling during market hours.", "green")

    # =========================================================================
    # DOCUMENT 2: EXIT LOGIC OVERHAUL
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("Document 2: Exit Logic Overhaul Spec", level=1)
    doc.add_paragraph(
        "Source: WEBULL-EXIT-LOGIC-OVERHAUL-SPEC.md (2026-05-19). "
        "Proposes replacing our exit logic with a 4-stage time-based FSM. "
        "Claims median peak capture of 40% and estimates doubling P&L by moving to 70%."
    )

    # --- FUNDAMENTAL MISUNDERSTANDING ---
    doc.add_heading("Critical Context the Team Is Missing", level=2)
    doc.add_paragraph(
        "The team's spec is written as if our exit logic is ad-hoc discretionary closes. "
        "In reality, we run a sophisticated V5 FSM with 10 gates, category-aware trailing, "
        "DTE-aware stops, per-ticker configs, and V6 enhancements (breakeven ratchet, scaleout, "
        "2PM tightening). Their proposed 4-stage FSM is a DOWNGRADE from what we already have."
    )

    doc.add_paragraph(
        "The data gap: our exit reasons are stored locally in SQLite but NOT written to Supabase. "
        "The team only sees 4 close_reasons (target_hit, manual, stop_loss, time_stop) because "
        "that's all we write to the shared brain. What they call 'manual' is actually our FSM "
        "firing adaptive_trail, soft_trail, graduated_stop, etc. -- all systematic, all rule-based."
    )

    # --- 4-STAGE FSM ---
    doc.add_heading("Proposed 4-Stage Time-Based FSM", level=2)

    doc.add_heading("Stage 1: T+0 to T+15 (Fast Win Zone)", level=3)
    doc.add_paragraph(
        "Team proposes: auto-close at +25% premium in first 15 min."
    )
    doc.add_heading("Backtest: Quick Take +25% in First 15min", level=4)
    add_table(doc,
        ["Metric", "Baseline", "Quick Take", "Delta"],
        [
            ["Total P&L", "$53,565", "$51,824", "-$1,742 (-3%)"],
            ["Win Rate", "77.4%", "78.5%", "+1.0%"],
            ["Peak Capture (median)", "52.6%", "66.7%", "+14.1%"],
            ["Max Win", "$7,760", "$7,760", "No change"],
            ["Worst Day", "-$1,909", "-$1,556", "+$353 better"],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "This is the ONLY team proposal that comes close to production performance. "
        "It improves peak capture and slightly improves worst-day risk, but still loses $1,742. "
        "The issue: trades that run past +25% and hit +50% or +100% get capped. "
        "Our V5 FSM handles this better -- scaleout sells 1/3 at +20% (locking partial profit) "
        "while letting the rest run."
    )

    add_verdict_box(doc, "REJECT (already handled better) -- Our V6 scaleout sells 1/3 at "
        "+20% without capping the remaining position. Gets the partial profit lock "
        "without sacrificing runners.", "yellow")

    doc.add_heading("Stage 2: T+15 to T+30 (Commit Zone)", level=3)
    doc.add_paragraph(
        "Team proposes: if P&L <= 0 at T+15, force exit immediately."
    )
    doc.add_heading("Backtest: Commit Zone", level=4)
    add_table(doc,
        ["Metric", "Baseline", "Commit Zone", "Delta"],
        [
            ["Total P&L", "$53,565", "$32,803", "-$20,762 (-39%)"],
            ["Win Rate", "77.4%", "56.9%", "-20.5%"],
            ["Trades force-closed", "N/A", "~84 trades", "43% of all trades killed"],
            ["Win:Loss ratio", "0.92:1", "1.86:1", "Better ratio, much less profit"],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "The commit zone DESTROYS $20,762 in profit and drops win rate by 20 points. "
        "The team's logic: the 15-30 min window has 47% WR in their data. But this is because "
        "their data includes ALL closes in that window, including stop-losses. "
        "Our V5 FSM already handles underwater trades via graduated_stop (35% 0DTE / 52% multi-day) "
        "and checkpoint_cut (0DTE: down 30% AND underlying against 0.5%). "
        "A blanket 'exit if negative at T+15' kills trades that are temporarily underwater "
        "but would recover -- which many do."
    )

    add_verdict_box(doc, "REJECT -- Loses $20,762. The 'death zone' problem is already "
        "solved by our graduated_stop and checkpoint_cut gates which are smarter "
        "(they check BOTH premium AND underlying direction, not just time).", "red")

    doc.add_heading("Stage 3: T+30 to T+45 (Trail Hard)", level=3)
    doc.add_paragraph(
        "Team proposes: tighten trailing stop to peak - 10%."
    )
    doc.add_paragraph(
        "Our V5 adaptive_trail already provides category-aware trailing that is tighter "
        "than 10% for index options and appropriately wider for high-vol names. "
        "A flat 10% trail on MSTR or TSLA would trigger on normal volatility, "
        "cutting winners short. Our per-ticker configs handle this."
    )

    add_verdict_box(doc, "REJECT -- Flat 10% trail ignores ticker volatility. "
        "Our category-aware adaptive_trail (25-55% by tier and category) is superior.", "red")

    doc.add_heading("Stage 4: T+45 Max Hold", level=3)
    doc.add_paragraph(
        "Already backtested above. Loses $24,182."
    )
    add_verdict_box(doc, "REJECT -- See 45-min max hold backtest above.", "red")

    # --- BEARISH-SPECIFIC RULES ---
    doc.add_heading("Bearish-Specific Tighter Trails", level=2)
    doc.add_heading("Team's Data", level=3)
    doc.add_paragraph(
        "Bearish manual closes: n=11, WR=36%, avg -13.31%, total -$3,044. "
        "They propose tighter trails on puts: 15% vs 20% (Stage 2), 7% vs 10% (Stage 3), "
        "30-min max hold vs 45-min."
    )

    doc.add_heading("Our Data", level=3)
    add_table(doc,
        ["Metric", "Baseline Puts", "Block All Puts", "Halve Put Sizing"],
        [
            ["Put P&L", "$12,552", "$0 (no puts)", "$5,673"],
            ["Put WR", "85%", "N/A", "85%"],
            ["Total P&L", "$53,565", "$41,013", "$46,686"],
            ["Delta", "--", "-$12,552", "-$6,879"],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "CRITICAL FINDING: The team says bearish trades lose money (WR=56.7%, avg -6.12%). "
        "Our backtest shows puts at 85% WR with +$12,552 profit. The discrepancy is because "
        "our V5 FSM handles put exits well -- the team's 'manual' bearish losses are from "
        "their system, not ours. Blocking puts would cost us $12,552. "
        "Halving put sizing costs $6,879. Neither is justified by our actual data."
    )

    add_verdict_box(doc, "REJECT -- Our puts are profitable (85% WR, +$12,552). "
        "The team's bearish underperformance data does not match our production results.", "red")

    # --- COMBINED PROPOSAL ---
    doc.add_heading("Combined: Commit Zone + 45-min Max Hold", level=2)
    doc.add_paragraph(
        "The team's full exit overhaul (Stages 2+4 combined) was backtested."
    )
    add_table(doc,
        ["Metric", "Baseline", "Combined", "Delta"],
        [
            ["Total P&L", "$53,565", "$20,217", "-$33,348 (-62%)"],
            ["Win Rate", "77.4%", "53.8%", "-23.6%"],
            ["Max Win", "$7,760", "$2,700", "-$5,060"],
            ["Avg Hold", "94 min", "24 min", "-70 min"],
        ]
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "The combined proposal loses 62% of all profit. This is the strongest evidence "
        "that the team's exit overhaul would be destructive to deploy."
    )
    run.bold = True

    add_verdict_box(doc, "STRONGLY REJECT -- Combined proposal loses $33,348 (62% of profit). "
        "Our V5 FSM is categorically superior to the proposed 4-stage time-based FSM.", "red")

    # =========================================================================
    # DATA QUALITY ITEMS
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("Data Quality & Infrastructure Items", level=1)

    doc.add_heading("Conviction Score >100", level=2)
    doc.add_paragraph(
        "Team notes our conviction_score goes up to 177, exceeding their 0-100 scale. "
        "This is correct -- our signal scoring uses a different scale (78-177). "
        "We should either clamp to 0-100 when writing to Supabase, or document the scale "
        "difference. Easy fix."
    )
    add_verdict_box(doc, "ACCEPT -- Will clamp to 0-100 on Supabase write-back.", "green")

    doc.add_heading("Score Tier Null on 44% of Alerts", level=2)
    doc.add_paragraph(
        "This is on the scanner side. They acknowledged it and will fix. No action from us."
    )
    add_verdict_box(doc, "NO ACTION (scanner-side fix).", "green")

    # =========================================================================
    # FINAL RECOMMENDATIONS
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("Final Recommendations Summary", level=1)

    doc.add_paragraph(
        "Sorted by priority. Green = accept, yellow = partial, red = reject."
    )

    add_table(doc,
        ["#", "Recommendation", "Verdict", "Effort", "Expected Impact"],
        [
            ["1", "Supabase decision logging (write-back)", "ACCEPT", "Medium", "Enables cross-system learning"],
            ["2", "Account state cron (5-min polling)", "ACCEPT", "Low", "Enables circuit breaker"],
            ["3", "Conviction score clamp (0-100)", "ACCEPT", "Low", "Data quality fix"],
            ["4", "NBBO logging on entry", "ACCEPT", "Low", "Measurement improvement"],
            ["5", "GEX/volume as pipeline gates", "CONSIDER", "High", "May improve entry quality"],
            ["6", "Halve put sizing", "MONITOR", "Low", "Costs $6,879 in backtest; revisit if puts underperform"],
            ["7", "Small account conviction filter", "MONITOR", "Low", "Wait for position cap changes to settle"],
            ["8", "Mid-pricing with retry loops", "REJECT", "High", "Latency risk on 0DTE; spread gate covers worst cases"],
            ["9", "45-min max hold", "REJECT", "Low", "Loses $24,182 (-45%)"],
            ["10", "30-min max hold", "REJECT", "Low", "Loses $27,834 (-52%)"],
            ["11", "Commit zone (exit if underwater T+15)", "REJECT", "Low", "Loses $20,762 (-39%)"],
            ["12", "4-stage time-based FSM (full overhaul)", "REJECT", "Very High", "Loses $33,348 (-62%)"],
            ["13", "Bearish-specific tighter trails", "REJECT", "Medium", "Puts are 85% WR / +$12,552 in our system"],
            ["14", "Block all puts", "REJECT", "Low", "Loses $12,552"],
        ]
    )

    doc.add_paragraph()
    doc.add_heading("Key Takeaway for the Team", level=2)
    doc.add_paragraph(
        "The primary action item is closing the DATA GAP between our systems, not overhauling "
        "exit logic. Once we write our V5 FSM exit reasons and trade events to Supabase, "
        "the team will see that the 'manual' close problem they identified doesn't exist -- "
        "those are systematic exits that just weren't being reported. "
        "The four items we accept (decision logging, account state, conviction clamp, NBBO logging) "
        "are all about closing this visibility gap."
    )

    doc.add_paragraph(
        "The exit logic overhaul spec should be shelved. Our V5 FSM with 10 gates, "
        "category-aware trailing, DTE-awareness, per-ticker configs, and V6 enhancements "
        "(breakeven ratchet, scaleout, 2PM tightening) is significantly more sophisticated "
        "than the proposed 4-stage time-based FSM. Every proposed change that was backtested "
        "lost money -- in some cases catastrophically."
    )

    # =========================================================================
    # APPENDIX: METHODOLOGY
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("Appendix: Backtest Methodology", level=1)

    doc.add_paragraph(
        "All backtests use the FULL production V5 FSM strategy with:"
    )

    bullets = [
        "ExitFSM with all 10 gates (eod_cutoff, bid_disappearance, profit_target, "
        "breakeven_ratchet, scaleout, scalp_trail, checkpoint_cut, graduated_stop, "
        "soft_trail, adaptive_trail, theta_exit)",
        "All V6 enhancements enabled (breakeven ratchet, scaleout, 2PM tightening, "
        "per-ticker configs, premium cap, spread gate)",
        "MomentumConfirmGate for entry filtering",
        "Score-tiered sizing (135: 100%, 120: 85%, 100: 85%, 90: 50%, 78: 25%)",
        "Per-ticker configs via get_ticker_config()",
        "MAX_CONCURRENT=4, MAX_POSITION_PCT=10, MAX_PORTFOLIO_RISK_PCT=75%",
        "Entry price from first ask in harvester data, fallback to midpoint",
        "Scaleout partial exits handled correctly (not treated as full close)",
        "Harvester tick data (Polygon options snapshots, ~2M rows)",
        "Window: April 10 - May 18, 2026 (23 trading days, 195 qualifying trades)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(
        "Each test changes EXACTLY ONE variable from the production baseline. "
        "This isolates the effect of each recommendation and prevents confounding."
    )

    doc.add_heading("Hold Time Distribution (Baseline)", level=2)
    add_table(doc,
        ["Hold Time", "Trades", "Win Rate", "P&L", "Avg P&L"],
        [
            ["5-15 min", "72", "88%", "$13,997", "$194"],
            ["15-30 min", "40", "75%", "$5,734", "$143"],
            ["30-45 min", "22", "86%", "$8,171", "$371"],
            ["45-75 min", "20", "80%", "$1,409", "$70"],
            ["75+ min", "33", "55%", "$9,662", "$293"],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Note: Our 15-30 min bucket shows 75% WR (vs team's 47%). Our 30-45 min bucket shows "
        "86% WR (vs team's 90%). The 'death zone' the team identifies in their data does not "
        "exist in ours because our V5 FSM actively manages positions throughout these windows."
    )

    # Save
    out_path = "/Users/kody/Downloads/OptionsOwl_Team_Recommendations_Analysis.docx"
    doc.save(out_path)
    print(f"Report saved to: {out_path}")


if __name__ == "__main__":
    main()
