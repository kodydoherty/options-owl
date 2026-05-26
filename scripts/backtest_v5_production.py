"""Backtest using the LIVE production v5 FSM code against historical Discord signals.

This uses the actual ExitFSM class deployed to all owlets — not a separate
simulation function. If the FSM code changes, this backtest automatically
reflects those changes.

Outputs:
  - Per-trade results table
  - Daily P&L summary
  - Daily cumulative P&L chart (saved as PNG)

Usage:
    python scripts/backtest_v5_production.py
"""

from __future__ import annotations

import sqlite3
import sys
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path

import numpy as np
import pandas as pd

# Add project root to path
PROJECT_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_DIR))

from types import SimpleNamespace

from zoneinfo import ZoneInfo

from options_owl.risk.exit_v5.config import V5Config, get_ticker_config
from options_owl.risk.exit_v5.fsm import ExitFSM, TradeState
from options_owl.collectors.support_levels import find_support_levels, is_at_support

ET = ZoneInfo("America/New_York")
UTC = ZoneInfo("UTC")

# Mock settings matching production docker-compose V6 flags (ALL enabled flags)
_V6_SETTINGS = SimpleNamespace(
    ENABLE_V6_BREAKEVEN_RATCHET=True,
    V6_BREAKEVEN_TRIGGER_PCT=20.0,
    ENABLE_V6_SCALEOUT=True,
    V6_SCALEOUT_GAIN_PCT=20.0,
    V6_SCALEOUT_FRACTION=0.333,
    V6_SCALEOUT_MIN_CONTRACTS=3,
    ENABLE_V6_2PM_TIGHTEN=True,
    V6_2PM_TRAIL_TIGHTEN_FACTOR=0.7,
    V6_2PM_SOFT_TRAIL_BOOST=0.15,
    ENABLE_V6_PER_TICKER_CONFIG=True,
    # V6 entry gates (applied in backtest entry filter, not FSM)
    ENABLE_V6_PREMIUM_CAP=True,
    V6_PREMIUM_CAP=6.0,
    V6_PREMIUM_CAP_MID=7.0,    # score 120+
    V6_PREMIUM_CAP_HIGH=9.0,   # score 150+
    ENABLE_V6_SPREAD_GATE=True,
    V6_MAX_SPREAD_PCT=15.0,
    ENABLE_V6_EARLY_POP_GATE=True,
    ENABLE_V6_DCA=True,
    V6_DCA_TICKERS="MSFT,IWM,SPY,QQQ,AMZN,NVDA",
    V6_DCA_MIN_MINUTES=8.0,
    V6_DCA_MAX_MINUTES=20.0,
    V6_DCA_MIN_DIP_PCT=15.0,
    V6_DCA_MAX_DIP_PCT=35.0,
    V6_DCA_UNDERLYING_THRESHOLD=0.5,
    # Scalp target gate — take +25% unless confirmed runner
    ENABLE_SCALP_TARGET=True,
    SCALP_TARGET_PCT=25.0,
    SCALP_RUNNER_CONFIRM_PCT=40.0,
    # Sideways scalp
    ENABLE_V6_SIDEWAYS_SCALP=True,
)

SIGNALS_DB = str(PROJECT_DIR / "journal" / "owlet-kody" / "raw_messages.db")
HARVESTER_DB = str(PROJECT_DIR / "journal" / "owlet-harvester" / "options_data.db")

POLYGON_API_KEY = "Vk7gXTz6dbp_F69UmmqIx9BDEasHfExb"
PORTFOLIO = 20000

# Morning cutoff: block entries after configured time ET
# Set via command line: --morning-cutoff 11:00 (or --no-morning-cutoff)
ENABLE_MORNING_CUTOFF = False  # Default OFF for backtest comparison
MORNING_CUTOFF_HOUR = 11
MORNING_CUTOFF_MINUTE = 0

# Toggle VWAP+Support gate A/B test
ENABLE_VWAP_SUPPORT_GATE = True


# ── VWAP + Multi-TF Support Gate helpers ─────────────────────────────────────

import time as _time
import requests as _requests

_polygon_cache: dict[tuple[str, str], list[dict]] = {}


@dataclass(slots=True)
class _CandleBar:
    timestamp: float
    open: float
    high: float
    low: float
    close: float
    volume: float
    vwap: float = 0.0


def _fetch_polygon_5m(ticker: str, from_date: str, to_date: str) -> list[dict]:
    """Fetch 5m candles from Polygon with caching."""
    key = (ticker, from_date)
    if key in _polygon_cache:
        return _polygon_cache[key]
    url = (
        f"https://api.polygon.io/v2/aggs/ticker/{ticker}/range/5/minute"
        f"/{from_date}/{to_date}"
        f"?adjusted=true&sort=asc&limit=50000&apiKey={POLYGON_API_KEY}"
    )
    for attempt in range(3):
        try:
            resp = _requests.get(url, timeout=30)
            if resp.status_code == 429:
                _time.sleep(12)
                continue
            results = resp.json().get("results", [])
            _polygon_cache[key] = results
            return results
        except Exception:
            _time.sleep(2)
    _polygon_cache[key] = []
    return []


def _aggregate_to_tf(bars: list[_CandleBar], tf_minutes: int) -> list[_CandleBar]:
    bucket_ms = tf_minutes * 60 * 1000
    buckets: dict[int, list[_CandleBar]] = {}
    for b in bars:
        bk = int((b.timestamp // bucket_ms) * bucket_ms)
        buckets.setdefault(bk, []).append(b)
    out = []
    for bk_ts in sorted(buckets):
        g = buckets[bk_ts]
        out.append(_CandleBar(
            timestamp=bk_ts, open=g[0].open,
            high=max(b.high for b in g), low=min(b.low for b in g),
            close=g[-1].close, volume=sum(b.volume for b in g),
            vwap=g[-1].vwap if g[-1].vwap else 0.0,
        ))
    return out


def _check_vwap_support(ticker: str, date_str: str, entry_ts_ms: int, direction: str):
    """Check VWAP and multi-TF support at entry time.

    Returns (above_vwap, at_support, detail_str) or None if no data.
    """
    dt = datetime.strptime(date_str, "%Y-%m-%d")
    from_date = (dt - timedelta(days=7)).strftime("%Y-%m-%d")
    raw = _fetch_polygon_5m(ticker, from_date, date_str)
    if not raw:
        return None
    _time.sleep(0.15)

    all_5m = [
        _CandleBar(
            timestamp=c["t"], open=c["o"], high=c["h"], low=c["l"],
            close=c["c"], volume=c.get("v", 0), vwap=c.get("vw", 0),
        )
        for c in raw if c["t"] <= entry_ts_ms
    ]
    if len(all_5m) < 6:
        return None

    entry_price = all_5m[-1].close
    candle_data = {
        "5m": all_5m,
        "15m": _aggregate_to_tf(all_5m, 15),
        "1h": _aggregate_to_tf(all_5m, 60),
        "4h": _aggregate_to_tf(all_5m, 240),
    }

    # VWAP from today's session
    today_start_ms = dt.replace(hour=0, minute=0).timestamp() * 1000
    total_vp = total_vol = 0.0
    for b in all_5m:
        if b.timestamp < today_start_ms or b.volume <= 0:
            continue
        typical = (b.high + b.low + b.close) / 3
        total_vp += typical * b.volume
        total_vol += b.volume
    vwap = total_vp / total_vol if total_vol > 0 else None
    above_vwap = entry_price >= vwap if vwap else None

    at_support_result, support_detail = is_at_support(
        candle_data, current_price=entry_price,
        max_distance_pct=0.3, min_strength=3, min_confluence=1,
    )

    return above_vwap, at_support_result, f"vwap={'above' if above_vwap else 'below'} support={at_support_result} {support_detail[:60]}"


# ── Data loading (reused from backtest_ml_v5b.py) ────────────────────────────


def load_signals():
    conn = sqlite3.connect(SIGNALS_DB)
    conn.row_factory = sqlite3.Row
    rows = conn.execute("""
        SELECT id, ticker, direction, sentiment, score,
               atm_premium, otm_premium, strike, expiry,
               entry_price, created_at
        FROM trade_signals
        WHERE score >= 70
        ORDER BY created_at
    """).fetchall()
    signals = []
    for r in rows:
        sig = dict(r)
        sig["premium"] = sig["atm_premium"] or sig["otm_premium"]
        sent = (sig.get("sentiment") or sig.get("direction") or "bullish").lower()
        sig["option_type"] = "put" if sent in ("bearish", "put") else "call"
        if sig["premium"] and sig["premium"] > 0 and sig["strike"]:
            signals.append(sig)
    conn.close()
    return signals


def build_contract_ticker(ticker, expiry, strike, option_type):
    if not expiry:
        return ""
    try:
        exp_dt = datetime.strptime(expiry, "%Y-%m-%d")
    except ValueError:
        return ""
    exp_str = exp_dt.strftime("%y%m%d")
    ot = "C" if option_type.lower() in ("call", "bullish", "c") else "P"
    strike_int = int(strike * 1000)
    return f"O:{ticker}{exp_str}{ot}{strike_int:08d}"


def load_ticks(harvester_conn, signal):
    ticker = signal["ticker"]
    strike = signal["strike"]
    created_at = signal["created_at"]
    option_type = signal["option_type"]
    sig_date = created_at[:10]

    sig_dt = datetime.strptime(sig_date, "%Y-%m-%d").date()
    candidates = [sig_dt]
    for delta in range(1, 6):
        d = sig_dt + timedelta(days=delta)
        if d.weekday() < 5:
            candidates.append(d)
            if len(candidates) >= 4:
                break

    for exp_date in candidates:
        expiry = exp_date.strftime("%Y-%m-%d")
        ct = build_contract_ticker(ticker, expiry, strike, option_type)
        if not ct:
            continue
        rows = harvester_conn.execute("""
            SELECT captured_at, midpoint, bid, ask, underlying_price,
                   implied_volatility, delta, gamma, theta, vega,
                   day_volume
            FROM harvest_snapshots
            WHERE contract_ticker = ? AND captured_at >= ?
            ORDER BY captured_at
        """, (ct, created_at)).fetchall()
        if rows and len(rows) >= 10:
            signal["_dte"] = (exp_date - sig_dt).days
            signal["_expiry_date"] = expiry
            break
    else:
        return None

    df = pd.DataFrame(rows, columns=[
        "captured_at", "midpoint", "bid", "ask", "underlying_price",
        "iv", "delta", "gamma", "theta", "vega", "volume"
    ])
    df["premium"] = df["midpoint"].where(df["midpoint"] > 0, (df["bid"] + df["ask"]) / 2)
    df["premium"] = df["premium"].where(df["premium"] > 0, np.nan)
    df = df.dropna(subset=["premium"])
    if len(df) < 10:
        return None
    df["ts"] = pd.to_datetime(df["captured_at"])
    df = df.sort_values("ts").reset_index(drop=True)
    return df


# ── Production FSM simulation ────────────────────────────────────────────────


def simulate_with_production_fsm(df, entry_premium, contracts, direction, dte, expiry_date, ticker="SIM"):
    """Run the ACTUAL production ExitFSM against tick data.

    This is the code running on all owlets right now.
    """
    if entry_premium <= 0:
        return {"pnl": 0, "reason": "no_data", "hold": 0, "exit_prem": 0, "peak_gain": 0}

    cfg = get_ticker_config(ticker, use_per_ticker=True)
    fsm = ExitFSM(cfg, settings=_V6_SETTINGS)
    option_type = "put" if direction in ("bearish", "put") else "call"

    # Create TradeState matching what monitor_bridge does
    entry_ts = df["ts"].iloc[0]
    if hasattr(entry_ts, 'to_pydatetime'):
        entry_ts = entry_ts.to_pydatetime()
    if entry_ts.tzinfo is not None:
        entry_ts = entry_ts.replace(tzinfo=None)

    first_underlying = 0.0
    for i in range(min(5, len(df))):
        u = df["underlying_price"].iloc[i]
        if u and u > 0:
            first_underlying = float(u)
            break

    state = TradeState(
        trade_id=1,
        ticker=ticker,
        option_type=option_type,
        entry_premium=entry_premium,
        entry_time=entry_ts,
        contracts=contracts,
        peak_premium=entry_premium,
        entry_underlying_price=first_underlying,
        dte=dte,
        expiry_date=expiry_date or "",
    )

    # Track partial exits (V6 scaleout locks in profit on some contracts)
    locked_pnl = 0.0
    remaining = contracts

    for idx in range(1, len(df)):
        premium = df["premium"].iloc[idx]
        if np.isnan(premium) or premium <= 0:
            continue

        raw_bid = df["bid"].iloc[idx]
        raw_ask = df["ask"].iloc[idx]
        bid = float(raw_bid) if raw_bid and not pd.isna(raw_bid) else premium
        ask = float(raw_ask) if raw_ask and not pd.isna(raw_ask) else premium

        now = df["ts"].iloc[idx]
        if hasattr(now, 'to_pydatetime'):
            now = now.to_pydatetime()
        if now.tzinfo is not None:
            now = now.replace(tzinfo=None)

        underlying = df["underlying_price"].iloc[idx] or 0.0

        # Compute minutes_to_close (market closes at 16:00 ET)
        # Use proper timezone conversion (handles DST correctly)
        now_utc = now.replace(tzinfo=UTC) if now.tzinfo is None else now
        now_et = now_utc.astimezone(ET)
        et_hour = now_et.hour
        et_minute = now_et.minute
        minutes_to_close = max(0, (16 * 60) - (et_hour * 60 + et_minute))
        # Pass naive ET time to FSM (it expects naive datetimes)
        now = now_et.replace(tzinfo=None)

        action = fsm.evaluate(
            state, premium, bid, ask, now,
            current_underlying=underlying,
            minutes_to_close=minutes_to_close,
        )

        if action.should_exit:
            # V6 scaleout: partial exit — lock in profit, continue with remaining
            if action.contracts_to_close > 0 and action.contracts_to_close < remaining:
                closed = action.contracts_to_close
                locked_pnl += (premium - entry_premium) * closed * 100
                remaining -= closed
                state.contracts = remaining
                continue

            elapsed = (now - entry_ts).total_seconds() / 60
            peak_gain = (state.peak_premium - entry_premium) / entry_premium * 100
            pnl = locked_pnl + (premium - entry_premium) * remaining * 100
            return {
                "pnl": pnl,
                "reason": action.reason.value,
                "hold": elapsed,
                "exit_prem": premium,
                "peak_gain": peak_gain,
            }

    # End of data — force close at last tick
    last_prem = df["premium"].iloc[-1]
    last_ts = df["ts"].iloc[-1]
    if hasattr(last_ts, 'to_pydatetime'):
        last_ts = last_ts.to_pydatetime()
    if last_ts.tzinfo is not None:
        last_ts = last_ts.replace(tzinfo=None)
    elapsed = (last_ts - entry_ts).total_seconds() / 60
    peak_gain = (state.peak_premium - entry_premium) / entry_premium * 100
    pnl = locked_pnl + (last_prem - entry_premium) * remaining * 100
    return {
        "pnl": pnl,
        "reason": "eod_data_end",
        "hold": elapsed,
        "exit_prem": last_prem,
        "peak_gain": peak_gain,
    }


# ── Main ──────────────────────────────────────────────────────────────────────


def main():
    signals = load_signals()
    print(f"Loaded {len(signals)} signals from DB")

    harvester_conn = sqlite3.connect(HARVESTER_DB)
    results = []
    no_data = 0

    for sig in signals:
        ticker = sig["ticker"]
        direction = (sig["direction"] or "bullish").lower()
        score = sig["score"] or 80
        day = sig["created_at"][:10]
        entry_premium = sig["premium"]

        print(f"  [{len(results)+no_data+1}/{len(signals)}] {ticker} {direction} {day}...", end="", flush=True)

        df = load_ticks(harvester_conn, sig)
        if df is None:
            no_data += 1
            print(" no tick data")
            continue

        dte = sig.get("_dte", 0)
        expiry_date = sig.get("_expiry_date", "")

        # Use actual market price from harvester
        first_ask = df["ask"].iloc[0]
        first_mid = df["premium"].iloc[0]
        adj_entry = first_ask if first_ask and first_ask > 0 else first_mid
        if adj_entry <= 0:
            adj_entry = entry_premium

        # Dollar-target sizing — MUST match production docker-compose.yml
        max_risk_pct = 0.75
        max_concurrent = 4    # production: MAX_CONCURRENT=4
        max_position_pct = 0.15  # production: MAX_POSITION_PCT=15
        deployable = PORTFOLIO * max_risk_pct
        per_slot = deployable / max_concurrent

        # Production score tiers (vinny_strategy.py _SCORE_TIER_TABLE)
        # Includes per-tier position cap % — matches production exactly.
        SCORE_TIERS = [
            # (min_score, budget_mult, position_cap_pct)
            (135, 1.00, 0.15),  # elite:    100% of slot, up to 15% of portfolio
            (120, 0.85, 0.12),  # strong:   85% of slot, up to 12%
            (100, 0.85, 0.08),  # standard: 85% of slot, up to 8% (bulk of signals)
            (90, 0.50, 0.08),   # moderate: 50% of slot, up to 8%
            (78, 0.25, 0.08),   # marginal: 25% of slot, up to 8%
        ]
        score_mult = 0.25  # fallback below 78
        tier_pos_pct = 0.08
        for threshold, mult, pos_pct in SCORE_TIERS:
            if score >= threshold:
                score_mult = mult
                tier_pos_pct = pos_pct
                break

        # Production score floor = 78
        if score < 78:
            continue

        # Morning cutoff: block entries after 11:00 AM ET
        if ENABLE_MORNING_CUTOFF:
            sig_time = sig["created_at"]
            try:
                sig_dt_full = (
                    datetime.strptime(sig_time[:19], "%Y-%m-%dT%H:%M:%S")
                    if "T" in sig_time
                    else datetime.strptime(sig_time[:19], "%Y-%m-%d %H:%M:%S")
                )
                # DB times are UTC — convert properly to ET (handles DST)
                sig_utc = sig_dt_full.replace(tzinfo=UTC)
                sig_et = sig_utc.astimezone(ET)
                cutoff_minutes = MORNING_CUTOFF_HOUR * 60 + MORNING_CUTOFF_MINUTE
                signal_minutes = sig_et.hour * 60 + sig_et.minute
                if signal_minutes >= cutoff_minutes:
                    continue  # Skip — after morning cutoff
            except (ValueError, TypeError):
                pass

        # V6 premium cap gate (production: ENABLE_V6_PREMIUM_CAP=true)
        if _V6_SETTINGS.ENABLE_V6_PREMIUM_CAP:
            cap = _V6_SETTINGS.V6_PREMIUM_CAP
            if score >= 150:
                cap = _V6_SETTINGS.V6_PREMIUM_CAP_HIGH
            elif score >= 120:
                cap = _V6_SETTINGS.V6_PREMIUM_CAP_MID
            if adj_entry > cap:
                continue

        # V6 spread gate (production: ENABLE_V6_SPREAD_GATE=true)
        if _V6_SETTINGS.ENABLE_V6_SPREAD_GATE and len(df) > 0:
            first_bid = df["bid"].iloc[0]
            first_ask = df["ask"].iloc[0]
            if first_bid and first_ask and first_bid > 0 and first_ask > 0:
                spread_pct = (first_ask - first_bid) / first_ask * 100
                if spread_pct > _V6_SETTINGS.V6_MAX_SPREAD_PCT:
                    continue

        # Use the LARGER of tier cap and settings override (matches production)
        effective_pos_pct = max(tier_pos_pct, max_position_pct)
        position_cap = PORTFOLIO * effective_pos_pct

        cost_per = adj_entry * 100
        scaled_target = per_slot * score_mult
        raw_contracts = int(scaled_target / cost_per) if cost_per > 0 else 1
        pos_cap_contracts = int(position_cap / cost_per) if cost_per > 0 else 1
        if pos_cap_contracts == 0:
            continue
        contracts = max(1, min(raw_contracts, pos_cap_contracts))

        # Late-session 0DTE size reduction (matches production paper_trader)
        if dte == 0 and contracts > 1:
            sig_time = sig["created_at"]
            try:
                sig_dt_full = datetime.strptime(sig_time[:19], "%Y-%m-%dT%H:%M:%S") if "T" in sig_time else datetime.strptime(sig_time[:19], "%Y-%m-%d %H:%M:%S")
                # DB times are UTC — convert properly to ET (handles DST)
                sig_utc = sig_dt_full.replace(tzinfo=UTC)
                sig_et = sig_utc.astimezone(ET)
                if sig_et.hour >= 14:  # 2 PM ET or later
                    contracts = 1
                elif sig_et.hour >= 13:  # 1 PM ET or later
                    contracts = max(1, contracts // 2)
            except (ValueError, TypeError):
                pass

        # ── Simulate MomentumConfirmGate ────────────────────────────────
        # Check underlying price trend in the first ~15 min of ticks
        # to mimic what 5m candle data would show at entry time.
        is_call = direction in ("bullish", "call")
        momentum_blocked = False
        momentum_reason = ""

        # Look at underlying prices in the first 15 ticks (~15 min)
        window = min(15, len(df))
        underlying_prices = []
        for i in range(window):
            u = df["underlying_price"].iloc[i]
            if u and u > 0:
                underlying_prices.append(float(u))

        if len(underlying_prices) >= 5:
            # Simulate RSI-like check: is the underlying trending against us?
            first_half = underlying_prices[:len(underlying_prices)//2]
            second_half = underlying_prices[len(underlying_prices)//2:]
            avg_first = sum(first_half) / len(first_half)
            avg_second = sum(second_half) / len(second_half)
            pct_move = (avg_second - avg_first) / avg_first * 100

            # Check premium trend (first 5 ticks)
            prem_start = df["premium"].iloc[0]
            prem_5 = df["premium"].iloc[min(4, len(df)-1)]
            prem_fade = (prem_5 - prem_start) / prem_start * 100 if prem_start > 0 else 0

            neg_signals = 0

            # Signal 1: Underlying moving against direction
            if is_call and pct_move < -0.05:
                neg_signals += 1
                momentum_reason += f"underlying fading ({pct_move:+.2f}%); "
            elif not is_call and pct_move > 0.05:
                neg_signals += 1
                momentum_reason += f"underlying rising ({pct_move:+.2f}%); "

            # Signal 2: Premium fading in first 5 ticks
            if prem_fade < -5:
                neg_signals += 1
                momentum_reason += f"premium fading ({prem_fade:+.1f}%); "

            # Signal 3: All recent price bars against direction (3/3)
            against = 0
            for i in range(max(0, window-3), window):
                if i == 0:
                    continue
                prev_u = df["underlying_price"].iloc[i-1]
                cur_u = df["underlying_price"].iloc[i]
                if prev_u and cur_u:
                    if is_call and cur_u < prev_u:
                        against += 1
                    elif not is_call and cur_u > prev_u:
                        against += 1
            if against >= 3:
                neg_signals += 1
                momentum_reason += f"3/3 bars against; "

            if neg_signals >= 2:
                momentum_blocked = True

        # ── VWAP + Multi-TF Support Gate ─────────────────────────────────
        vwap_blocked = False
        vwap_detail = ""
        if ENABLE_VWAP_SUPPORT_GATE and not momentum_blocked:
            sig_time = sig["created_at"]
            try:
                sig_dt_full = (
                    datetime.strptime(sig_time[:19], "%Y-%m-%dT%H:%M:%S")
                    if "T" in sig_time
                    else datetime.strptime(sig_time[:19], "%Y-%m-%d %H:%M:%S")
                )
                entry_ts_ms = int(sig_dt_full.timestamp() * 1000)
                vs = _check_vwap_support(ticker, day, entry_ts_ms, direction)
                if vs is not None:
                    above_vwap, at_support_result, detail = vs
                    # Gate: block if NEITHER above_vwap NOR at_support
                    if not above_vwap and not at_support_result:
                        vwap_blocked = True
                        vwap_detail = detail
                else:
                    print(f"    [VWAP] {ticker} {day}: no data returned")
            except Exception as e:
                print(f"    [VWAP] {ticker} {day}: error: {e}")

        blk_str = ""
        if momentum_blocked:
            blk_str = " MOM_BLOCKED"
        elif vwap_blocked:
            blk_str = " VWAP_BLOCKED"
        print(f" {contracts}ct ${adj_entry:.2f}{blk_str}")

        result_base = {
            "ticker": ticker,
            "day": day,
            "score": score,
            "entry": adj_entry,
            "contracts": contracts,
            "direction": direction,
            "dte": dte,
            "momentum_blocked": momentum_blocked,
            "momentum_reason": momentum_reason.rstrip("; "),
            "vwap_blocked": vwap_blocked,
            "vwap_detail": vwap_detail,
        }

        if momentum_blocked:
            # Record what WOULD have happened (for comparison)
            sim_result = simulate_with_production_fsm(
                df, adj_entry, contracts, direction, dte, expiry_date, ticker=ticker
            )
            result_base.update({
                "pnl": 0,
                "reason": "momentum_blocked",
                "hold": 0,
                "exit_prem": adj_entry,
                "peak_gain": 0,
                "would_have_pnl": sim_result["pnl"],
                "would_have_reason": sim_result["reason"],
            })
            results.append(result_base)
            continue

        if vwap_blocked:
            sim_result = simulate_with_production_fsm(
                df, adj_entry, contracts, direction, dte, expiry_date, ticker=ticker
            )
            result_base.update({
                "pnl": 0,
                "reason": "vwap_support_blocked",
                "hold": 0,
                "exit_prem": adj_entry,
                "peak_gain": 0,
                "would_have_pnl": sim_result["pnl"],
                "would_have_reason": sim_result["reason"],
            })
            results.append(result_base)
            continue

        result = simulate_with_production_fsm(
            df, adj_entry, contracts, direction, dte, expiry_date, ticker=ticker
        )
        result_base.update(result)
        result_base["would_have_pnl"] = result["pnl"]
        result_base["would_have_reason"] = result["reason"]
        results.append(result_base)

    harvester_conn.close()

    if not results:
        print("No results — check that signals and harvester DBs exist")
        return

    # ── Results table ─────────────────────────────────────────────────────

    df_results = pd.DataFrame(results)
    pnls = df_results["pnl"]
    wins = (pnls > 0).sum()
    losses = (pnls <= 0).sum()
    total_pnl = pnls.sum()
    win_rate = wins / len(pnls) * 100

    print(f"\n{'=' * 90}")
    print(f"PRODUCTION V5 FSM BACKTEST — {len(results)} trades, {no_data} skipped (no tick data)")
    print(f"{'=' * 90}")
    print(f"Total P&L:   ${total_pnl:,.2f}")
    print(f"Win Rate:    {win_rate:.1f}% ({wins}W / {losses}L)")
    print(f"Avg Win:     ${pnls[pnls > 0].mean():,.2f}" if wins > 0 else "Avg Win:     N/A")
    print(f"Avg Loss:    ${pnls[pnls <= 0].mean():,.2f}" if losses > 0 else "Avg Loss:    N/A")
    print(f"Avg Hold:    {df_results['hold'].mean():.0f} min")
    print(f"Max Win:     ${pnls.max():,.2f}")
    print(f"Max Loss:    ${pnls.min():,.2f}")

    # ── Exit reason breakdown ─────────────────────────────────────────────

    print(f"\n{'Reason':<25} {'Count':>6} {'Total P&L':>12} {'Avg P&L':>10} {'Win%':>6}")
    print("-" * 62)
    for reason, group in df_results.groupby("reason"):
        gpnl = group["pnl"]
        gwins = (gpnl > 0).sum()
        gwr = gwins / len(gpnl) * 100
        print(f"{reason:<25} {len(gpnl):>6} ${gpnl.sum():>10,.2f} ${gpnl.mean():>8,.2f} {gwr:>5.0f}%")

    # ── Momentum Gate Analysis ────────────────────────────────────────────

    blocked = df_results[df_results["momentum_blocked"] == True]
    passed = df_results[df_results["momentum_blocked"] == False]

    print(f"\n{'=' * 90}")
    print(f"MOMENTUM CONFIRM GATE ANALYSIS")
    print(f"{'=' * 90}")
    print(f"Total signals:     {len(df_results)}")
    print(f"Blocked by gate:   {len(blocked)}")
    print(f"Passed gate:       {len(passed)}")
    if len(blocked) > 0:
        avoided_pnl = blocked["would_have_pnl"].sum()
        avoided_wins = (blocked["would_have_pnl"] > 0).sum()
        avoided_losses = (blocked["would_have_pnl"] <= 0).sum()
        print(f"\nBlocked trades would have P&L: ${avoided_pnl:,.2f}")
        print(f"  Would-be wins:   {avoided_wins} (${blocked[blocked['would_have_pnl'] > 0]['would_have_pnl'].sum():,.2f})")
        print(f"  Would-be losses: {avoided_losses} (${blocked[blocked['would_have_pnl'] <= 0]['would_have_pnl'].sum():,.2f})")
        savings = -avoided_pnl if avoided_pnl < 0 else 0
        cost = avoided_pnl if avoided_pnl > 0 else 0
        if avoided_pnl < 0:
            print(f"  NET SAVINGS:     ${savings:,.2f} (gate saved money!)")
        else:
            print(f"  NET COST:        ${cost:,.2f} (gate blocked profitable trades)")

        print(f"\n{'Day':<12} {'Ticker':<6} {'Dir':<5} {'Score':>5} {'WouldPnL':>10} {'WouldReason':<25} {'BlockReason'}")
        print("-" * 110)
        for _, r in blocked.iterrows():
            print(f"{r['day']:<12} {r['ticker']:<6} {r['direction'][:4]:<5} {r['score']:>5} "
                  f"${r['would_have_pnl']:>8.2f} {r['would_have_reason']:<25} {r.get('momentum_reason', '')}")

    # Compare with/without momentum gate
    pnl_with_gate = total_pnl
    pnl_without_gate = df_results["would_have_pnl"].sum()
    print(f"\nP&L WITH momentum gate:    ${pnl_with_gate:,.2f}")
    print(f"P&L WITHOUT momentum gate: ${pnl_without_gate:,.2f}")
    print(f"Gate impact:               ${pnl_with_gate - pnl_without_gate:+,.2f}")

    # ── VWAP + Support Gate Analysis ─────────────────────────────────────

    if ENABLE_VWAP_SUPPORT_GATE:
        vwap_blocked = df_results[df_results.get("vwap_blocked", False) == True] if "vwap_blocked" in df_results.columns else pd.DataFrame()
        vwap_passed = df_results[df_results.get("vwap_blocked", False) != True] if "vwap_blocked" in df_results.columns else df_results

        print(f"\n{'=' * 90}")
        print(f"VWAP + MULTI-TF SUPPORT GATE ANALYSIS")
        print(f"  Gate: BLOCK if below VWAP AND not at wick-cluster support (3+ touches, 0.3%)")
        print(f"{'=' * 90}")
        print(f"Total signals:     {len(df_results)}")
        print(f"Blocked by VWAP gate: {len(vwap_blocked)}")

        if len(vwap_blocked) > 0:
            avoided_pnl = vwap_blocked["would_have_pnl"].sum()
            avoided_wins = (vwap_blocked["would_have_pnl"] > 0).sum()
            avoided_losses = (vwap_blocked["would_have_pnl"] <= 0).sum()
            print(f"\nBlocked trades would have P&L: ${avoided_pnl:,.2f}")
            print(f"  Would-be wins:   {avoided_wins} (${vwap_blocked[vwap_blocked['would_have_pnl'] > 0]['would_have_pnl'].sum():,.2f})")
            print(f"  Would-be losses: {avoided_losses} (${vwap_blocked[vwap_blocked['would_have_pnl'] <= 0]['would_have_pnl'].sum():,.2f})")
            if avoided_pnl < 0:
                print(f"  NET SAVINGS:     ${-avoided_pnl:,.2f} (gate saved money!)")
            else:
                print(f"  NET COST:        ${avoided_pnl:,.2f} (gate blocked profitable trades)")

            print(f"\n{'Day':<12} {'Ticker':<6} {'Dir':<5} {'Score':>5} {'WouldPnL':>10} {'WouldReason':<25} {'Detail'}")
            print("-" * 110)
            for _, r in vwap_blocked.iterrows():
                print(f"{r['day']:<12} {r['ticker']:<6} {r['direction'][:4]:<5} {r['score']:>5} "
                      f"${r['would_have_pnl']:>8.2f} {r['would_have_reason']:<25} {r.get('vwap_detail', '')[:50]}")

        # Combined gate comparison
        both_blocked = df_results[
            (df_results.get("momentum_blocked", False) == True) |
            (df_results.get("vwap_blocked", False) == True)
        ] if "vwap_blocked" in df_results.columns else blocked
        pnl_both_gates = df_results[
            (df_results.get("momentum_blocked", False) != True) &
            (df_results.get("vwap_blocked", False) != True)
        ]["would_have_pnl"].sum() if "vwap_blocked" in df_results.columns else pnl_with_gate

        print(f"\n{'=' * 90}")
        print(f"COMBINED GATE COMPARISON")
        print(f"{'=' * 90}")
        print(f"P&L without any gates:     ${pnl_without_gate:,.2f}")
        print(f"P&L with momentum only:    ${pnl_with_gate:,.2f} ({len(blocked)} blocked)")
        print(f"P&L with BOTH gates:       ${total_pnl:,.2f} ({len(both_blocked)} blocked)")
        both_avoided = both_blocked["would_have_pnl"].sum() if len(both_blocked) > 0 else 0
        print(f"Total avoided P&L:         ${both_avoided:,.2f}")

    # ── Per-trade details ─────────────────────────────────────────────────

    print(f"\n{'Day':<12} {'Ticker':<6} {'Dir':<5} {'Score':>5} {'Entry':>7} {'Ct':>3} "
          f"{'Exit':>7} {'P&L':>9} {'Peak%':>6} {'Hold':>5} {'Reason':<20} {'MomBlk'}")
    print("-" * 115)
    for _, r in df_results.iterrows():
        blk = "BLOCKED" if r.get("momentum_blocked") else ""
        print(f"{r['day']:<12} {r['ticker']:<6} {r['direction'][:4]:<5} {r['score']:>5} "
              f"${r['entry']:>5.2f} {r['contracts']:>3} ${r['exit_prem']:>5.2f} "
              f"${r['pnl']:>8.2f} {r['peak_gain']:>5.0f}% {r['hold']:>4.0f}m {r['reason']:<20} {blk}")

    # ── Daily P&L ─────────────────────────────────────────────────────────

    daily = df_results.groupby("day").agg(
        trades=("pnl", "count"),
        pnl=("pnl", "sum"),
        wins=("pnl", lambda x: (x > 0).sum()),
    ).reset_index()
    daily["cum_pnl"] = daily["pnl"].cumsum()
    daily["win_rate"] = daily["wins"] / daily["trades"] * 100

    print(f"\n{'Day':<12} {'Trades':>6} {'Day P&L':>10} {'Cum P&L':>10} {'W/L':>6} {'Win%':>6}")
    print("-" * 55)
    for _, row in daily.iterrows():
        losses_d = row["trades"] - row["wins"]
        print(f"{row['day']:<12} {row['trades']:>6} ${row['pnl']:>8,.2f} "
              f"${row['cum_pnl']:>8,.2f} {int(row['wins'])}/{int(losses_d):>2} "
              f"{row['win_rate']:>5.0f}%")

    # ── Chart ─────────────────────────────────────────────────────────────

    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.dates as mdates

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 8), height_ratios=[2, 1])
        fig.suptitle("Production V5 FSM — Daily P&L (Historical Discord Signals)",
                     fontsize=14, fontweight="bold")

        dates = pd.to_datetime(daily["day"])

        # Top: cumulative P&L line
        ax1.plot(dates, daily["cum_pnl"], "b-o", markersize=5, linewidth=2)
        ax1.axhline(y=0, color="gray", linestyle="--", alpha=0.5)
        ax1.fill_between(dates, daily["cum_pnl"], 0,
                         where=daily["cum_pnl"] >= 0, alpha=0.15, color="green")
        ax1.fill_between(dates, daily["cum_pnl"], 0,
                         where=daily["cum_pnl"] < 0, alpha=0.15, color="red")
        ax1.set_ylabel("Cumulative P&L ($)")
        ax1.set_title(f"Total: ${total_pnl:,.2f} | Win Rate: {win_rate:.0f}% | "
                      f"{len(results)} trades over {len(daily)} days")
        ax1.grid(True, alpha=0.3)

        # Bottom: daily P&L bars
        colors = ["green" if p >= 0 else "red" for p in daily["pnl"]]
        ax2.bar(dates, daily["pnl"], color=colors, alpha=0.7, width=0.8)
        ax2.axhline(y=0, color="gray", linestyle="--", alpha=0.5)
        ax2.set_ylabel("Daily P&L ($)")
        ax2.set_xlabel("Date")
        ax2.grid(True, alpha=0.3)

        for ax in (ax1, ax2):
            ax.xaxis.set_major_formatter(mdates.DateFormatter("%m/%d"))
            ax.xaxis.set_major_locator(mdates.DayLocator(interval=max(1, len(daily) // 15)))
            plt.setp(ax.xaxis.get_majorticklabels(), rotation=45, ha="right")

        plt.tight_layout()
        chart_path = str(PROJECT_DIR / "v5_daily_pnl.png")
        plt.savefig(chart_path, dpi=150, bbox_inches="tight")
        print(f"\nChart saved: {chart_path}")
        plt.close()

    except ImportError:
        print("\nmatplotlib not installed — skipping chart generation")
        print("Install with: pip install matplotlib")

    # ── DOCX Report ─────────────────────────────────────────────────────
    try:
        _export_docx(df_results, daily, total_pnl, win_rate, wins, losses, no_data)
    except Exception as e:
        print(f"\nDOCX export failed: {e}")


def _export_docx(df_results, daily, total_pnl, win_rate, wins, losses, no_data):
    """Export comprehensive backtest report to DOCX."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Title
    title = doc.add_heading("OptionsOwl V5 FSM Backtest Report", 0)
    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M ET')} | "
        f"Portfolio: ${PORTFOLIO:,.0f} | "
        f"Period: {daily['day'].iloc[0]} to {daily['day'].iloc[-1]} ({len(daily)} trading days)"
    )

    # Executive summary
    doc.add_heading("Executive Summary", level=1)
    summary_data = [
        ("Total Trades", f"{len(df_results)}"),
        ("Skipped (no data)", f"{no_data}"),
        ("Total P&L", f"${total_pnl:,.2f}"),
        ("Win Rate", f"{win_rate:.1f}% ({wins}W / {losses}L)"),
        ("Avg Win", f"${df_results[df_results['pnl'] > 0]['pnl'].mean():,.2f}" if wins > 0 else "N/A"),
        ("Avg Loss", f"${df_results[df_results['pnl'] <= 0]['pnl'].mean():,.2f}" if losses > 0 else "N/A"),
        ("Max Win", f"${df_results['pnl'].max():,.2f}"),
        ("Max Loss", f"${df_results['pnl'].min():,.2f}"),
        ("Avg Hold Time", f"{df_results['hold'].mean():.0f} min"),
        ("Morning Cutoff", f"{'Enabled (11:00 AM ET)' if ENABLE_MORNING_CUTOFF else 'Disabled'}"),
        ("Scalp Target", f"{'Enabled (+25%, runner confirm +40%)' if _V6_SETTINGS.ENABLE_SCALP_TARGET else 'Disabled'}"),
    ]
    tbl = doc.add_table(rows=len(summary_data), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(summary_data):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = value

    # Exit reason breakdown
    doc.add_heading("Exit Reason Breakdown", level=1)
    reason_groups = df_results.groupby("reason")
    tbl = doc.add_table(rows=len(reason_groups) + 1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Exit Reason", "Count", "Total P&L", "Avg P&L", "Win %"]
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        for run in tbl.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    row_idx = 1
    for reason, group in reason_groups:
        gpnl = group["pnl"]
        gwins = (gpnl > 0).sum()
        gwr = gwins / len(gpnl) * 100 if len(gpnl) > 0 else 0
        tbl.rows[row_idx].cells[0].text = str(reason)
        tbl.rows[row_idx].cells[1].text = str(len(gpnl))
        tbl.rows[row_idx].cells[2].text = f"${gpnl.sum():,.2f}"
        tbl.rows[row_idx].cells[3].text = f"${gpnl.mean():,.2f}"
        tbl.rows[row_idx].cells[4].text = f"{gwr:.0f}%"
        row_idx += 1

    # Per-ticker breakdown
    doc.add_heading("Per-Ticker Breakdown", level=1)
    ticker_groups = df_results.groupby("ticker")
    tbl = doc.add_table(rows=len(ticker_groups) + 1, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Ticker", "Trades", "Total P&L", "Avg P&L", "Win %", "Top Exit Reason"]
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        for run in tbl.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    row_idx = 1
    for tkr, group in sorted(ticker_groups, key=lambda x: x[1]["pnl"].sum(), reverse=True):
        gpnl = group["pnl"]
        gwins = (gpnl > 0).sum()
        gwr = gwins / len(gpnl) * 100 if len(gpnl) > 0 else 0
        top_reason = group["reason"].value_counts().index[0] if len(group) > 0 else "N/A"
        tbl.rows[row_idx].cells[0].text = str(tkr)
        tbl.rows[row_idx].cells[1].text = str(len(gpnl))
        tbl.rows[row_idx].cells[2].text = f"${gpnl.sum():,.2f}"
        tbl.rows[row_idx].cells[3].text = f"${gpnl.mean():,.2f}"
        tbl.rows[row_idx].cells[4].text = f"{gwr:.0f}%"
        tbl.rows[row_idx].cells[5].text = str(top_reason)
        row_idx += 1

    # Daily P&L table
    doc.add_heading("Daily P&L Report", level=1)
    doc.add_paragraph(
        "Each trading day's performance with cumulative P&L. "
        "All times in ET (America/New_York)."
    )
    tbl = doc.add_table(rows=len(daily) + 1, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Date", "Trades", "Day P&L", "Cum P&L", "W/L", "Win %"]
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        for run in tbl.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for row_idx, (_, row) in enumerate(daily.iterrows(), start=1):
        losses_d = row["trades"] - row["wins"]
        tbl.rows[row_idx].cells[0].text = str(row["day"])
        tbl.rows[row_idx].cells[1].text = str(int(row["trades"]))
        tbl.rows[row_idx].cells[2].text = f"${row['pnl']:,.2f}"
        tbl.rows[row_idx].cells[3].text = f"${row['cum_pnl']:,.2f}"
        tbl.rows[row_idx].cells[4].text = f"{int(row['wins'])}/{int(losses_d)}"
        tbl.rows[row_idx].cells[5].text = f"{row['win_rate']:.0f}%"

    # Full trade log — every trade with exit logic detail
    doc.add_heading("Full Trade Log", level=1)
    doc.add_paragraph(
        "Every trade with entry/exit details and exit gate logic. "
        "All times in ET."
    )

    # Group by day for readability
    for day_str, day_group in df_results.groupby("day"):
        doc.add_heading(f"{day_str}", level=2)
        day_pnl = day_group["pnl"].sum()
        day_wins = (day_group["pnl"] > 0).sum()
        day_losses = (day_group["pnl"] <= 0).sum()
        doc.add_paragraph(
            f"Trades: {len(day_group)} | Day P&L: ${day_pnl:,.2f} | "
            f"W/L: {day_wins}/{day_losses}"
        )

        tbl = doc.add_table(rows=len(day_group) + 1, cols=9)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        headers = ["Ticker", "Dir", "Score", "Entry$", "Ct", "Exit$", "P&L", "Peak%", "Exit Reason"]
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)

        for t_idx, (_, trade) in enumerate(day_group.iterrows(), start=1):
            direction_str = trade["direction"][:4] if isinstance(trade["direction"], str) else "?"
            vals = [
                str(trade["ticker"]),
                direction_str,
                str(int(trade["score"])),
                f"${trade['entry']:.2f}",
                str(int(trade["contracts"])),
                f"${trade.get('exit_prem', 0):.2f}",
                f"${trade['pnl']:,.2f}",
                f"{trade.get('peak_gain', 0):.0f}%",
                str(trade["reason"]),
            ]
            for i, val in enumerate(vals):
                cell = tbl.rows[t_idx].cells[i]
                cell.text = val
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8)
                    # Color P&L
                    if i == 6:
                        if trade["pnl"] > 0:
                            run.font.color.rgb = RGBColor(0, 128, 0)
                        elif trade["pnl"] < 0:
                            run.font.color.rgb = RGBColor(200, 0, 0)

    # Add chart image if it exists
    chart_path = str(PROJECT_DIR / "v5_daily_pnl.png")
    if Path(chart_path).exists():
        doc.add_heading("Equity Curve", level=1)
        doc.add_picture(chart_path, width=Inches(6.5))

    # Strategy configuration appendix
    doc.add_heading("Strategy Configuration", level=1)
    config_items = [
        "V5 FSM Exit Engine (category-aware, DTE-aware)",
        f"Portfolio: ${PORTFOLIO:,.0f}",
        f"Max Concurrent: 4 trades",
        f"Max Position: 15% of portfolio",
        f"Max Portfolio Risk: 75%",
        f"Morning Cutoff: {'11:00 AM ET' if ENABLE_MORNING_CUTOFF else 'Disabled'}",
        f"Scalp Target: +{_V6_SETTINGS.SCALP_TARGET_PCT}% (runner confirm at +{_V6_SETTINGS.SCALP_RUNNER_CONFIRM_PCT}%)",
        f"Break-even Ratchet: +{_V6_SETTINGS.V6_BREAKEVEN_TRIGGER_PCT}%",
        f"Scaleout: 33% at +{_V6_SETTINGS.V6_SCALEOUT_GAIN_PCT}%",
        f"2PM Trail Tighten: 30% (factor {_V6_SETTINGS.V6_2PM_TRAIL_TIGHTEN_FACTOR})",
        f"Per-ticker Configs: Enabled",
        f"DCA: Enabled (dip 15-35%)",
        f"Score Floor: 78",
    ]
    for item in config_items:
        doc.add_paragraph(item, style="List Bullet")

    docx_path = str(PROJECT_DIR / "v5_backtest_report.docx")
    doc.save(docx_path)
    print(f"\nDOCX report saved: {docx_path}")


if __name__ == "__main__":
    main()
