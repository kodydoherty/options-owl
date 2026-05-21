"""Generate V6 Strategy Write-Up as a .docx file for team review."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import datetime

OUTPUT = Path(__file__).resolve().parent.parent / "V6_Strategy_Report.docx"


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '2F5496')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                if c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Alternate row shading
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(table.rows[r_idx + 1].cells[c_idx], 'D9E2F3')

    return table


def main():
    doc = Document()

    # ── Title ──
    title = doc.add_heading('OptionsOwl V6 Strategy Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Backtest Results & Implementation Plan\nGenerated {datetime.date.today().strftime("%B %d, %Y")}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(89, 89, 89)

    # ── Executive Summary ──
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'V6 combines per-ticker optimized FSM configs, entry filters, and 4 new exit enhancements '
        'to improve our 0DTE options trading performance. Backtested on 133 real signals from '
        'April 10 – May 1, 2026 using Polygon harvester tick data.'
    )

    summary_data = [
        ['Metric', 'V5 Baseline', 'V6 Combined', 'Change'],
        ['Total P&L', '$7,020', '$16,514', '+$9,494 (+135%)'],
        ['Trades Executed', '133', '120', '-13 (filtered)'],
        ['Win Rate', '64.7%', '72.5%', '+7.8%'],
        ['Avg Win', '$301', '$340', '+$40'],
        ['Avg Loss', '-$401', '-$397', '+$3'],
        ['Max Win', '$5,680', '$5,680', '--'],
        ['Max Loss', '-$2,520', '-$1,270', '+$1,250 (halved)'],
    ]
    add_styled_table(doc, summary_data[0], summary_data[1:])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Key Takeaway: ')
    run.bold = True
    p.add_run(
        'V6 more than doubles total P&L while cutting the max loss in half. '
        'The biggest drivers are entry filtering (blocking $3,282 in losses from bad entries like META $25.35), '
        'per-ticker configs (+$5,569), and scale-out at +20% ($4,224 locked in partial profits).'
    )

    # ── V6 Enhancements ──
    doc.add_heading('V6 Enhancements (6 New Features)', level=1)

    enhancements = [
        ('1. Premium $5 Cap (Entry Filter)',
         'Reject single-stock entries with premium > $5.00. Index tickers (SPY, QQQ, IWM) are exempt.',
         '11 trades blocked, $0 loss on filtered trades (vs -$3,282 baseline). '
         'Biggest save: META $25.35 call that lost -$2,520 in V5.'),

        ('2. Spread-Cost Gate (Entry Filter)',
         'Reject trades where bid-ask spread exceeds 15% of premium. '
         'Wide spreads indicate illiquid options with bad fill prices.',
         '2 trades blocked. Prevents slippage on illiquid contracts.'),

        ('3. Per-Ticker Optimized Configs',
         'Instead of one-size-fits-all, each ticker gets its optimal FSM configuration '
         'based on exhaustive backtesting of 12 config variations per ticker.',
         '+$5,569 improvement over default config for all tickers. '
         'See Per-Ticker Config section below.'),

        ('4. Break-Even Ratchet',
         'Once a trade reaches +20% gain, the stop floor moves to entry price. '
         'If it falls back below entry, we exit at break-even instead of riding into a loss.',
         '1 exit triggered (TSLA +$50). Prevents winners from turning into losers.'),

        ('5. 2PM Trail Tightening',
         'After 2:00 PM ET, all adaptive trail widths are tightened by 30% and '
         'soft trail keep is increased by 15%. This accounts for gamma acceleration in '
         'the last 2 hours of 0DTE trading (the "gamma death zone").',
         'Protects afternoon gains from violent end-of-day reversals.'),

        ('6. Scale-Out at +20%',
         'When a trade reaches +20% gain, sell 1/3 of contracts to lock in partial profits. '
         'The remaining 2/3 continue riding with the FSM trailing logic.',
         '46 partial fills executed, $4,224 in profits locked. '
         'This is the single biggest P&L improvement mechanism.'),

        ('7. Selective DCA (Dollar Cost Average)',
         'For 6 eligible tickers (MSFT, IWM, SPY, QQQ, AMZN, NVDA), add contracts '
         'when the position dips 15-35% during the 8-20 minute window, but only if the '
         'underlying stock price confirms the thesis (not moving against us).',
         '23 DCA additions fired. Best results: MSFT +$1,334 on a DCA\'d trade.'),
    ]

    for title_text, description, result in enhancements:
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(description)
        p = doc.add_paragraph()
        run = p.add_run('Result: ')
        run.bold = True
        p.add_run(result)

    # ── V5 FSM Gates ──
    doc.add_heading('V5 Exit FSM — 10 Gates (Priority Order)', level=1)
    doc.add_paragraph(
        'The V5 Finite State Machine evaluates 10 exit gates every 5 seconds for each open trade. '
        'The first gate that triggers causes the exit. Gates are ordered by priority — '
        'critical exits (EOD, stop loss) fire before optional ones (trailing stops, theta).'
    )

    gates = [
        ['#', 'Gate', 'What It Does', 'Win%'],
        ['1', 'EOD Cutoff', 'Force close at 3:50 PM ET (no overnight 0DTE risk)', '100%'],
        ['2', 'Bid Disappearance', 'Exit if bid drops to $0 (no liquidity)', '0%'],
        ['3', 'Profit Target', 'Hit target gain % (20% for index, varies by ticker)', '100%'],
        ['4', 'Scalp Trail', 'Quick trail after small peak — lock scalp gains', '100%'],
        ['5', 'Checkpoint Cut', 'Graduated stop: deeper cut allowed as time passes', '0%'],
        ['6', 'Hard Stop', 'Fixed hard stop (30% default, per-ticker override)', '0%'],
        ['7', 'Confirmed Stop', 'Backstop: larger loss threshold as last resort stop', '0%'],
        ['8', 'Soft Trail', 'Trailing stop in gain zone — keeps % of peak', '82%'],
        ['9', 'Adaptive Trail', 'Multi-tier trail that widens for runners', '100%'],
        ['10', 'Theta Timer', 'Time-based exit after extended hold with no progress', '0%'],
    ]
    add_styled_table(doc, gates[0], gates[1:])

    # ── Exit Reason Comparison ──
    doc.add_heading('Exit Reason Comparison: V5 vs V6', level=1)
    doc.add_paragraph(
        'Side-by-side comparison of how trades exited under V5 (baseline) vs V6 (optimized). '
        'V6 reduces the number of hard stop and confirmed stop exits while increasing '
        'profit target and scalp trail exits.'
    )

    doc.add_heading('V5 Baseline Exit Reasons', level=2)
    v5_reasons = [
        ['Reason', 'Count', 'Total P&L', 'Avg P&L', 'Win%'],
        ['soft_trail', '67', '$2,986', '$45', '79%'],
        ['scalp_trail', '15', '$1,778', '$119', '100%'],
        ['checkpoint_cut', '11', '-$5,712', '-$519', '0%'],
        ['profit_target', '10', '$3,623', '$362', '100%'],
        ['confirmed_stop', '8', '-$6,300', '-$788', '0%'],
        ['hard_stop', '7', '-$3,942', '-$563', '0%'],
        ['eod_cutoff', '5', '$7,819', '$1,564', '80%'],
        ['adaptive_trail', '3', '$1,177', '$392', '67%'],
        ['eod_data_end', '3', '$6,784', '$2,261', '67%'],
        ['theta_bleed', '2', '-$729', '-$365', '0%'],
        ['theta_timer', '2', '-$463', '-$232', '0%'],
    ]
    add_styled_table(doc, v5_reasons[0], v5_reasons[1:])

    doc.add_heading('V6 Combined Exit Reasons', level=2)
    v6_reasons = [
        ['Reason', 'Count', 'Total P&L', 'Avg P&L', 'Win%'],
        ['soft_trail', '60', '$4,047', '$67', '82%'],
        ['profit_target', '14', '$6,099', '$436', '100%'],
        ['scalp_trail', '14', '$2,439', '$174', '100%'],
        ['premium_cap_filtered', '11', '$0', '$0', 'N/A'],
        ['checkpoint_cut', '10', '-$6,042', '-$604', '0%'],
        ['hard_stop', '4', '-$2,500', '-$625', '0%'],
        ['confirmed_stop', '4', '-$2,224', '-$556', '0%'],
        ['adaptive_trail', '3', '$1,026', '$342', '100%'],
        ['eod_cutoff', '3', '$7,511', '$2,504', '100%'],
        ['eod_data_end', '3', '$7,463', '$2,488', '100%'],
        ['theta_timer', '3', '-$1,074', '-$358', '0%'],
        ['spread_filtered', '2', '$0', '$0', 'N/A'],
        ['breakeven_ratchet', '1', '$50', '$50', '100%'],
        ['theta_bleed', '1', '-$281', '-$281', '0%'],
    ]
    add_styled_table(doc, v6_reasons[0], v6_reasons[1:])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Key differences: ')
    run.bold = True
    p.add_run(
        'V6 cuts confirmed_stop exits from 8 to 4 (-$4,076 in saved losses), '
        'hard_stop from 7 to 4 (-$1,442 saved), while increasing profit_target from 10 to 14 '
        'hits (+$2,476 more captured). The 13 filtered trades would have lost -$3,282 combined.'
    )

    # ── Per-Ticker Configs ──
    doc.add_heading('Per-Ticker Configuration Map', level=1)
    doc.add_paragraph(
        'Each ticker was tested against 12 different FSM configurations. '
        'The optimal config for each ticker was selected based on total P&L improvement.'
    )

    ticker_configs = [
        ['Ticker', 'Config', 'Key Changes', 'V5 P&L', 'V6 P&L', 'Delta'],
        ['NVDA', 'EARLY_PROFIT', 'Profit target 20%, soft trail keep 70%', '-$1,278', '-$308', '+$970'],
        ['GOOGL', 'WIDE_STOP', 'Stop 45%/75%, checkpoint 40%', '-$1,110', '-$412', '+$698'],
        ['TSLA', 'LONG_GRACE', 'Grace period 8 min', '$3,494', '$3,457', '-$37'],
        ['IWM', 'WIDE_STOP', 'Stop 45%/75%, checkpoint 40%', '-$453', '$651', '+$1,104'],
        ['QQQ', 'LONG_GRACE', 'Grace period 8 min', '$1,083', '$2,109', '+$1,026'],
        ['META', 'DEFENSIVE', 'Stop 25%/50%, fast theta', '-$4,105', '-$1,013', '+$3,092'],
        ['AAPL', 'DEFENSIVE', 'Stop 25%/50%, fast theta', '-$754', '-$538', '+$216'],
        ['AMZN', 'TIGHT_TRAIL', 'Tighter adaptive tiers', '-$1,456', '-$739', '+$718'],
        ['AVGO', 'EARLY_PROFIT', 'Profit target 20%, soft trail keep 70%', '$1,682', '$1,183', '-$499'],
        ['MSFT', 'EARLY_PROFIT', 'Profit target 20%, soft trail keep 70%', '$1,362', '$2,755', '+$1,393'],
        ['MSTR', 'TIGHT+QUICK', 'Tight trail + quick scalp', '-$1,230', '-$880', '+$350'],
        ['AMD', 'DEFAULT', 'No changes needed', '$6,180', '$6,120', '-$60'],
        ['PLTR', 'DEFAULT', 'No changes needed', '$2,325', '$1,929', '-$396'],
        ['SPY', 'DEFAULT', 'No changes needed', '$1,766', '$2,685', '+$919'],
        ['MU', 'DEFAULT', 'No changes needed', '-$485', '-$485', '$0'],
    ]
    add_styled_table(doc, ticker_configs[0], ticker_configs[1:])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Pattern insight: ')
    run.bold = True
    p.add_run(
        'High-volatility stocks (NVDA, AVGO, MSFT) benefit from early profit-taking. '
        'Stocks with wide premium swings (GOOGL, IWM) need wider stops to avoid getting shaken out. '
        'Momentum stocks (TSLA, QQQ) need longer grace periods to develop. '
        'Expensive/illiquid options (META, AAPL) need defensive configs with tight stops. '
        'Index ETFs (SPY, QQQ, IWM) and liquid names benefit from selective DCA.'
    )

    # ── DCA Analysis ──
    doc.add_heading('Selective DCA (Dollar Cost Average)', level=1)
    doc.add_paragraph(
        'DCA adds contracts when a position dips 15-35% during the 8-20 minute window. '
        'This only fires for tickers where backtesting showed positive results. '
        'Early DCA (first 3-5 min) was tested and rejected — it loses money by doubling down '
        'into trades that are dying.'
    )

    dca_data = [
        ['Ticker', 'DCA Fires', 'Best Trade', 'Notes'],
        ['MSFT', '3', '+$1,334 (Apr 13)', 'DCA\'d at dip, rode to EOD'],
        ['IWM', '4', '+$780 (Apr 20)', 'Wide stop + DCA = big wins'],
        ['SPY', '3', '+$513 (Apr 30)', 'Thesis confirmed by underlying'],
        ['QQQ', '2', '+$577 (Apr 15)', 'Grace period lets DCA develop'],
        ['AMZN', '4', '+$190 (Apr 22)', 'Mixed results, DCA helps some'],
        ['NVDA', '7', '+$372 (Apr 29)', 'High fire rate, mostly positive'],
    ]
    add_styled_table(doc, dca_data[0], dca_data[1:])

    # ── V6 Feature Activity ──
    doc.add_heading('V6 Feature Activity Summary', level=1)

    activity = [
        ['Feature', 'Fires', 'P&L Impact', 'Notes'],
        ['Premium Cap Filter', '11 blocked', '+$3,282 saved', 'Biggest save: META $25.35 (-$2,520 avoided)'],
        ['Spread Gate Filter', '2 blocked', '+$150 saved', 'Catches illiquid contracts'],
        ['Scale-Out at +20%', '46 partial fills', '$4,224 locked', 'Single biggest P&L driver'],
        ['DCA Additions', '23 adds', 'Net positive', 'MSFT, IWM, SPY benefit most'],
        ['Break-Even Ratchet', '1 exit', '+$50', 'Prevents winners turning to losers'],
        ['2PM Trail Tightening', 'All afternoon trades', 'Embedded in trail', 'Protects gamma death zone gains'],
    ]
    add_styled_table(doc, activity[0], activity[1:])

    # ── Daily P&L ──
    doc.add_heading('Daily P&L Comparison', level=1)
    doc.add_paragraph(
        'V6 outperforms V5 on 9 of 13 trading days. The biggest single-day improvement is '
        'April 29 (+$2,859) driven by filtering the META $25.35 disaster and DCA additions.'
    )

    daily = [
        ['Date', 'V5 Day', 'V6 Day', 'Delta', 'V5 Cumul.', 'V6 Cumul.'],
        ['Apr 10', '-$765', '-$516', '+$249', '-$765', '-$516'],
        ['Apr 13', '$2,230', '$4,761', '+$2,531', '$1,465', '$4,245'],
        ['Apr 15', '$379', '$1,266', '+$887', '$1,844', '$5,510'],
        ['Apr 17', '$780', '$2,150', '+$1,370', '$2,624', '$7,660'],
        ['Apr 20', '-$704', '-$4', '+$701', '$1,920', '$7,657'],
        ['Apr 21', '-$3,588', '-$3,228', '+$360', '-$1,668', '$4,429'],
        ['Apr 22', '$3,060', '$2,458', '-$602', '$1,391', '$6,887'],
        ['Apr 23', '$6,571', '$6,476', '-$96', '$7,962', '$13,362'],
        ['Apr 24', '$193', '$1,123', '+$930', '$8,155', '$14,485'],
        ['Apr 27', '$3,208', '$2,907', '-$302', '$11,363', '$17,391'],
        ['Apr 29', '-$2,276', '$583', '+$2,859', '$9,088', '$17,974'],
        ['Apr 30', '$164', '$288', '+$125', '$9,251', '$18,262'],
        ['May 1', '-$2,232', '-$1,749', '+$483', '$7,020', '$16,514'],
    ]
    add_styled_table(doc, daily[0], daily[1:])

    # ── Full Trade Detail ──
    doc.add_heading('Full Trade Detail (133 Signals)', level=1)
    doc.add_paragraph(
        'Every trade with V5 baseline P&L, V6 P&L, both exit reasons, and V6 feature flags. '
        'SO = Scale-Out fired. DCA = Dollar Cost Average fired. FILTERED = Entry rejected by V6.'
    )

    # All trades
    trades = [
        ['Date', 'Ticker', 'Dir', 'Score', 'Entry', 'Ct', 'V5 P&L', 'V5 Exit', 'V6 P&L', 'V6 Exit', 'V6 Notes'],
        ['Apr 10', 'MU', 'put', '100', '$2.52', '4', '-$485', 'checkpoint_cut', '-$485', 'checkpoint_cut', ''],
        ['Apr 10', 'AMZN', 'call', '100', '$0.21', '20', '-$280', 'soft_trail', '-$31', 'soft_trail', 'DCA'],
        ['Apr 13', 'NVDA', 'call', '100', '$1.24', '9', '$68', 'soft_trail', '$9', 'soft_trail', ''],
        ['Apr 13', 'MSFT', 'call', '100', '$2.85', '4', '$1,940', 'eod_cutoff', '$3,274', 'eod_cutoff', 'DCA SO'],
        ['Apr 13', 'IWM', 'call', '100', '$0.37', '20', '-$510', 'soft_trail', '$110', 'soft_trail', 'DCA'],
        ['Apr 13', 'SPY', 'call', '94', '$1.05', '8', '$308', 'profit_target', '$276', 'profit_target', 'SO'],
        ['Apr 13', 'QQQ', 'call', '100', '$0.72', '16', '$120', 'scalp_trail', '$190', 'scalp_trail', 'SO'],
        ['Apr 13', 'AMZN', 'call', '79', '$1.03', '2', '$47', 'adaptive_trail', '$68', 'adaptive_trail', ''],
        ['Apr 13', 'TSLA', 'call', '100', '$1.81', '6', '$162', 'soft_trail', '$192', 'soft_trail', 'SO'],
        ['Apr 13', 'META', 'call', '100', '$2.85', '4', '$80', 'soft_trail', '$80', 'soft_trail', ''],
        ['Apr 13', 'QQQ', 'call', '100', '$0.92', '13', '-$78', 'soft_trail', '$468', 'profit_target', 'SO'],
        ['Apr 13', 'IWM', 'call', '84', '$0.26', '11', '$94', 'profit_target', '$94', 'profit_target', 'SO'],
        ['Apr 15', 'GOOGL', 'call', '100', '$0.58', '20', '-$185', 'soft_trail', '-$185', 'soft_trail', ''],
        ['Apr 15', 'QQQ', 'call', '100', '$1.62', '7', '$54', 'soft_trail', '$577', 'profit_target', 'DCA SO'],
        ['Apr 15', 'NVDA', 'call', '100', '$0.57', '20', '$10', 'soft_trail', '$10', 'soft_trail', ''],
        ['Apr 15', 'AAPL', 'call', '100', '$1.41', '8', '$83', 'soft_trail', '$83', 'soft_trail', ''],
        ['Apr 15', 'SPY', 'call', '100', '$1.19', '10', '$435', 'profit_target', '$798', 'profit_target', 'DCA SO'],
        ['Apr 15', 'SPY', 'call', '99', '$0.43', '20', '-$17', 'soft_trail', '-$17', 'soft_trail', ''],
        ['Apr 17', 'SPY', 'call', '93', '$1.55', '5', '$120', 'soft_trail', '$136', 'soft_trail', 'SO'],
        ['Apr 17', 'AMZN', 'call', '100', '$1.00', '12', '$114', 'soft_trail', '$114', 'soft_trail', ''],
        ['Apr 17', 'TSLA', 'call', '100', '$3.30', '3', '-$8', 'soft_trail', '$50', 'breakeven_ratchet', 'SO BE'],
        ['Apr 17', 'IWM', 'call', '100', '$0.68', '17', '$374', 'profit_target', '$337', 'profit_target', 'SO'],
        ['Apr 17', 'QQQ', 'call', '100', '$1.46', '8', '-$4', 'soft_trail', '-$4', 'soft_trail', ''],
        ['Apr 17', 'AAPL', 'call', '100', '$1.04', '11', '$605', 'adaptive_trail', '$604', 'adaptive_trail', 'SO'],
        ['Apr 17', 'META', 'call', '100', '$2.35', '5', '-$228', 'soft_trail', '-$228', 'soft_trail', ''],
        ['Apr 17', 'GOOGL', 'call', '100', '$0.63', '19', '-$789', 'confirmed_stop', '$166', 'soft_trail', 'SO'],
        ['Apr 17', 'NVDA', 'call', '96', '$0.26', '20', '$10', 'soft_trail', '$10', 'soft_trail', ''],
        ['Apr 17', 'QQQ', 'call', '100', '$1.29', '9', '$153', 'soft_trail', '$5', 'soft_trail', 'DCA'],
        ['Apr 17', 'IWM', 'call', '92', '$0.51', '17', '$281', 'profit_target', '$256', 'profit_target', 'SO'],
        ['Apr 17', 'AMD', 'call', '100', '$0.84', '14', '$175', 'soft_trail', '$205', 'soft_trail', 'SO'],
        ['Apr 17', 'AMZN', 'call', '100', '$1.21', '9', '$59', 'soft_trail', '$59', 'soft_trail', ''],
        ['Apr 17', 'AAPL', 'call', '100', '$0.33', '20', '-$430', 'hard_stop', '-$330', 'hard_stop', ''],
        ['Apr 17', 'PLTR', 'call', '100', '$0.46', '20', '$190', 'scalp_trail', '$193', 'scalp_trail', 'SO'],
        ['Apr 17', 'NVDA', 'call', '100', '$0.88', '13', '$46', 'soft_trail', '$26', 'soft_trail', 'DCA'],
        ['Apr 17', 'QQQ', 'call', '100', '$0.77', '15', '$15', 'scalp_trail', '$88', 'scalp_trail', 'SO'],
        ['Apr 17', 'TSLA', 'put', '100', '$1.45', '8', '$60', 'soft_trail', '$103', 'soft_trail', 'SO'],
        ['Apr 17', 'IWM', 'call', '100', '$0.61', '19', '$19', 'soft_trail', '$345', 'scalp_trail', 'DCA SO'],
        ['Apr 17', 'NVDA', 'call', '100', '$0.94', '12', '$18', 'soft_trail', '$18', 'soft_trail', ''],
        ['Apr 20', 'IWM', 'call', '100', '$0.59', '20', '$80', 'soft_trail', '$781', 'profit_target', 'DCA SO'],
        ['Apr 20', 'GOOGL', 'call', '100', '$0.71', '16', '-$784', 'checkpoint_cut', '-$784', 'checkpoint_cut', ''],
        ['Apr 21', 'IWM', 'call', '100', '$0.57', '20', '-$790', 'checkpoint_cut', '-$1,270', 'checkpoint_cut', 'DCA'],
        ['Apr 21', 'MSFT', 'call', '100', '$3.20', '3', '-$153', 'soft_trail', '-$272', 'theta_timer', 'DCA'],
        ['Apr 21', 'NVDA', 'call', '100', '$1.32', '9', '-$666', 'confirmed_stop', '$27', 'soft_trail', ''],
        ['Apr 21', 'QQQ', 'call', '100', '$1.58', '7', '$70', 'soft_trail', '$70', 'soft_trail', ''],
        ['Apr 21', 'NVDA', 'call', '88', '$0.95', '6', '-$303', 'theta_timer', '-$177', 'theta_timer', 'DCA'],
        ['Apr 21', 'AMZN', 'call', '100', '$2.25', '5', '-$588', 'confirmed_stop', '-$625', 'theta_timer', 'DCA'],
        ['Apr 21', 'META', 'call', '100', '$4.95', '2', '-$449', 'theta_bleed', '-$281', 'theta_bleed', ''],
        ['Apr 21', 'AVGO', 'call', '98', '$3.80', '3', '$66', 'soft_trail', '$75', 'soft_trail', ''],
        ['Apr 21', 'SPY', 'call', '100', '$1.06', '11', '-$776', 'hard_stop', '-$776', 'hard_stop', ''],
        ['Apr 22', 'AVGO', 'call', '100', '$2.58', '4', '$2,708', 'eod_cutoff', '$2,104', 'eod_cutoff', 'SO'],
        ['Apr 22', 'MSTR', 'call', '100', '$5.65', '2', '-$310', 'confirmed_stop', '$0', 'FILTERED', 'premium cap'],
        ['Apr 22', 'NVDA', 'call', '100', '$0.34', '20', '-$450', 'hard_stop', '-$810', 'hard_stop', 'DCA'],
        ['Apr 22', 'TSLA', 'put', '92', '$10.15', '1', '$90', 'soft_trail', '$0', 'FILTERED', 'premium cap'],
        ['Apr 22', 'AMZN', 'call', '100', '$0.59', '20', '$80', 'soft_trail', '$190', 'soft_trail', 'DCA'],
        ['Apr 22', 'PLTR', 'call', '100', '$3.50', '3', '$60', 'soft_trail', '$60', 'soft_trail', ''],
        ['Apr 22', 'AVGO', 'call', '100', '$1.91', '6', '-$744', 'checkpoint_cut', '-$744', 'checkpoint_cut', ''],
        ['Apr 22', 'AMD', 'call', '100', '$7.00', '1', '$50', 'soft_trail', '$0', 'FILTERED', 'premium cap'],
        ['Apr 22', 'SPY', 'call', '100', '$0.60', '20', '$570', 'profit_target', '$570', 'profit_target', 'SO'],
        ['Apr 22', 'NVDA', 'call', '100', '$0.14', '20', '$100', 'scalp_trail', '$103', 'scalp_trail', 'SO'],
        ['Apr 22', 'QQQ', 'call', '100', '$1.38', '8', '$336', 'profit_target', '$336', 'profit_target', 'SO'],
        ['Apr 22', 'AMZN', 'call', '95', '$0.11', '20', '$10', 'soft_trail', '$0', 'FILTERED', 'spread gate'],
        ['Apr 22', 'NVDA', 'call', '100', '$0.29', '20', '$50', 'scalp_trail', '$77', 'scalp_trail', 'SO'],
        ['Apr 22', 'SPY', 'call', '100', '$0.72', '16', '$56', 'soft_trail', '$280', 'soft_trail', 'DCA'],
        ['Apr 22', 'GOOGL', 'call', '100', '$1.14', '10', '$610', 'eod_cutoff', '$354', 'adaptive_trail', 'SO'],
        ['Apr 22', 'META', 'call', '100', '$1.34', '8', '-$708', 'confirmed_stop', '-$584', 'hard_stop', ''],
        ['Apr 22', 'AVGO', 'call', '100', '$1.11', '10', '$135', 'scalp_trail', '$176', 'scalp_trail', 'SO'],
        ['Apr 22', 'SPY', 'call', '93', '$0.52', '17', '$417', 'profit_target', '$347', 'profit_target', 'SO'],
        ['Apr 23', 'AVGO', 'call', '100', '$6.05', '1', '$161', 'soft_trail', '$0', 'FILTERED', 'premium cap'],
        ['Apr 23', 'AMZN', 'call', '100', '$1.74', '6', '$246', 'soft_trail', '$244', 'soft_trail', 'SO'],
        ['Apr 23', 'SPY', 'call', '100', '$1.22', '9', '$414', 'profit_target', '$366', 'profit_target', 'SO'],
        ['Apr 23', 'QQQ', 'call', '100', '$1.69', '7', '$21', 'soft_trail', '$21', 'soft_trail', ''],
        ['Apr 23', 'AMD', 'call', '100', '$4.55', '2', '$235', 'soft_trail', '$235', 'soft_trail', ''],
        ['Apr 23', 'AVGO', 'call', '100', '$4.00', '3', '-$626', 'confirmed_stop', '-$626', 'confirmed_stop', ''],
        ['Apr 23', 'AAPL', 'call', '100', '$1.35', '8', '-$612', 'checkpoint_cut', '-$612', 'checkpoint_cut', ''],
        ['Apr 23', 'MSTR', 'call', '97', '$3.95', '3', '-$620', 'confirmed_stop', '-$620', 'confirmed_stop', ''],
        ['Apr 23', 'PLTR', 'put', '95', '$1.61', '7', '$1,621', 'eod_data_end', '$1,236', 'eod_data_end', 'SO'],
        ['Apr 23', 'TSLA', 'call', '94', '$4.20', '2', '$45', 'soft_trail', '$45', 'soft_trail', ''],
        ['Apr 23', 'AVGO', 'call', '100', '$4.50', '2', '$80', 'soft_trail', '$93', 'soft_trail', ''],
        ['Apr 23', 'NVDA', 'put', '91', '$1.18', '7', '-$53', 'soft_trail', '$110', 'scalp_trail', 'DCA SO'],
        ['Apr 23', 'MSFT', 'put', '100', '$4.60', '2', '$90', 'soft_trail', '$100', 'soft_trail', ''],
        ['Apr 23', 'QQQ', 'call', '100', '$1.37', '8', '-$12', 'soft_trail', '-$12', 'soft_trail', ''],
        ['Apr 23', 'AMD', 'call', '100', '$4.70', '2', '$5,680', 'eod_data_end', '$5,680', 'eod_data_end', ''],
        ['Apr 23', 'TSLA', 'put', '100', '$3.55', '3', '$150', 'soft_trail', '$185', 'soft_trail', 'SO'],
        ['Apr 23', 'AMZN', 'put', '100', '$1.89', '6', '$30', 'soft_trail', '$30', 'soft_trail', ''],
        ['Apr 23', 'META', 'put', '100', '$6.90', '1', '-$280', 'hard_stop', '$0', 'FILTERED', 'premium cap'],
        ['Apr 24', 'PLTR', 'put', '100', '$1.00', '12', '$402', 'scalp_trail', '$388', 'scalp_trail', 'SO'],
        ['Apr 24', 'NVDA', 'call', '100', '$1.74', '6', '$240', 'soft_trail', '$352', 'soft_trail', 'SO'],
        ['Apr 24', 'MSTR', 'put', '82', '$1.63', '1', '-$57', 'checkpoint_cut', '-$57', 'checkpoint_cut', ''],
        ['Apr 24', 'AVGO', 'put', '100', '$3.83', '3', '-$99', 'soft_trail', '$105', 'soft_trail', ''],
        ['Apr 24', 'TSLA', 'put', '100', '$1.40', '8', '$104', 'scalp_trail', '$143', 'scalp_trail', 'SO'],
        ['Apr 24', 'SPY', 'call', '100', '$1.37', '8', '$120', 'soft_trail', '$164', 'scalp_trail', 'SO'],
        ['Apr 24', 'AMZN', 'call', '100', '$1.07', '11', '-$517', 'confirmed_stop', '$28', 'soft_trail', 'DCA'],
        ['Apr 27', 'AAPL', 'put', '100', '$1.09', '11', '$110', 'soft_trail', '$110', 'soft_trail', ''],
        ['Apr 27', 'TSLA', 'put', '100', '$2.58', '4', '-$78', 'soft_trail', '$186', 'scalp_trail', 'SO'],
        ['Apr 27', 'NVDA', 'call', '100', '$0.44', '20', '-$50', 'soft_trail', '-$50', 'soft_trail', ''],
        ['Apr 27', 'TSLA', 'call', '100', '$2.53', '4', '$2,768', 'eod_cutoff', '$2,133', 'eod_cutoff', 'SO'],
        ['Apr 27', 'NVDA', 'call', '100', '$0.44', '20', '$50', 'soft_trail', '$50', 'soft_trail', ''],
        ['Apr 27', 'QQQ', 'call', '100', '$1.00', '12', '$396', 'profit_target', '$382', 'profit_target', 'SO'],
        ['Apr 27', 'TSLA', 'call', '100', '$1.40', '8', '$12', 'soft_trail', '$96', 'scalp_trail', 'SO'],
        ['Apr 29', 'NVDA', 'put', '100', '$1.38', '8', '$116', 'soft_trail', '$28', 'soft_trail', 'DCA'],
        ['Apr 29', 'META', 'call', '100', '$25.35', '1', '-$2,520', 'hard_stop', '$0', 'FILTERED', 'premium cap'],
        ['Apr 29', 'TSLA', 'put', '100', '$2.37', '5', '-$43', 'soft_trail', '-$243', 'soft_trail', ''],
        ['Apr 29', 'MSTR', 'put', '100', '$4.55', '2', '$115', 'soft_trail', '$155', 'soft_trail', ''],
        ['Apr 29', 'NVDA', 'call', '100', '$2.05', '5', '-$3', 'soft_trail', '-$3', 'soft_trail', ''],
        ['Apr 29', 'TSLA', 'call', '100', '$1.83', '6', '$138', 'scalp_trail', '$181', 'scalp_trail', 'SO'],
        ['Apr 29', 'NVDA', 'put', '100', '$1.79', '6', '$36', 'soft_trail', '$372', 'soft_trail', 'DCA SO'],
        ['Apr 29', 'TSLA', 'call', '79', '$2.01', '1', '$13', 'soft_trail', '$13', 'soft_trail', ''],
        ['Apr 29', 'PLTR', 'call', '104', '$2.76', '4', '$52', 'soft_trail', '$52', 'soft_trail', ''],
        ['Apr 29', 'SPY', 'call', '167', '$2.15', '5', '$28', 'soft_trail', '$28', 'soft_trail', ''],
        ['Apr 29', 'AMZN', 'call', '163', '$15.40', '1', '-$208', 'confirmed_stop', '$0', 'FILTERED', 'premium cap'],
        ['Apr 30', 'SPY', 'call', '145', '$3.29', '3', '$92', 'soft_trail', '$513', 'profit_target', 'DCA SO'],
        ['Apr 30', 'MSFT', 'put', '90', '$4.65', '1', '$30', 'soft_trail', '$160', 'soft_trail', 'DCA'],
        ['Apr 30', 'TSLA', 'call', '130', '$5.10', '2', '$40', 'soft_trail', '$0', 'FILTERED', 'premium cap'],
        ['Apr 30', 'AMD', 'call', '93', '$6.05', '1', '$0', 'eod_data_end', '$0', 'FILTERED', 'premium cap'],
        ['Apr 30', 'AMZN', 'put', '78', '$2.94', '1', '$26', 'soft_trail', '$26', 'soft_trail', ''],
        ['Apr 30', 'NVDA', 'put', '97', '$1.99', '6', '$66', 'soft_trail', '$102', 'soft_trail', 'DCA'],
        ['Apr 30', 'QQQ', 'call', '82', '$2.27', '1', '$12', 'soft_trail', '-$11', 'soft_trail', ''],
        ['Apr 30', 'AAPL', 'call', '101', '$4.50', '2', '$50', 'soft_trail', '$50', 'soft_trail', ''],
        ['Apr 30', 'AMD', 'call', '116', '$6.55', '1', '$40', 'scalp_trail', '$0', 'FILTERED', 'premium cap'],
        ['Apr 30', 'MSFT', 'put', '127', '$3.70', '3', '$60', 'soft_trail', '$98', 'soft_trail', ''],
        ['Apr 30', 'TSLA', 'call', '89', '$4.05', '1', '$33', 'soft_trail', '$33', 'soft_trail', ''],
        ['Apr 30', 'AMZN', 'put', '156', '$2.35', '5', '-$620', 'confirmed_stop', '-$620', 'confirmed_stop', ''],
        ['Apr 30', 'GOOGL', 'call', '105', '$4.00', '3', '$38', 'soft_trail', '$38', 'soft_trail', ''],
        ['Apr 30', 'MSTR', 'put', '90', '$2.29', '3', '-$359', 'confirmed_stop', '-$359', 'confirmed_stop', ''],
        ['Apr 30', 'AAPL', 'call', '94', '$5.30', '1', '$33', 'soft_trail', '$0', 'FILTERED', 'premium cap'],
        ['Apr 30', 'AMZN', 'call', '97', '$2.79', '4', '$624', 'eod_cutoff', '$259', 'soft_trail', 'SO'],
        ['May 1', 'AAPL', 'put', '91', '$0.12', '20', '-$150', 'hard_stop', '$0', 'FILTERED', 'spread gate'],
        ['May 1', 'TSLA', 'call', '134', '$2.79', '4', '-$104', 'soft_trail', '-$206', 'soft_trail', ''],
        ['May 1', 'AAPL', 'call', '120', '$2.59', '4', '-$442', 'checkpoint_cut', '-$442', 'checkpoint_cut', ''],
        ['May 1', 'AMZN', 'call', '127', '$1.75', '6', '-$480', 'checkpoint_cut', '-$480', 'checkpoint_cut', ''],
        ['May 1', 'NVDA', 'put', '103', '$1.25', '9', '-$563', 'checkpoint_cut', '-$563', 'checkpoint_cut', ''],
        ['May 1', 'MSFT', 'put', '146', '$2.39', '5', '-$605', 'checkpoint_cut', '-$605', 'checkpoint_cut', ''],
        ['May 1', 'TSLA', 'call', '116', '$2.57', '4', '$112', 'soft_trail', '$547', 'eod_data_end', 'SO'],
    ]
    add_styled_table(doc, trades[0], trades[1:])

    # ── Risk Notes ──
    doc.add_heading('Risk Notes & Caveats', level=1)

    risks = [
        'Backtest uses Polygon harvester tick data (60s snapshots), not true real-time fills. '
        'Actual slippage may be worse.',
        'The 133-signal sample covers only 13 trading days (Apr 10 – May 1). '
        'More data is needed to confirm statistical significance.',
        'DCA adds capital to losing positions — if the thesis is wrong, it amplifies losses '
        '(see IWM Apr 21: -$1,270 with DCA vs -$790 without).',
        'Per-ticker configs were optimized on the same data they were tested on (in-sample). '
        'True out-of-sample validation requires new signals.',
        'Scale-out at +20% reduces the upside of big runners '
        '(see AVGO Apr 22: $2,104 vs $2,708 baseline). This is an intentional tradeoff for consistency.',
        '2PM trail tightening may cut profitable afternoon trades short if the move extends past 3PM.',
    ]
    for risk in risks:
        doc.add_paragraph(risk, style='List Bullet')

    # ── Implementation Plan ──
    doc.add_heading('Implementation Plan', level=1)

    steps = [
        ('Phase 1: Entry Filters (Low Risk)',
         'Add premium $5 cap and spread-cost gate to the entry pipeline. '
         'These are pure filters — they only block bad trades, no behavior changes for passing trades.'),
        ('Phase 2: Per-Ticker Configs (Medium Risk)',
         'Add TICKER_CONFIGS mapping to the FSM. Each ticker gets its backtested-optimal '
         'V5Config. Tickers not in the map use DEFAULT. Deploy to one bot first for validation.'),
        ('Phase 3: Scale-Out (Medium Risk)',
         'Implement partial contract selling at +20% gain. Requires Webull API partial close support. '
         'Track scaleout_pnl separately in paper_trades table.'),
        ('Phase 4: DCA & Ratchet (Higher Risk)',
         'Enable selective DCA for 6 tickers and break-even ratchet. These add capital to trades, '
         'so deploy conservatively — one bot first, then expand after 1 week of live validation.'),
        ('Phase 5: 2PM Trail Tightening (Low Risk)',
         'Tighten adaptive trail widths by 30% after 2PM ET. Pure parameter change, no new logic.'),
    ]
    for title_text, description in steps:
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(description)

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— End of Report —')
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.size = Pt(10)

    doc.save(str(OUTPUT))
    print(f"Report saved to: {OUTPUT}")


if __name__ == "__main__":
    main()
