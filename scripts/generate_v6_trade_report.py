"""Generate V6 Backtest Trade-by-Trade Report as .docx

Runs the production V6 FSM + DCA on all historical signals and generates a
Word document with every trade, organized by day, with P&L, exit reasons,
and V6 feature annotations.

Usage:
    python scripts/generate_v6_trade_report.py
    # => reports/V6_Trade_Report_YYYY-MM-DD.docx
"""

from __future__ import annotations

import sqlite3
import sys
from datetime import datetime, timedelta
from pathlib import Path
from types import SimpleNamespace

import numpy as np
import pandas as pd

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx required. pip install python-docx")
    sys.exit(1)

PROJECT_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_DIR))

from options_owl.risk.exit_v5.config import (
    INDEX_TICKERS,
    TICKER_CONFIGS,
    get_ticker_config,
)
from options_owl.risk.exit_v5.fsm import ExitFSM, TradeState

SIGNALS_DB = str(PROJECT_DIR / "journal" / "owlet-kody" / "raw_messages.db")
HARVESTER_DB = str(PROJECT_DIR / "journal" / "owlet-harvester" / "options_data.db")
REPORTS_DIR = PROJECT_DIR / "reports"
PORTFOLIO = 8000


# ── Styling helpers ──────────────────────────────────────────────────────


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    elm = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex,
    })
    shading.append(elm)


def add_table(doc, headers, rows, pnl_col=None):
    """Add a styled table. If pnl_col is set, color that column green/red."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(7)
        set_cell_shading(cell, "2F5496")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7)
                if c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Color P&L columns
            if pnl_col is not None and c_idx in (pnl_col if isinstance(pnl_col, (list, tuple)) else [pnl_col]):
                try:
                    v = float(str(val).replace("$", "").replace(",", "").replace("+", ""))
                    if v > 0:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(0, 128, 0)
                    elif v < 0:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(200, 0, 0)
                except ValueError:
                    pass

        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(table.rows[r_idx + 1].cells[c_idx], "D9E2F3")

    return table


# ── Settings ─────────────────────────────────────────────────────────────


def _v6_settings():
    return SimpleNamespace(
        ENABLE_V6_PER_TICKER_CONFIG=True,
        ENABLE_V6_BREAKEVEN_RATCHET=True,
        V6_BREAKEVEN_TRIGGER_PCT=20.0,
        ENABLE_V6_2PM_TIGHTEN=True,
        V6_2PM_TRAIL_TIGHTEN_FACTOR=0.7,
        V6_2PM_SOFT_TRAIL_BOOST=0.15,
        ENABLE_V6_PREMIUM_CAP=True,
        V6_PREMIUM_CAP=5.0,
        ENABLE_V6_SPREAD_GATE=True,
        V6_MAX_SPREAD_PCT=15.0,
        ENABLE_V6_SCALEOUT=True,
        V6_SCALEOUT_GAIN_PCT=20.0,
        V6_SCALEOUT_FRACTION=0.333,
        V6_SCALEOUT_MIN_CONTRACTS=3,
        ENABLE_V6_DCA=True,
        V6_DCA_TICKERS="MSFT,IWM,SPY,QQQ,AMZN,NVDA",
        V6_DCA_MIN_MINUTES=8.0,
        V6_DCA_MAX_MINUTES=20.0,
        V6_DCA_MIN_DIP_PCT=15.0,
        V6_DCA_MAX_DIP_PCT=35.0,
        V6_DCA_UNDERLYING_THRESHOLD=0.5,
    )


def _v5_settings():
    return SimpleNamespace(
        ENABLE_V6_PER_TICKER_CONFIG=False,
        ENABLE_V6_BREAKEVEN_RATCHET=False,
        ENABLE_V6_2PM_TIGHTEN=False,
        ENABLE_V6_SCALEOUT=False,
        ENABLE_V6_DCA=False,
    )


# ── Data loading (same as backtest_v6_full_report.py) ────────────────────


def load_signals():
    conn = sqlite3.connect(SIGNALS_DB)
    conn.row_factory = sqlite3.Row
    rows = conn.execute("""
        SELECT id, ticker, direction, sentiment, score,
               atm_premium, otm_premium, strike, expiry,
               entry_price, created_at
        FROM trade_signals
        WHERE score >= 78
        ORDER BY created_at
    """).fetchall()
    signals = []
    for r in rows:
        sig = dict(r)
        sig["premium"] = sig["atm_premium"] or sig["otm_premium"]
        sent = (sig.get("sentiment") or sig.get("direction") or "bullish").lower()
        sig["option_type"] = "put" if sent in ("bearish", "put") else "call"
        if sig["premium"] and sig["premium"] > 0 and sig["strike"]:
            signals.append(sig)
    conn.close()
    return signals


def build_contract_ticker(ticker, expiry, strike, option_type):
    if not expiry:
        return ""
    try:
        exp_dt = datetime.strptime(expiry, "%Y-%m-%d")
    except ValueError:
        return ""
    exp_str = exp_dt.strftime("%y%m%d")
    ot = "C" if option_type.lower() in ("call", "bullish", "c") else "P"
    strike_int = int(strike * 1000)
    return f"O:{ticker}{exp_str}{ot}{strike_int:08d}"


def load_ticks(harvester_conn, signal):
    ticker = signal["ticker"]
    strike = signal["strike"]
    created_at = signal["created_at"]
    option_type = signal["option_type"]
    sig_date = created_at[:10]
    sig_dt = datetime.strptime(sig_date, "%Y-%m-%d").date()
    candidates = [sig_dt]
    for delta in range(1, 6):
        d = sig_dt + timedelta(days=delta)
        if d.weekday() < 5:
            candidates.append(d)
            if len(candidates) >= 4:
                break
    for exp_date in candidates:
        expiry = exp_date.strftime("%Y-%m-%d")
        ct = build_contract_ticker(ticker, expiry, strike, option_type)
        if not ct:
            continue
        rows = harvester_conn.execute("""
            SELECT captured_at, midpoint, bid, ask, underlying_price,
                   implied_volatility, delta, gamma, theta, vega, day_volume
            FROM harvest_snapshots
            WHERE contract_ticker = ? AND captured_at >= ?
            ORDER BY captured_at
        """, (ct, created_at)).fetchall()
        if rows and len(rows) >= 10:
            signal["_dte"] = (exp_date - sig_dt).days
            signal["_expiry_date"] = expiry
            break
    else:
        return None
    df = pd.DataFrame(rows, columns=[
        "captured_at", "midpoint", "bid", "ask", "underlying_price",
        "iv", "delta", "gamma", "theta", "vega", "volume"
    ])
    df["premium"] = df["midpoint"].where(df["midpoint"] > 0, (df["bid"] + df["ask"]) / 2)
    df["premium"] = df["premium"].where(df["premium"] > 0, np.nan)
    df = df.dropna(subset=["premium"])
    if len(df) < 10:
        return None
    df["ts"] = pd.to_datetime(df["captured_at"])
    df = df.sort_values("ts").reset_index(drop=True)
    return df


def compute_contracts(entry_premium, score):
    deployable = PORTFOLIO * 0.75
    per_slot = deployable / 5
    position_cap = PORTFOLIO * 0.15
    if score >= 95: sm = 1.0
    elif score >= 90: sm = 0.75
    elif score >= 85: sm = 0.50
    else: sm = 0.25
    cost_per = entry_premium * 100
    scaled = per_slot * sm
    raw = int(scaled / cost_per) if cost_per > 0 else 1
    pos_cap = int(position_cap / cost_per) if cost_per > 0 else 1
    return max(1, min(raw, pos_cap, 20))


# ── Simulation ───────────────────────────────────────────────────────────


def simulate(df, entry_premium, contracts, direction, dte, expiry_date,
             ticker, settings):
    if entry_premium <= 0:
        return {
            "pnl": 0, "reason": "no_data", "hold_min": 0, "peak_gain": 0,
            "exit_prem": 0, "entry_prem": entry_premium,
            "scaleout_pnl": 0, "dca_fired": False, "dca_add": 0,
            "final_contracts": contracts, "entry_time": "", "exit_time": "",
        }

    option_type = "put" if direction in ("bearish", "put") else "call"
    is_call = option_type in ("call", "bullish")

    use_per_ticker = getattr(settings, "ENABLE_V6_PER_TICKER_CONFIG", False)
    cfg = get_ticker_config(ticker, use_per_ticker=use_per_ticker)
    fsm = ExitFSM(cfg, settings=settings)

    entry_ts = df["ts"].iloc[0]
    if hasattr(entry_ts, "to_pydatetime"):
        entry_ts = entry_ts.to_pydatetime()
    if entry_ts.tzinfo is not None:
        entry_ts = entry_ts.replace(tzinfo=None)

    first_underlying = 0.0
    for i in range(min(5, len(df))):
        u = df["underlying_price"].iloc[i]
        if u and u > 0:
            first_underlying = float(u)
            break

    state = TradeState(
        trade_id=1, ticker=ticker, option_type=option_type,
        entry_premium=entry_premium, entry_time=entry_ts,
        contracts=contracts, peak_premium=entry_premium,
        entry_underlying_price=first_underlying,
        dte=dte, expiry_date=expiry_date or "",
    )

    remaining_contracts = contracts
    scaleout_pnl = 0.0
    avg_entry = entry_premium
    total_cost = entry_premium * contracts * 100

    enable_dca = getattr(settings, "ENABLE_V6_DCA", False)
    dca_tickers_str = getattr(settings, "V6_DCA_TICKERS", "")
    dca_tickers = {t.strip().upper() for t in dca_tickers_str.split(",") if t.strip()}
    dca_fired = False
    dca_add = 0

    def _result(pnl, reason, exit_prem, now):
        elapsed = (now - entry_ts).total_seconds() / 60
        peak_gain = (state.peak_premium - avg_entry) / avg_entry * 100 if avg_entry > 0 else 0
        return {
            "pnl": round(pnl, 2), "reason": reason,
            "hold_min": round(elapsed, 1),
            "peak_gain": round(peak_gain, 1),
            "exit_prem": round(exit_prem, 4),
            "entry_prem": round(avg_entry, 4),
            "scaleout_pnl": round(scaleout_pnl, 2),
            "dca_fired": dca_fired, "dca_add": dca_add,
            "final_contracts": remaining_contracts,
            "entry_time": entry_ts.strftime("%H:%M") if entry_ts else "",
            "exit_time": now.strftime("%H:%M") if now else "",
        }

    for idx in range(1, len(df)):
        premium = df["premium"].iloc[idx]
        if np.isnan(premium) or premium <= 0:
            continue

        raw_bid = df["bid"].iloc[idx]
        raw_ask = df["ask"].iloc[idx]
        bid = float(raw_bid) if raw_bid and not pd.isna(raw_bid) else premium
        ask = float(raw_ask) if raw_ask and not pd.isna(raw_ask) else premium

        now = df["ts"].iloc[idx]
        if hasattr(now, "to_pydatetime"):
            now = now.to_pydatetime()
        if now.tzinfo is not None:
            now = now.replace(tzinfo=None)

        underlying = df["underlying_price"].iloc[idx] or 0.0
        elapsed_min = (now - entry_ts).total_seconds() / 60.0
        et_hour = now.hour - 4
        if et_hour < 0:
            et_hour += 24
        et_minute = now.minute
        minutes_to_close = max(0, (16 * 60) - (et_hour * 60 + et_minute))

        # V6 DCA
        if enable_dca and not dca_fired and ticker in dca_tickers:
            dca_min = getattr(settings, "V6_DCA_MIN_MINUTES", 8.0)
            dca_max = getattr(settings, "V6_DCA_MAX_MINUTES", 20.0)
            dca_min_dip = getattr(settings, "V6_DCA_MIN_DIP_PCT", 15.0)
            dca_max_dip = getattr(settings, "V6_DCA_MAX_DIP_PCT", 35.0)
            u_thresh = getattr(settings, "V6_DCA_UNDERLYING_THRESHOLD", 0.5)

            if dca_min <= elapsed_min <= dca_max and avg_entry > 0:
                dip_pct = (avg_entry - premium) / avg_entry * 100
                if dca_min_dip <= dip_pct <= dca_max_dip:
                    underlying_ok = True
                    if first_underlying > 0 and underlying > 0:
                        u_move = (underlying - first_underlying) / first_underlying * 100
                        if is_call and u_move < -u_thresh:
                            underlying_ok = False
                        elif not is_call and u_move > u_thresh:
                            underlying_ok = False
                    if underlying_ok:
                        dca_fired = True
                        dca_add = contracts
                        total_cost += premium * dca_add * 100
                        remaining_contracts += dca_add
                        avg_entry = total_cost / (remaining_contracts * 100)
                        state.entry_premium = avg_entry
                        state.contracts = remaining_contracts
                        state.peak_premium = max(avg_entry, premium)

        action = fsm.evaluate(
            state, premium, bid, ask, now,
            current_underlying=underlying,
            minutes_to_close=minutes_to_close,
        )

        if action.should_exit:
            if action.contracts_to_close > 0:
                close_qty = min(action.contracts_to_close, remaining_contracts)
                scaleout_pnl += (premium - avg_entry) * close_qty * 100
                remaining_contracts -= close_qty
                state.contracts = remaining_contracts
                if remaining_contracts <= 0:
                    return _result(scaleout_pnl, action.reason.value, premium, now)
                continue

            pnl = (premium - avg_entry) * remaining_contracts * 100 + scaleout_pnl
            return _result(pnl, action.reason.value, premium, now)

    last_prem = df["premium"].iloc[-1]
    last_ts = df["ts"].iloc[-1]
    if hasattr(last_ts, "to_pydatetime"):
        last_ts = last_ts.to_pydatetime()
    if last_ts.tzinfo is not None:
        last_ts = last_ts.replace(tzinfo=None)
    pnl = (last_prem - avg_entry) * remaining_contracts * 100 + scaleout_pnl
    return _result(pnl, "eod_data_end", last_prem, last_ts)


# ── DOCX generation ─────────────────────────────────────────────────────


def fmt_pnl(v):
    return f"${v:+,.2f}" if v != 0 else "$0.00"


def main():
    print("Loading signals...")
    signals = load_signals()
    harvester_conn = sqlite3.connect(HARVESTER_DB)
    v6_settings = _v6_settings()
    v5_settings = _v5_settings()

    results = []
    no_data = []

    for sig in signals:
        ticker = sig["ticker"]
        direction = (sig["direction"] or "bullish").lower()
        score = sig["score"] or 80
        day = sig["created_at"][:10]
        sig_time = sig["created_at"][11:16] if len(sig["created_at"]) > 11 else ""

        df = load_ticks(harvester_conn, sig)
        if df is None:
            no_data.append({
                "day": day, "time": sig_time, "ticker": ticker,
                "dir": direction[:4].upper(), "score": score,
                "strike": sig["strike"], "premium": sig["premium"],
            })
            continue

        dte = sig.get("_dte", 0)
        expiry_date = sig.get("_expiry_date", "")

        first_ask = df["ask"].iloc[0]
        first_mid = df["premium"].iloc[0]
        adj_entry = first_ask if first_ask and first_ask > 0 else first_mid
        if adj_entry <= 0:
            adj_entry = sig["premium"]

        first_bid = df["bid"].iloc[0]
        first_ask_val = df["ask"].iloc[0]
        contracts = compute_contracts(adj_entry, score)

        # V6 entry filters
        is_index = ticker in INDEX_TICKERS
        filtered_reason = None
        if not is_index and adj_entry > 5.0:
            filtered_reason = "premium_cap"
        elif (first_bid and first_ask_val and not pd.isna(first_bid)
                and not pd.isna(first_ask_val)
                and first_ask_val > 0 and first_bid > 0):
            spread_pct = (first_ask_val - first_bid) / adj_entry * 100
            if spread_pct > 15:
                filtered_reason = f"spread_{spread_pct:.0f}%"

        v5 = simulate(df, adj_entry, contracts, direction, dte, expiry_date,
                       ticker, v5_settings)

        if filtered_reason:
            v6 = {
                "pnl": 0, "reason": f"FILTERED:{filtered_reason}",
                "hold_min": 0, "peak_gain": 0, "exit_prem": 0,
                "entry_prem": adj_entry, "scaleout_pnl": 0,
                "dca_fired": False, "dca_add": 0,
                "final_contracts": 0, "entry_time": "", "exit_time": "",
            }
        else:
            v6 = simulate(df, adj_entry, contracts, direction, dte, expiry_date,
                           ticker, v6_settings)

        cfg_label = "DEFAULT"
        if ticker in TICKER_CONFIGS:
            cfg_names = {
                "NVDA": "EARLY_PROFIT", "GOOGL": "WIDE_STOP", "TSLA": "LONG_GRACE",
                "IWM": "WIDE_STOP", "QQQ": "LONG_GRACE", "META": "DEFENSIVE",
                "AAPL": "DEFENSIVE", "AMZN": "TIGHT_TRAIL", "AVGO": "EARLY_PROFIT",
                "MSFT": "EARLY_PROFIT", "MSTR": "TIGHT+QUICK",
            }
            cfg_label = cfg_names.get(ticker, "CUSTOM")

        results.append({
            "day": day, "time": sig_time, "ticker": ticker,
            "dir": direction[:4].upper(), "score": score,
            "strike": sig["strike"], "contracts": contracts,
            "entry": adj_entry, "dte": dte,
            "v5_pnl": v5["pnl"], "v5_reason": v5["reason"], "v5_hold": v5["hold_min"],
            "v6_pnl": v6["pnl"], "v6_reason": v6["reason"], "v6_hold": v6["hold_min"],
            "v6_peak": v6["peak_gain"], "v6_scaleout": v6["scaleout_pnl"],
            "v6_dca": v6["dca_fired"], "v6_dca_add": v6["dca_add"],
            "v6_final_ct": v6["final_contracts"],
            "cfg": cfg_label,
            "delta": v6["pnl"] - v5["pnl"],
        })

    harvester_conn.close()

    if not results:
        print("No results to report.")
        return

    # ── Compute summary stats ────────────────────────────────────────────

    total_v5 = sum(r["v5_pnl"] for r in results)
    total_v6 = sum(r["v6_pnl"] for r in results)
    traded = [r for r in results if not r["v6_reason"].startswith("FILTERED")]
    v6_wins = sum(1 for r in traded if r["v6_pnl"] > 0)
    v6_losses = sum(1 for r in traded if r["v6_pnl"] <= 0)
    v5_wins = sum(1 for r in results if r["v5_pnl"] > 0)
    filtered_count = sum(1 for r in results if r["v6_reason"].startswith("FILTERED"))
    dca_count = sum(1 for r in results if r["v6_dca"])

    avg_v6_win = np.mean([r["v6_pnl"] for r in traded if r["v6_pnl"] > 0]) if v6_wins else 0
    avg_v6_loss = np.mean([r["v6_pnl"] for r in traded if r["v6_pnl"] <= 0]) if v6_losses else 0

    # ── Build DOCX ───────────────────────────────────────────────────────

    print("Building .docx...")
    doc = Document()

    # Title
    title = doc.add_heading("V6 Production Build — Trade-by-Trade Backtest", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"Generated {datetime.now().strftime('%B %d, %Y')}\n"
        f"Portfolio: ${PORTFOLIO:,}  |  {len(results)} signals  |  "
        f"{len(traded)} traded  |  {filtered_count} filtered  |  "
        f"{len(no_data)} skipped (no data)"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(89, 89, 89)

    # ── Executive Summary ────────────────────────────────────────────────

    doc.add_heading("Executive Summary", level=1)

    summary_headers = ["Metric", "V5 Baseline", "V6 Production", "Delta"]
    wr_v5 = f"{v5_wins}/{len(results)} ({v5_wins/len(results)*100:.0f}%)" if results else "-"
    wr_v6 = f"{v6_wins}/{len(traded)} ({v6_wins/len(traded)*100:.0f}%)" if traded else "-"
    summary_rows = [
        ["Total P&L", fmt_pnl(total_v5), fmt_pnl(total_v6), fmt_pnl(total_v6 - total_v5)],
        ["Win Rate", wr_v5, wr_v6, ""],
        ["Avg Win", fmt_pnl(np.mean([r["v5_pnl"] for r in results if r["v5_pnl"] > 0]) if v5_wins else 0),
         fmt_pnl(avg_v6_win), ""],
        ["Avg Loss", fmt_pnl(np.mean([r["v5_pnl"] for r in results if r["v5_pnl"] <= 0]) if len(results) - v5_wins else 0),
         fmt_pnl(avg_v6_loss), ""],
        ["DCA Fires", "N/A", str(dca_count), ""],
        ["Filtered", "N/A", str(filtered_count), ""],
    ]
    add_table(doc, summary_headers, summary_rows, pnl_col=[1, 2, 3])

    # ── Daily Summary ────────────────────────────────────────────────────

    doc.add_heading("Daily P&L Summary", level=1)

    daily_headers = ["Date", "Signals", "Filtered", "V5 Day", "V6 Day", "Delta", "V6 Cum", "V6 WR"]
    daily_rows = []
    days = sorted(set(r["day"] for r in results))
    cum_v6 = 0
    for day in days:
        dr = [r for r in results if r["day"] == day]
        dv5 = sum(r["v5_pnl"] for r in dr)
        dv6 = sum(r["v6_pnl"] for r in dr)
        cum_v6 += dv6
        filt = sum(1 for r in dr if r["v6_reason"].startswith("FILTERED"))
        dt = [r for r in dr if not r["v6_reason"].startswith("FILTERED")]
        dw = sum(1 for r in dt if r["v6_pnl"] > 0)
        wr = f"{dw}/{len(dt)}" if dt else "-"
        daily_rows.append([
            day, str(len(dr)), str(filt),
            fmt_pnl(dv5), fmt_pnl(dv6), fmt_pnl(dv6 - dv5),
            fmt_pnl(cum_v6), wr,
        ])
    add_table(doc, daily_headers, daily_rows, pnl_col=[3, 4, 5, 6])

    # ── Per-Day Trade Tables ─────────────────────────────────────────────

    doc.add_heading("Trade-by-Trade Detail", level=1)

    trade_headers = [
        "#", "Time", "Ticker", "Dir", "Score", "Strike", "Ct",
        "Entry", "DTE", "V5 P&L", "V5 Exit", "V6 P&L", "V6 Exit",
        "Hold", "Peak%", "Notes",
    ]

    cum_v6 = 0
    for day in days:
        dr = [r for r in results if r["day"] == day]
        dv6 = sum(r["v6_pnl"] for r in dr)
        cum_v6 += dv6
        dt = [r for r in dr if not r["v6_reason"].startswith("FILTERED")]
        dw = sum(1 for r in dt if r["v6_pnl"] > 0)
        wr = f"{dw}/{len(dt)}" if dt else "-"

        h = doc.add_heading(f"{day}  —  V6: {fmt_pnl(dv6)}  |  WR: {wr}  |  Cum: {fmt_pnl(cum_v6)}", level=2)

        trade_rows = []
        for i, r in enumerate(dr, 1):
            notes = []
            if r["v6_reason"].startswith("FILTERED"):
                notes.append(r["v6_reason"].replace("FILTERED:", ""))
            if r["v6_dca"]:
                notes.append(f"DCA+{r['v6_dca_add']}")
            if r["v6_scaleout"] != 0:
                notes.append(f"SO:{fmt_pnl(r['v6_scaleout'])}")
            if r["cfg"] != "DEFAULT":
                notes.append(r["cfg"])

            trade_rows.append([
                str(i),
                r["time"],
                r["ticker"],
                r["dir"],
                str(r["score"]),
                f"${r['strike']:.2f}",
                str(r["contracts"]),
                f"${r['entry']:.2f}",
                str(r["dte"]),
                fmt_pnl(r["v5_pnl"]),
                r["v5_reason"][:15],
                fmt_pnl(r["v6_pnl"]),
                r["v6_reason"][:15],
                f"{r['v6_hold']:.0f}m" if r["v6_hold"] > 0 else "-",
                f"{r['v6_peak']:.0f}%" if r["v6_peak"] != 0 else "-",
                " ".join(notes),
            ])

        add_table(doc, trade_headers, trade_rows, pnl_col=[9, 11])
        doc.add_paragraph()  # spacing

    # ── Per-Ticker Summary ───────────────────────────────────────────────

    doc.add_heading("Per-Ticker Summary", level=1)

    ticker_headers = ["Ticker", "Config", "N", "Filt", "DCA", "V5 P&L", "V6 P&L", "Delta", "V6 WR"]
    ticker_rows = []
    tickers = sorted(set(r["ticker"] for r in results))
    for t in tickers:
        tr = [r for r in results if r["ticker"] == t]
        t_traded = [r for r in tr if not r["v6_reason"].startswith("FILTERED")]
        t_v5 = sum(r["v5_pnl"] for r in tr)
        t_v6 = sum(r["v6_pnl"] for r in tr)
        t_filt = sum(1 for r in tr if r["v6_reason"].startswith("FILTERED"))
        t_dca = sum(1 for r in tr if r["v6_dca"])
        t_w = sum(1 for r in t_traded if r["v6_pnl"] > 0)
        wr = f"{t_w}/{len(t_traded)}" if t_traded else "-"
        cfg = "DEFAULT"
        if t in TICKER_CONFIGS:
            cfg_names = {
                "NVDA": "EARLY_PROFIT", "GOOGL": "WIDE_STOP", "TSLA": "LONG_GRACE",
                "IWM": "WIDE_STOP", "QQQ": "LONG_GRACE", "META": "DEFENSIVE",
                "AAPL": "DEFENSIVE", "AMZN": "TIGHT_TRAIL", "AVGO": "EARLY_PROFIT",
                "MSFT": "EARLY_PROFIT", "MSTR": "TIGHT+QUICK",
            }
            cfg = cfg_names.get(t, "CUSTOM")
        ticker_rows.append([
            t, cfg, str(len(tr)), str(t_filt), str(t_dca),
            fmt_pnl(t_v5), fmt_pnl(t_v6), fmt_pnl(t_v6 - t_v5), wr,
        ])
    add_table(doc, ticker_headers, ticker_rows, pnl_col=[5, 6, 7])

    # ── Exit Reason Breakdown ────────────────────────────────────────────

    doc.add_heading("V6 Exit Reason Breakdown", level=1)

    reason_headers = ["Exit Reason", "Count", "Wins", "Losses", "Win Rate", "Total P&L", "Avg P&L", "Avg Hold"]
    reason_rows = []
    df_all = pd.DataFrame(results)
    for reason, group in df_all.groupby("v6_reason"):
        cnt = len(group)
        wins = int((group["v6_pnl"] > 0).sum())
        losses = cnt - wins
        wr = f"{wins/cnt*100:.0f}%" if cnt else "-"
        total = group["v6_pnl"].sum()
        avg = group["v6_pnl"].mean()
        avg_hold = group["v6_hold"].mean()
        reason_rows.append([
            reason, str(cnt), str(wins), str(losses), wr,
            fmt_pnl(total), fmt_pnl(avg), f"{avg_hold:.0f}m",
        ])
    add_table(doc, reason_headers, reason_rows, pnl_col=[5, 6])

    # ── V6 Feature Impact ────────────────────────────────────────────────

    doc.add_heading("V6 Feature Impact", level=1)

    # DCA
    dca_trades = [r for r in results if r["v6_dca"]]
    if dca_trades:
        dca_v5 = sum(r["v5_pnl"] for r in dca_trades)
        dca_v6 = sum(r["v6_pnl"] for r in dca_trades)
        doc.add_heading("DCA (Dollar Cost Averaging)", level=2)
        p = doc.add_paragraph()
        p.add_run(f"{len(dca_trades)} DCA fires  |  ").bold = False
        p.add_run(f"V5 P&L: {fmt_pnl(dca_v5)}  |  V6 P&L: {fmt_pnl(dca_v6)}  |  ")
        r = p.add_run(f"Benefit: {fmt_pnl(dca_v6 - dca_v5)}")
        r.bold = True

    # Per-ticker configs
    custom_cfg = [r for r in results if r["cfg"] != "DEFAULT" and not r["v6_reason"].startswith("FILTERED")]
    if custom_cfg:
        cv5 = sum(r["v5_pnl"] for r in custom_cfg)
        cv6 = sum(r["v6_pnl"] for r in custom_cfg)
        doc.add_heading("Per-Ticker Configs", level=2)
        p = doc.add_paragraph()
        p.add_run(f"{len(custom_cfg)} trades with custom config  |  ")
        p.add_run(f"V5 P&L: {fmt_pnl(cv5)}  |  V6 P&L: {fmt_pnl(cv6)}  |  ")
        r = p.add_run(f"Benefit: {fmt_pnl(cv6 - cv5)}")
        r.bold = True

    # Filters
    filt_trades = [r for r in results if r["v6_reason"].startswith("FILTERED")]
    if filt_trades:
        fv5 = sum(r["v5_pnl"] for r in filt_trades)
        doc.add_heading("Entry Filters", level=2)
        p = doc.add_paragraph()
        p.add_run(f"{len(filt_trades)} blocked  |  ")
        p.add_run(f"V5 would-have P&L: {fmt_pnl(fv5)}  |  ")
        r = p.add_run(f"Filter benefit: {fmt_pnl(-fv5)}")
        r.bold = True

    # Scale-out
    so_trades = [r for r in results if r["v6_scaleout"] != 0]
    if so_trades:
        doc.add_heading("Scale-Out", level=2)
        p = doc.add_paragraph()
        p.add_run(f"{len(so_trades)} fires  |  ")
        r = p.add_run(f"Locked profit: {fmt_pnl(sum(r['v6_scaleout'] for r in so_trades))}")
        r.bold = True

    # ── Skipped signals ──────────────────────────────────────────────────

    if no_data:
        doc.add_heading("Skipped Signals (No Harvester Data)", level=1)
        skip_headers = ["Date", "Time", "Ticker", "Dir", "Score", "Strike", "Premium"]
        skip_rows = [
            [s["day"], s["time"], s["ticker"], s["dir"], str(s["score"]),
             f"${s['strike']:.2f}", f"${s['premium']:.2f}"]
            for s in no_data
        ]
        add_table(doc, skip_headers, skip_rows)

    # ── Save ─────────────────────────────────────────────────────────────

    REPORTS_DIR.mkdir(exist_ok=True)
    out_path = REPORTS_DIR / f"V6_Trade_Report_{datetime.now().strftime('%Y-%m-%d')}.docx"
    doc.save(str(out_path))
    print(f"\nSaved: {out_path}")
    print(f"  {len(results)} signals  |  V5: {fmt_pnl(total_v5)}  |  V6: {fmt_pnl(total_v6)}  |  Delta: {fmt_pnl(total_v6 - total_v5)}")


if __name__ == "__main__":
    main()
